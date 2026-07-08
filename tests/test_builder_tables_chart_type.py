"""MNT-33 — the legacy `table` CHART TYPE must render as a native Word table.

MNT-30 made the dedicated ``tables:`` config section render a native
``w:tbl``. But a ``charts:`` entry with ``type: table`` still flowed through
the chart engine (``_generate_charts`` -> ``generate_chart`` -> matplotlib PNG
-> ``InlineImage``), producing an unsearchable, non-editable, screen-reader-
invisible image table -- a second, UI-exposed way to make exactly the image
tables MNT-30 set out to eliminate.

Acceptance criteria encoded here (from docs/ROADMAP.md, MNT-33):
  AC1 -- A ``charts:`` entry with ``type: table`` renders in the built ``.docx``
        as a native python-docx table (``document.tables`` gains a ``w:tbl``)
        with borders and NO InlineImage / embedded PNG for it.
  AC2 -- Real chart types (bar/pie/line/...) still render as InlineImage PNGs
        (no chart regression).
  AC4 -- Loading a legacy config that contains a ``type: table`` chart does not
        error and produces a native table (backward compatible, no migration).

These tests build a report from a ``charts:`` config (NOT a ``tables:`` recipe)
so they exercise the legacy chart-type path specifically.
"""
import zipfile
from pathlib import Path

import pandas as pd
import pytest
import yaml
from docx import Document
from docx.oxml.ns import qn

from src.reports.builder import ReportBuilder


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _embedded_media(path: Path) -> list[str]:
    """Names of embedded binary media parts (e.g. word/media/image1.png)."""
    with zipfile.ZipFile(path) as z:
        return [n for n in z.namelist() if n.startswith("word/media/")]


def _make_template(path: Path, placeholders: list[str]) -> None:
    """A .docx whose content is one single-run paragraph per placeholder."""
    doc = Document()
    for ph in placeholders:
        p = doc.add_paragraph()
        p.add_run(ph)
    doc.save(str(path))


def _base_cfg(ws: Path, template_path: Path, alias: str) -> dict:
    return {
        "api": {
            "url": "https://kf.kobotoolbox.org/api/v2",
            "token": "dummy",
            "platform": "kobo",
        },
        "form": {"alias": alias, "uid": "x"},
        "questions": [
            {
                "kobo_key": "region",
                "label": "Region",
                "type": "select_one",
                "category": "categorical",
                "group": "",
                "export_label": "Region",
            },
            {
                "kobo_key": "age",
                "label": "Age",
                "type": "integer",
                "category": "quantitative",
                "group": "",
                "export_label": "Age",
            },
        ],
        "filters": [],
        "charts": [],
        "tables": [],
        "report": {
            "template": str(template_path),
            "output_dir": str(ws / "reports"),
            "title": "MNT-33 legacy table chart type",
            "period": "Q2 2026",
        },
        "export": {
            "format": "csv",
            "output_dir": str(ws / "data" / "processed"),
        },
    }


def _write_data(ws: Path, alias: str) -> None:
    csv_path = ws / "data" / "processed" / f"{alias}_data_20260101_120000.csv"
    pd.DataFrame(
        {
            "Region": ["Nairobi", "Mombasa", "Nairobi", "Kisumu"],
            "Age": [10, 20, 30, 40],
        }
    ).to_csv(csv_path, index=False)


@pytest.fixture
def chart_table_workspace(tmp_path, monkeypatch):
    """Factory: a workspace whose ``charts:`` section contains a ``type: table``
    entry (and optionally a sibling ``type: bar`` chart), with a template that
    references the matching placeholders. Ready to build."""

    def _make(*, include_bar=False):
        alias = "mnt33" + ("b" if include_bar else "")
        ws = tmp_path / f"ws_{alias}"
        (ws / "data" / "processed").mkdir(parents=True)
        (ws / "templates").mkdir()
        (ws / "reports").mkdir()

        _write_data(ws, alias)
        template_path = ws / "templates" / "t.docx"

        table_chart_name = "region_table"
        bar_chart_name = "region_bar"

        cfg = _base_cfg(ws, template_path, alias)
        charts = [
            {
                "name": table_chart_name,
                "title": "By Region",
                "type": "table",
                "questions": ["Region", "Age"],
                "options": {},
            }
        ]
        # A legacy config referenced the table chart via the chart placeholder
        # `{{ chart_<name> }}`. The native-table bridge must resolve it to a
        # w:tbl regardless of which placeholder carries it, so seed both the
        # legacy chart placeholder AND the native table sentinel in the
        # template -- exactly one must materialise as a table.
        placeholders = [
            "{{ chart_%s }}" % table_chart_name,
            "{{ table_%s }}" % table_chart_name,
        ]
        if include_bar:
            charts.append(
                {
                    "name": bar_chart_name,
                    "title": "Region",
                    "type": "bar",
                    "questions": ["Region"],
                    "options": {},
                }
            )
            placeholders.append("{{ chart_%s }}" % bar_chart_name)

        cfg["charts"] = charts
        _make_template(template_path, placeholders)

        (ws / "config.yml").write_text(yaml.dump(cfg, allow_unicode=True))
        monkeypatch.chdir(ws)
        return {
            "ws": ws,
            "cfg": cfg,
            "out_dir": ws / "reports",
            "table_chart_name": table_chart_name,
            "bar_chart_name": bar_chart_name,
        }

    return _make


def _first_output(out_dir: Path) -> Path:
    docs = list(out_dir.glob("*.docx"))
    assert docs, "build() produced no .docx output"
    return docs[0]


def _table_has_visible_borders(tbl) -> bool:
    """True if the table declares borders -- either a Grid table-style reference
    or a direct ``w:tblBorders`` element in tblPr."""
    tblPr = tbl._tbl.tblPr
    if tblPr is not None:
        if tblPr.find(qn("w:tblBorders")) is not None:
            return True
        style_el = tblPr.find(qn("w:tblStyle"))
        if style_el is not None:
            val = style_el.get(qn("w:val")) or ""
            if "Grid" in val:
                return True
    try:
        if tbl.style is not None and "Grid" in (tbl.style.name or ""):
            return True
    except Exception:
        pass
    return False


# ---------------------------------------------------------------------------
# AC1 -- a `type: table` CHART renders as a native w:tbl, not an image
# ---------------------------------------------------------------------------

def test_table_chart_type_renders_as_native_table(chart_table_workspace):
    """AC1: a ``charts:`` entry with ``type: table`` must resolve to a native
    Word table -- python-docx ``document.tables`` contains at least one w:tbl.
    Currently the table chart type routes through the PNG chart engine, so no
    native table is produced and this assertion fails."""
    workspace = chart_table_workspace()
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert len(doc.tables) >= 1, (
        "A `charts:` entry with type: table produced no native Word table "
        "(document.tables is empty). The legacy table chart type must be routed "
        "to the native-table path, not rendered as an InlineImage PNG."
    )


def test_table_chart_type_emits_no_embedded_image(chart_table_workspace):
    """AC1: with ONLY a ``type: table`` chart (no real chart), the output .docx
    must embed no image media -- proving the table is native text, not a
    flattened matplotlib PNG."""
    workspace = chart_table_workspace()  # no bar chart
    ReportBuilder(workspace["cfg"]).build()

    media = _embedded_media(_first_output(workspace["out_dir"]))
    assert media == [], (
        "The report embeds image media even though its only chart is a "
        f"type: table entry: {media}. A table chart must render as a native "
        "python-docx table, not an InlineImage PNG."
    )


def test_table_chart_type_native_table_has_borders(chart_table_workspace):
    """AC1: the native table produced from a ``type: table`` chart must have
    visible borders (Table Grid style or a manual w:tblBorders element)."""
    workspace = chart_table_workspace()
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert doc.tables, "no native table found for the type: table chart"
    assert _table_has_visible_borders(doc.tables[0]), (
        "The native table produced from a type: table chart has no visible "
        "borders. It must apply Table Grid (or a manual w:tblBorders element)."
    )


def test_table_chart_type_populates_group_values_as_text(chart_table_workspace):
    """AC1: the native table must carry the tabulated data as selectable text
    cells -- the Region group values must appear as cell text, not be drawn
    into a figure."""
    workspace = chart_table_workspace()
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert doc.tables, "no native table found for the type: table chart"
    all_cell_text = {
        c.text.strip() for t in doc.tables for r in t.rows for c in r.cells
    }
    for value in ("Nairobi", "Mombasa", "Kisumu"):
        assert value in all_cell_text, (
            f"Region value {value!r} not found as a text cell in the native "
            f"table. Cell texts were: {sorted(all_cell_text)}. The table chart "
            "data must be populated as text cells, not rendered as an image."
        )


# ---------------------------------------------------------------------------
# AC2 -- no chart regression: a real chart type is still an embedded PNG
# ---------------------------------------------------------------------------

def test_bar_chart_still_embedded_as_image(chart_table_workspace):
    """AC2: a sibling ``type: bar`` chart in the same config still renders as an
    InlineImage / embedded PNG (charts stay images -- python-docx has no chart
    API; explicitly out of scope). Regression guard: routing the table type to
    native must not affect real chart types."""
    workspace = chart_table_workspace(include_bar=True)
    ReportBuilder(workspace["cfg"]).build()

    out = _first_output(workspace["out_dir"])
    media = _embedded_media(out)
    assert media, (
        "The report embeds no image media even though it contains a bar chart. "
        "Real chart types must still render as InlineImage PNGs -- only the "
        "table type changes to a native table."
    )

    # And the table chart must still materialise as a native table alongside it.
    doc = Document(str(out))
    assert len(doc.tables) >= 1, (
        "The type: table chart did not render as a native table when a real "
        "bar chart was also present in the report."
    )


# ---------------------------------------------------------------------------
# AC4 -- legacy config with a type: table chart loads without error -> native
# ---------------------------------------------------------------------------

def test_legacy_table_chart_config_builds_without_error(chart_table_workspace):
    """AC4: building a legacy config that contains a ``type: table`` chart must
    not raise and must produce a native table -- backward compatible, no
    migration step required of the user."""
    workspace = chart_table_workspace()
    outputs = ReportBuilder(workspace["cfg"]).build()

    assert outputs, "build() returned no output for a legacy type: table config"
    doc = Document(str(_first_output(workspace["out_dir"])))
    assert len(doc.tables) >= 1, (
        "A legacy config with a type: table chart built without error but "
        "produced no native table. The bridge must render it native with no "
        "config migration."
    )
