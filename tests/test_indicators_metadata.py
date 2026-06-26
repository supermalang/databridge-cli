"""ME-3 — Indicator metadata catalog.

Contract under test (derived from the ME-3 acceptance criteria):

1. Indicators accept four NEW optional metadata fields — `unit`, `source`,
   `frequency`, `responsible`. Carrying them must not break indicator
   computation (the fields are accepted/optional).

   NOTE: `source` already exists on indicators as a data-source selector
   ("main" / a repeat-group path). ME-3 reuses it as the donor-style
   "data source / means of verification" metadata column. A free-text value
   that is not a known repeat table must fail-soft to the main df (it already
   does) and therefore must NOT break computation.

2. `generate-template` emits an indicator reference annex built from those
   metadata fields — for EACH indicator the generated template surfaces its
   `unit`, `source`, `frequency` and `responsible` values.

3. Indicators WITHOUT the metadata fields still generate without error
   (graceful absence — the annex degrades cleanly).
"""
import pandas as pd
from docx import Document

from src.reports.indicators import compute_indicators
from src.reports.template_generator import generate_template


def _doc_text(path):
    """All visible text in the generated .docx — paragraphs AND table cells.

    The annex may be emitted as paragraphs or as a table; capture both so the
    test asserts the *requirement* (the metadata is surfaced) without pinning
    the implementer to one layout.
    """
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _df():
    return pd.DataFrame({"Doses": [10, 20, 30, 40]})


def _indicator_with_metadata():
    return {
        "name": "doses",
        "label": "Doses administered",
        "stat": "sum",
        "question": "Doses",
        "unit": "doses",
        "source": "Vaccination register",
        "frequency": "Monthly",
        "responsible": "Health Officer",
    }


# --- AC 1: metadata fields accepted (optional, do not break computation) -----

def test_indicator_with_metadata_fields_computes_without_error():
    ctx = compute_indicators([_indicator_with_metadata()], _df())
    # The indicator still computes its scalar normally — fields are accepted.
    assert ctx["ind_doses"] == "100"


def test_metadata_fields_do_not_alter_computed_value():
    with_meta = compute_indicators([_indicator_with_metadata()], _df())
    plain = compute_indicators(
        [{"name": "doses", "label": "Doses administered", "stat": "sum", "question": "Doses"}],
        _df(),
    )
    # Adding metadata must not change the computed indicator value.
    assert with_meta["ind_doses"] == plain["ind_doses"]


# --- AC 2: generate-template emits an indicator reference annex --------------

def test_template_emits_metadata_for_each_indicator(tmp_path):
    cfg = {
        "charts": [],
        "indicators": [
            {
                "name": "doses",
                "label": "Doses administered",
                "unit": "doses",
                "source": "Vaccination register",
                "frequency": "Monthly",
                "responsible": "Health Officer",
            },
            {
                "name": "coverage",
                "label": "Coverage rate",
                "unit": "%",
                "source": "DHIS2 export",
                "frequency": "Quarterly",
                "responsible": "M&E Lead",
            },
        ],
    }
    out = tmp_path / "tpl.docx"
    generate_template(cfg, out)
    text = _doc_text(out)

    # Each indicator's four metadata values must appear in the generated annex.
    for ind in cfg["indicators"]:
        assert ind["unit"] in text, f"missing unit for {ind['name']}"
        assert ind["source"] in text, f"missing source for {ind['name']}"
        assert ind["frequency"] in text, f"missing frequency for {ind['name']}"
        assert ind["responsible"] in text, f"missing responsible for {ind['name']}"


def test_template_annex_labels_the_metadata_columns(tmp_path):
    cfg = {
        "charts": [],
        "indicators": [_indicator_with_metadata()],
    }
    out = tmp_path / "tpl.docx"
    generate_template(cfg, out)
    text = _doc_text(out).lower()

    # The annex names the four metadata dimensions so the reference is readable.
    assert "unit" in text
    assert "source" in text
    assert "frequency" in text
    assert "responsible" in text


# --- AC 3: graceful absence of metadata fields ------------------------------

def test_template_generates_without_error_when_metadata_absent(tmp_path):
    cfg = {
        "charts": [],
        "indicators": [
            {"name": "doses", "label": "Doses administered", "stat": "sum", "question": "Doses"},
        ],
    }
    out = tmp_path / "tpl.docx"
    # Must not raise when indicators carry no metadata fields.
    result = generate_template(cfg, out)
    assert result.exists()
    text = _doc_text(out)
    # The indicator scalar placeholder is still emitted (existing behaviour intact).
    assert "{{ ind_doses }}" in text
