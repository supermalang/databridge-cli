"""MNT-30 — Report tables must render as native Word tables, not PNG images.

`{{ table_<name> }}` recipe placeholders currently render as a flattened
InlineImage produced by the `table` chart type. Tabular output must instead be
a native python-docx table (a ``w:tbl`` element): selectable, editable,
accessible text — with visible borders. Charts (bar/pie/line/etc.) stay as
image renders (python-docx has no chart API), which is explicitly out of scope.

Acceptance criteria encoded here:
  AC1 — A `{{ table_<name> }}` recipe renders as a native Word table: the output
         .docx contains a ``w:tbl`` for that placeholder (visible in
         python-docx ``document.tables``) and NO embedded image for it.
  AC2 — The table has one header row (column labels) plus one row per record,
         populated as text cells.
  AC3 — The table has visible borders — the ``Table Grid`` style when present in
         the template, otherwise a manually-applied ``w:tblBorders`` element.
         Covered BOTH ways (style present / style absent → manual fallback).
  AC4 — No regression to charts: a `bar` chart in the same report is still
         inserted as an InlineImage / embedded PNG.
  AC5 — No regression to table data resolution: one header row + one row per
         record populated as native cells (same rows/columns, now as text).
"""
import zipfile
from pathlib import Path

import pandas as pd
import pytest
import yaml
from docx import Document
from docx.oxml.ns import qn

from src.reports.builder import ReportBuilder


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _embedded_media(path: Path) -> list[str]:
    """Names of embedded binary media parts (e.g. word/media/image1.png)."""
    with zipfile.ZipFile(path) as z:
        return [n for n in z.namelist() if n.startswith("word/media/")]


def _make_table_only_template(path: Path, table_name: str, *, with_table_grid: bool) -> None:
    """A .docx whose only content is a single-run ``{{ table_<name> }}``
    placeholder. When ``with_table_grid`` is False the ``Table Grid`` style is
    stripped from the template's styles part so the implementation must fall
    back to a manually-applied ``w:tblBorders`` element for borders."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{ table_%s }}" % table_name)

    if not with_table_grid:
        styles = doc.styles
        # python-docx default template ships a 'Table Grid' style; remove it so
        # the manual-borders fallback path is exercised.
        if "Table Grid" in [s.name for s in styles]:
            el = styles["Table Grid"].element
            el.getparent().remove(el)
    doc.save(str(path))


def _make_table_and_chart_template(path: Path, table_name: str, chart_name: str) -> None:
    """A .docx containing a ``{{ table_<name> }}`` placeholder AND a
    ``{{ chart_<name> }}`` placeholder, each in its own single-run paragraph."""
    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("{{ table_%s }}" % table_name)
    p2 = doc.add_paragraph()
    p2.add_run("{{ chart_%s }}" % chart_name)
    doc.save(str(path))


def _base_cfg(ws: Path, template_path: Path, alias: str) -> dict:
    return {
        "api": {
            "url": "https://kf.kobotoolbox.org/api/v2",
            "token": "dummy",
            "platform": "kobo",
        },
        "form": {"alias": alias, "uid": "x"},
        "questions": [
            {
                "kobo_key": "region",
                "label": "Region",
                "type": "select_one",
                "category": "categorical",
                "group": "",
                "export_label": "Region",
            },
            {
                "kobo_key": "age",
                "label": "Age",
                "type": "integer",
                "category": "quantitative",
                "group": "",
                "export_label": "Age",
            },
        ],
        "filters": [],
        "charts": [],
        "tables": [],
        "report": {
            "template": str(template_path),
            "output_dir": str(ws / "reports"),
            "title": "MNT-30 Native Tables",
            "period": "Q2 2026",
        },
        "export": {
            "format": "csv",
            "output_dir": str(ws / "data" / "processed"),
        },
    }


def _write_data(ws: Path, alias: str) -> None:
    csv_path = ws / "data" / "processed" / f"{alias}_data_20260101_120000.csv"
    pd.DataFrame(
        {
            "Region": ["Nairobi", "Mombasa", "Nairobi", "Kisumu"],
            "Age": [10, 20, 30, 40],
        }
    ).to_csv(csv_path, index=False)


@pytest.fixture
def table_workspace(tmp_path, monkeypatch):
    """Factory: workspace with one `tables:` recipe (and optional `bar` chart),
    a template controlled by the caller, ready to build."""

    def _make(*, with_table_grid=True, include_chart=False):
        alias = "mnt30" + ("g" if with_table_grid else "n") + ("c" if include_chart else "")
        ws = tmp_path / f"ws_{alias}"
        (ws / "data" / "processed").mkdir(parents=True)
        (ws / "templates").mkdir()
        (ws / "reports").mkdir()

        _write_data(ws, alias)
        template_path = ws / "templates" / "t.docx"

        table_name = "region_breakdown"
        chart_name = "region_bar"
        if include_chart:
            _make_table_and_chart_template(template_path, table_name, chart_name)
        else:
            _make_table_only_template(
                template_path, table_name, with_table_grid=with_table_grid
            )

        cfg = _base_cfg(ws, template_path, alias)
        cfg["tables"] = [
            {"name": table_name, "title": "By Region", "questions": ["Region"]}
        ]
        if include_chart:
            cfg["charts"] = [
                {
                    "name": chart_name,
                    "title": "Region",
                    "type": "bar",
                    "questions": ["Region"],
                    "options": {},
                }
            ]

        (ws / "config.yml").write_text(yaml.dump(cfg, allow_unicode=True))
        monkeypatch.chdir(ws)
        return {
            "ws": ws,
            "cfg": cfg,
            "out_dir": ws / "reports",
            "table_name": table_name,
            "chart_name": chart_name,
        }

    return _make


def _first_output(out_dir: Path) -> Path:
    docs = list(out_dir.glob("*.docx"))
    assert docs, "build() produced no .docx output"
    return docs[0]


# ---------------------------------------------------------------------------
# AC1 — recipe placeholder resolves to a native w:tbl (not an inline image)
# ---------------------------------------------------------------------------

def test_recipe_placeholder_renders_as_native_table(table_workspace):
    """AC1: the `{{ table_<name> }}` recipe placeholder resolves to a native
    Word table — python-docx ``document.tables`` contains at least one table.
    The template has no other tabular content, so any table present is the
    recipe's."""
    workspace = table_workspace()
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert len(doc.tables) >= 1, (
        "The report contains no native Word table (document.tables is empty). "
        "The `{{ table_<name> }}` recipe must be substituted with a python-docx "
        "w:tbl, not rendered as an InlineImage PNG."
    )


def test_table_recipe_produces_no_embedded_image(table_workspace):
    """AC1: with a `tables:` recipe and NO charts, the output .docx must embed
    no image media — proving the table is text, not a flattened PNG figure."""
    workspace = table_workspace()  # no chart in this workspace
    ReportBuilder(workspace["cfg"]).build()

    media = _embedded_media(_first_output(workspace["out_dir"]))
    assert media == [], (
        "The report embeds image media even though it contains only a "
        f"`tables:` recipe (no charts): {media}. A native table must not emit "
        "an InlineImage — the table PNG render path is the bug being fixed."
    )


# ---------------------------------------------------------------------------
# AC2 / AC5 — header row + one row per record, populated as text cells
# ---------------------------------------------------------------------------

def test_table_has_header_plus_one_row_per_record(table_workspace):
    """AC2/AC5: the native table has exactly one header row plus one row per
    aggregated record. The recipe groups the 4-row dataset by Region into 3
    distinct groups (Nairobi, Mombasa, Kisumu) → 1 header + 3 data rows."""
    workspace = table_workspace()
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert doc.tables, "no native table found"
    tbl = doc.tables[0]

    n_rows = len(tbl.rows)
    # 3 distinct Region groups + 1 header row.
    assert n_rows == 4, (
        f"Expected 1 header row + 1 row per record (3 Region groups) = 4 rows, "
        f"got {n_rows}. The table must be built one header row plus one row per "
        "resolved record."
    )

    # The three group values must appear as selectable text cells.
    all_cell_text = {c.text.strip() for r in tbl.rows for c in r.cells}
    for value in ("Nairobi", "Mombasa", "Kisumu"):
        assert value in all_cell_text, (
            f"Group value {value!r} not found as a text cell in the native "
            f"table. Cell texts were: {sorted(all_cell_text)}. Data rows must "
            "be populated as text cells, not drawn as a figure."
        )


def test_table_header_row_carries_column_labels(table_workspace):
    """AC2: the first row is a header row carrying the field/column labels."""
    workspace = table_workspace()
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert doc.tables, "no native table found"
    header_cells = [c.text.strip() for c in doc.tables[0].rows[0].cells]

    assert any(h for h in header_cells), (
        f"Header row cells are all empty: {header_cells!r}. The first table row "
        "must carry the field/column labels as text."
    )
    assert "Region" in header_cells, (
        f"Expected the header row to include the 'Region' column label, got "
        f"{header_cells!r}."
    )


# ---------------------------------------------------------------------------
# AC3 — visible borders, BOTH via Table Grid style and manual w:tblBorders
# ---------------------------------------------------------------------------

def _table_has_visible_borders(tbl) -> bool:
    """True if the table declares borders — either via a 'Grid'/'Table Grid'
    table style reference or a direct ``w:tblBorders`` element in tblPr."""
    tblPr = tbl._tbl.tblPr
    if tblPr is not None:
        if tblPr.find(qn("w:tblBorders")) is not None:
            return True
        style_el = tblPr.find(qn("w:tblStyle"))
        if style_el is not None:
            val = style_el.get(qn("w:val")) or ""
            if "Grid" in val or "Table Grid" in val:
                return True
    # python-docx surfaces the applied style name too.
    try:
        if tbl.style is not None and "Grid" in (tbl.style.name or ""):
            return True
    except Exception:
        pass
    return False


def test_table_borders_present_with_table_grid_style(table_workspace):
    """AC3 (style path): when the template defines the 'Table Grid' style, the
    native table must be bordered (Table Grid style applied)."""
    workspace = table_workspace(with_table_grid=True)
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert doc.tables, "no native table found"
    assert _table_has_visible_borders(doc.tables[0]), (
        "The native table has no visible borders even though the template "
        "defines the 'Table Grid' style. The implementation must apply "
        "'Table Grid' (or an equivalent bordered style) when it exists."
    )


def test_table_borders_present_without_table_grid_style(table_workspace):
    """AC3 (manual-fallback path): when the template does NOT define the
    'Table Grid' style, the native table must still be bordered via a
    manually-applied ``w:tblBorders`` element (add_table() adds none by
    default)."""
    workspace = table_workspace(with_table_grid=False)
    ReportBuilder(workspace["cfg"]).build()

    doc = Document(str(_first_output(workspace["out_dir"])))
    assert doc.tables, "no native table found"
    tbl = doc.tables[0]

    tblPr = tbl._tbl.tblPr
    has_manual_borders = (
        tblPr is not None and tblPr.find(qn("w:tblBorders")) is not None
    )
    assert has_manual_borders, (
        "When the 'Table Grid' style is absent from the template, the native "
        "table must carry a manually-applied w:tblBorders element (add_table() "
        "adds no borders by default). None was found — the manual-borders "
        "fallback path is missing."
    )


# ---------------------------------------------------------------------------
# AC4 — no chart regression: a bar chart is still an embedded InlineImage PNG
# ---------------------------------------------------------------------------

def test_bar_chart_still_embedded_as_image(table_workspace):
    """AC4: a `bar` chart in the same report is still inserted as an
    InlineImage / embedded PNG (charts stay images — out of scope)."""
    workspace = table_workspace(include_chart=True)
    ReportBuilder(workspace["cfg"]).build()

    out = _first_output(workspace["out_dir"])
    media = _embedded_media(out)
    assert media, (
        "The report embeds no image media even though it contains a `bar` "
        "chart. Charts must still render as InlineImage PNGs — only tables "
        "change to native tables."
    )

    # And the native table must still be present alongside the chart image.
    doc = Document(str(out))
    assert len(doc.tables) >= 1, (
        "The `{{ table_<name> }}` recipe placeholder did not render as a native "
        "table when a chart was also present in the report."
    )
