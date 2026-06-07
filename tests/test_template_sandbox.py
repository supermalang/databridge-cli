"""Security: Word templates are user-uploaded and rendered with Jinja2 via
docxtpl. Rendering must use a sandboxed environment so a malicious template
cannot reach Python internals and execute code (SSTI -> RCE). Audit finding #2."""
from pathlib import Path

import pytest
from docx import Document
from docxtpl import DocxTemplate
from jinja2.exceptions import SecurityError

from src.reports.builder import sandboxed_jinja_env


def _make_docx(text: str, path: Path) -> Path:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_legitimate_placeholder_renders(tmp_path):
    tpl_path = _make_docx("Title: {{ report_title }}", tmp_path / "ok.docx")
    tpl = DocxTemplate(str(tpl_path))
    tpl.render({"report_title": "Quarterly Report"}, jinja_env=sandboxed_jinja_env())
    out = tmp_path / "ok_rendered.docx"
    tpl.save(out)
    rendered = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Quarterly Report" in rendered


def test_dunder_gadget_chain_is_blocked(tmp_path):
    """The classic SSTI gadget chain must raise SecurityError, not execute."""
    payload = "{{ ''.__class__.__mro__[1].__subclasses__() }}"
    tpl_path = _make_docx(payload, tmp_path / "evil.docx")
    tpl = DocxTemplate(str(tpl_path))
    with pytest.raises(SecurityError):
        tpl.render({}, jinja_env=sandboxed_jinja_env())


def test_code_execution_payload_does_not_run(tmp_path):
    """A payload that would write a file must not write it."""
    marker = tmp_path / "pwned.txt"
    payload = (
        "{{ cycler.__init__.__globals__.os.system('touch %s') }}" % marker
    )
    tpl_path = _make_docx(payload, tmp_path / "rce.docx")
    tpl = DocxTemplate(str(tpl_path))
    with pytest.raises(SecurityError):
        tpl.render({}, jinja_env=sandboxed_jinja_env())
    assert not marker.exists()
