"""Unit tests for XTF-1 — placeholder extraction from .docx.

These tests are the spec for ``extract_placeholders(docx_path) -> List[Token]``
(``src/reports/template_inference.py``), derived strictly from the XTF-1
acceptance criteria and design spec §4.1. They are written before the
implementation exists and are expected to be RED until it lands.

Contract committed to here (minimal, AC-derived):

``extract_placeholders(docx_path)`` returns a list of ``Token`` objects, in
document order, one per matched placeholder. Each ``Token`` exposes:

* ``raw``        -- the full delimited string, e.g. ``"[[Total]]"`` / ``"{{ x }}"``
* ``inner``      -- the trimmed inner text, e.g. ``"Total"`` / ``"x"``
* ``delimiter``  -- one of ``"[["``, ``"["``, ``"{{"`` (the opening delimiter)
* ``kind``       -- ``"literal"`` for a known ``{{ }}`` literal placeholder;
                    any other (non-``"literal"``) value for NL tokens to infer
* ``location``   -- an object/mapping carrying enough to rewrite the token later:
                    a ``runs`` sequence of integer run indices the token spans,
                    and ``paragraph_text`` = the reconstructed full paragraph text.

Attribute access is used throughout; if the implementer chooses a dataclass
these read as attributes. (A dict-shaped Token would fail these tests, which is
an intentional contract choice.)
"""
import pytest

from docx import Document

# Import the unit under test. If the module / symbol does not exist yet this
# raises at collection time, which is an acceptable RED for XTF-1 (the function
# is unimplemented). It is NOT a typo / fixture error.
from src.reports.template_inference import extract_placeholders


# --------------------------------------------------------------------------- #
# Fixtures / helpers (build .docx programmatically with python-docx)
# --------------------------------------------------------------------------- #

def _save(doc, tmp_path, name="t.docx"):
    path = tmp_path / name
    doc.save(str(path))
    return str(path)


def _by_raw(tokens):
    """Index returned tokens by their ``raw`` string for order-independent asserts."""
    out = {}
    for t in tokens:
        out.setdefault(t.raw, []).append(t)
    return out


@pytest.fixture
def one_per_delimiter_docx(tmp_path):
    """A body with exactly one placeholder of each delimiter, each in its own
    paragraph (each placeholder is a single run -> trivially matched)."""
    doc = Document()
    doc.add_paragraph("[[Total Beneficiaries]]")
    doc.add_paragraph("[Average Age]")
    doc.add_paragraph("{{ region breakdown }}")
    return _save(doc, tmp_path)


# --------------------------------------------------------------------------- #
# AC: each delimiter matched individually
# --------------------------------------------------------------------------- #

def test_extract_double_square_delimiter(tmp_path):
    doc = Document()
    doc.add_paragraph("[[Total Beneficiaries]]")
    tokens = extract_placeholders(_save(doc, tmp_path))
    assert len(tokens) == 1
    tok = tokens[0]
    assert tok.raw == "[[Total Beneficiaries]]"
    assert tok.inner == "Total Beneficiaries"
    assert tok.delimiter == "[["


def test_extract_single_square_delimiter(tmp_path):
    doc = Document()
    doc.add_paragraph("[Average Age]")
    tokens = extract_placeholders(_save(doc, tmp_path))
    assert len(tokens) == 1
    tok = tokens[0]
    assert tok.raw == "[Average Age]"
    assert tok.inner == "Average Age"
    assert tok.delimiter == "["


def test_extract_double_brace_delimiter(tmp_path):
    doc = Document()
    doc.add_paragraph("{{ region breakdown }}")
    tokens = extract_placeholders(_save(doc, tmp_path))
    assert len(tokens) == 1
    tok = tokens[0]
    assert tok.raw == "{{ region breakdown }}"
    assert tok.inner == "region breakdown"
    assert tok.delimiter == "{{"


def test_extract_all_three_delimiters_in_one_doc(one_per_delimiter_docx):
    tokens = extract_placeholders(one_per_delimiter_docx)
    raws = {t.raw for t in tokens}
    assert raws == {
        "[[Total Beneficiaries]]",
        "[Average Age]",
        "{{ region breakdown }}",
    }


# --------------------------------------------------------------------------- #
# AC: precedence -- [[x]] matched once as [[x]], never double-matched as [x]
# --------------------------------------------------------------------------- #

def test_extract_double_square_is_one_token_not_inner_single_square(tmp_path):
    doc = Document()
    doc.add_paragraph("[[Total]]")
    tokens = extract_placeholders(_save(doc, tmp_path))
    # Exactly one token, and it is the [[ ]] token -- the [Total] substring
    # inside must NOT also be returned as a [ ] token.
    assert len(tokens) == 1
    tok = tokens[0]
    assert tok.raw == "[[Total]]"
    assert tok.delimiter == "[["
    assert tok.inner == "Total"
    assert all(t.raw != "[Total]" for t in tokens)


# --------------------------------------------------------------------------- #
# AC: a token whose characters span multiple runs is matched as one token
# --------------------------------------------------------------------------- #

def test_extract_token_split_across_runs_matched_as_single_token(tmp_path):
    """Hand-typed placeholders are almost always split across runs by Word.
    Simulate that: the characters of one [[ ]] placeholder are spread over
    several runs in the same paragraph. It must come back as ONE token."""
    doc = Document()
    para = doc.add_paragraph()
    for chunk in ["[[Tot", "al Benef", "iciaries", "]]"]:
        para.add_run(chunk)
    tokens = extract_placeholders(_save(doc, tmp_path))
    assert len(tokens) == 1
    tok = tokens[0]
    assert tok.raw == "[[Total Beneficiaries]]"
    assert tok.inner == "Total Beneficiaries"
    assert tok.delimiter == "[["


# --------------------------------------------------------------------------- #
# AC: tokens in a table cell, a header, and a footer are all extracted
# --------------------------------------------------------------------------- #

def test_tokens_in_table_header_and_footer_extracted(tmp_path):
    doc = Document()
    doc.add_paragraph("[[Body Token]]")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("{{ cell token }}")

    section = doc.sections[0]
    section.header.paragraphs[0].add_run("[Header Token]")
    section.footer.paragraphs[0].add_run("{{ footer token }}")

    tokens = extract_placeholders(_save(doc, tmp_path))
    raws = {t.raw for t in tokens}
    assert "[[Body Token]]" in raws
    assert "{{ cell token }}" in raws
    assert "[Header Token]" in raws
    assert "{{ footer token }}" in raws


# --------------------------------------------------------------------------- #
# AC: known {{ }} literal -> kind "literal", raw unchanged
# --------------------------------------------------------------------------- #

@pytest.mark.parametrize("raw,inner", [
    ("{{ report_title }}", "report_title"),
    ("{{ chart_sales }}", "chart_sales"),
])
def test_extract_known_literal_placeholder_marked_literal_and_unchanged(tmp_path, raw, inner):
    doc = Document()
    doc.add_paragraph(raw)
    tokens = extract_placeholders(_save(doc, tmp_path))
    by_raw = _by_raw(tokens)
    assert raw in by_raw, f"{raw!r} not returned"
    tok = by_raw[raw][0]
    assert tok.kind == "literal"
    assert tok.inner == inner
    assert tok.raw == raw  # left untouched


# --------------------------------------------------------------------------- #
# AC: non-literal {{ }} -> returned as an NL (non-literal) token for inference
# --------------------------------------------------------------------------- #

def test_extract_unknown_double_brace_is_non_literal_nl_token(tmp_path):
    doc = Document()
    doc.add_paragraph("{{ unknown thing }}")
    tokens = extract_placeholders(_save(doc, tmp_path))
    by_raw = _by_raw(tokens)
    assert "{{ unknown thing }}" in by_raw
    tok = by_raw["{{ unknown thing }}"][0]
    assert tok.inner == "unknown thing"
    assert tok.delimiter == "{{"
    # The whole point of the express path: it is offered up for inference,
    # i.e. it is NOT a passthrough literal.
    assert tok.kind != "literal"


# --------------------------------------------------------------------------- #
# AC: location records a run-span reference sufficient to rewrite the token
# --------------------------------------------------------------------------- #

def test_extract_location_runspan_round_trips_to_same_runs(tmp_path):
    """The recorded run-span must identify the same runs the token occupies.
    Build a paragraph whose runs are known, then verify location.runs indexes
    back into those runs to reconstruct the token's raw text."""
    doc = Document()
    para = doc.add_paragraph()
    chunks = ["intro ", "[[Tot", "al]]", " outro"]
    for chunk in chunks:
        para.add_run(chunk)

    path = _save(doc, tmp_path)
    tokens = extract_placeholders(path)
    assert len(tokens) == 1
    loc = tokens[0].location

    # location must carry an integer run-span and the reconstructed paragraph text.
    run_indices = list(loc.runs)
    assert run_indices, "location.runs is empty"
    assert all(isinstance(i, int) for i in run_indices)
    assert loc.paragraph_text == "".join(chunks)

    # Re-open the document and confirm the recorded run indices select the same
    # runs whose concatenation contains the token's raw text -- i.e. the span is
    # sufficient to rewrite the token in place.
    reopened = Document(path)
    target_para = next(
        p for p in reopened.paragraphs
        if "".join(r.text for r in p.runs) == "".join(chunks)
    )
    spanned = "".join(target_para.runs[i].text for i in run_indices)
    assert "[[Total]]" in spanned


# =========================================================================== #
# XTF-2 — Batched inference + local validation
# =========================================================================== #
# These tests are the spec for two new functions appended to
# ``src/reports/template_inference.py`` (XTF-2). They are derived strictly from
# the XTF-2 acceptance criteria and design spec §4.2 / §4.3. Written before the
# implementation lands; expected RED until ``infer_specs`` / ``annotate_proposals``
# exist.
#
# Contract committed to here (AC-derived, mirroring ask_engine shapes):
#
# ``infer_specs(nl_tokens, catalog, ai_cfg) -> List[Proposal]``
#   * makes exactly ONE batched ``lf_client.chat`` call (trace_name=
#     "template_inference", json_mode=True) over ALL nl_tokens + the catalog;
#   * returns one Proposal per token.
#
# A ``Proposal`` is a mapping (dict access) carrying at least:
#   * ``token_index`` -- int, index into the input token list
#   * ``kind``        -- one of chart | indicator | summary | table | narrative | metadata
#   * ``spec``        -- a config-shaped dict (chart: {name,title,type,questions,…};
#                        indicator: {name,stat,question,…}; summary: {name,stat,questions,…})
#   * ``name``        -- canonical slug (str)
#   * ``confidence``  -- float in 0..1
#   * ``reason``      -- str
#
# ``annotate_proposals(proposals, profile) -> List[Proposal]`` is local +
# deterministic (no AI). It reuses ``ask_engine.validate_recipe`` / ``CHART_REQS``
# / ``INDICATOR_STATS`` and adds:
#   * ``status``      -- "ok" or "needs_attention"
#   * ``reason``      -- human-readable (overwritten/augmented with the failure)
# and dedupes canonical ``name``s (suffix on collision). ``needs_attention`` is
# set when confidence is low, validation fails, or a referenced column is absent.
# Narrative tokens map to a fixed slot (recommendations/observations/summary_text)
# when the text clearly matches, else a ``summaries`` entry with ``stat: "ai"``.
#
# These assertions intentionally pin dict-shaped Proposals.

from src.reports import template_inference as ti
from src.reports import ask_engine


# Confidence threshold used by the tests. The implementation must treat a
# proposal *below* this as low-confidence (needs_attention). 0.5 is a midpoint
# clearly below "high"; the AC only requires "low confidence" be flagged, so the
# tests use values at the extremes (0.1 low, 0.95 high) to stay robust to the
# implementation's exact cutoff.
_LOW_CONF = 0.1
_HIGH_CONF = 0.95


def _profile_xtf2():
    """A profile shaped exactly like ``ask_engine.validate_recipe`` expects:
    keyed by table name, each table {name, rows, columns:[{name, role, …}]}.
    Mirrors tests/test_ask_engine.py::_profile_fixture."""
    return {
        "main": {
            "name": "main", "rows": 3,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "Region", "role": "categorical", "distinct": 2, "missing_pct": 0.0,
                 "top_values": [{"value": "N", "count": 2}, {"value": "S", "count": 1}]},
                {"name": "Age", "role": "quantitative", "distinct": 3, "missing_pct": 0.0,
                 "min": 10.0, "max": 30.0, "mean": 20.0, "median": 20.0},
                {"name": "Income", "role": "quantitative", "distinct": 3, "missing_pct": 0.0,
                 "min": 100.0, "max": 900.0, "mean": 500.0, "median": 500.0},
                {"name": "Story", "role": "qualitative", "distinct": 3, "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        }
    }


def _proposal(kind, spec, name, confidence=_HIGH_CONF, token_index=0, reason="proposed"):
    """Build a Proposal dict in the shape ``infer_specs`` returns."""
    return {
        "token_index": token_index,
        "kind": kind,
        "spec": dict(spec),
        "name": name,
        "confidence": confidence,
        "reason": reason,
    }


def _get(proposal, key):
    """Read a Proposal field whether the impl returns dicts or objects."""
    if isinstance(proposal, dict):
        return proposal[key]
    return getattr(proposal, key)


# --------------------------------------------------------------------------- #
# annotate_proposals — confidence gate
# --------------------------------------------------------------------------- #
def test_annotate_flags_low_confidence_as_needs_attention():
    """AC: needs_attention is set when confidence is low. A bar proposal that is
    otherwise valid but has a low confidence score must be flagged."""
    proposals = [
        _proposal("chart", {"name": "by_region", "title": "By region",
                            "type": "bar", "questions": ["Region"]},
                  name="by_region", confidence=_LOW_CONF),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    assert _get(out[0], "status") == "needs_attention"
    assert isinstance(_get(out[0], "reason"), str) and _get(out[0], "reason")


# --------------------------------------------------------------------------- #
# annotate_proposals — missing column
# --------------------------------------------------------------------------- #
def test_annotate_flags_missing_column():
    """AC: needs_attention when a referenced column is absent from the data.
    A bar chart on a column not present in the profile must be flagged, and the
    reason should name the offending column."""
    proposals = [
        _proposal("chart", {"name": "ghost", "title": "Ghost",
                            "type": "bar", "questions": ["NotAColumn"]},
                  name="ghost", confidence=_HIGH_CONF),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    assert _get(out[0], "status") == "needs_attention"
    assert "NotAColumn" in _get(out[0], "reason")


# --------------------------------------------------------------------------- #
# annotate_proposals — bad type/column combo via CHART_REQS
# --------------------------------------------------------------------------- #
def test_annotate_flags_scatter_with_one_quantitative():
    """AC: bad type/column combo flagged via validate_recipe/CHART_REQS. A
    scatter needs >=2 quantitative columns; one quantitative + one categorical
    must fail and the reason should mention the quantitative requirement."""
    proposals = [
        _proposal("chart", {"name": "scatter_bad", "title": "Scatter",
                            "type": "scatter", "questions": ["Age", "Region"]},
                  name="scatter_bad", confidence=_HIGH_CONF),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    assert _get(out[0], "status") == "needs_attention"
    assert "quantitative" in _get(out[0], "reason")


# --------------------------------------------------------------------------- #
# annotate_proposals — MNT-28: single-column chart type + group_by/2+ questions
# --------------------------------------------------------------------------- #
def test_annotate_flags_table_with_group_by_and_extra_question():
    """AC: this validator is shared by both call sites -- Express Fill inference
    (annotate_proposals -> _validate_data_proposal -> validate_recipe ->
    _validate_chart) must also reject a 'table' proposal that sets group_by
    and/or supplies 2+ questions, exactly like the direct ask_engine.validate_recipe
    call. Reproduces the live bug: a '[table of Satisfaction by Region]'
    placeholder inferred as {type: table, questions: [Region, Age], group_by:
    Region} must be flagged needs_attention, not silently passed through to a
    misleadingly-titled, factually wrong table render."""
    proposals = [
        _proposal("chart", {"name": "satisfaction_by_region", "title": "Satisfaction by Region",
                            "type": "table", "questions": ["Region", "Age"],
                            "group_by": "Region"},
                  name="satisfaction_by_region", confidence=_HIGH_CONF),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    assert _get(out[0], "status") == "needs_attention", (
        "a table proposal with group_by + 2 questions must not silently pass "
        "Express Fill validation"
    )
    reason = _get(out[0], "reason")
    assert "table" in reason, reason


# --------------------------------------------------------------------------- #
# annotate_proposals — valid proposals pass
# --------------------------------------------------------------------------- #
def test_annotate_passes_valid_bar_indicator_summary():
    """AC: a valid bar / indicator / summary proposal is status ok. All three
    reference real columns, satisfy their role requirements, and have high
    confidence."""
    proposals = [
        _proposal("chart", {"name": "by_region", "title": "By region",
                            "type": "bar", "questions": ["Region"]},
                  name="by_region", confidence=_HIGH_CONF, token_index=0),
        _proposal("indicator", {"name": "mean_age", "stat": "mean",
                                "question": "Age"},
                  name="mean_age", confidence=_HIGH_CONF, token_index=1),
        _proposal("summary", {"name": "income_summary", "stat": "sum",
                              "questions": ["Income"]},
                  name="income_summary", confidence=_HIGH_CONF, token_index=2),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    statuses = [_get(p, "status") for p in out]
    assert statuses == ["ok", "ok", "ok"], statuses


# --------------------------------------------------------------------------- #
# annotate_proposals — dedupe canonical names
# --------------------------------------------------------------------------- #
def test_annotate_dedupes_colliding_names_with_suffix():
    """AC: canonical names are deduped (suffix on collision). Two valid proposals
    that resolve to the same slug must end with distinct ``name`` values."""
    proposals = [
        _proposal("chart", {"name": "by_region", "title": "By region",
                            "type": "bar", "questions": ["Region"]},
                  name="by_region", confidence=_HIGH_CONF, token_index=0),
        _proposal("chart", {"name": "by_region", "title": "By region again",
                            "type": "bar", "questions": ["Region"]},
                  name="by_region", confidence=_HIGH_CONF, token_index=1),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    names = [_get(p, "name") for p in out]
    assert len(set(names)) == 2, f"names not deduped: {names}"
    # The original slug is preserved on (at least) one; the other is suffixed.
    assert "by_region" in names


# --------------------------------------------------------------------------- #
# annotate_proposals — narrative routing
# --------------------------------------------------------------------------- #
def test_annotate_narrative_recommendations_maps_to_slot():
    """AC: a narrative token clearly matching 'recommendations' maps to the fixed
    ``recommendations`` slot."""
    proposals = [
        _proposal("narrative", {}, name="recommendations",
                  confidence=_HIGH_CONF, reason="Recommendations"),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    p = out[0]
    # The canonical name resolves to the fixed slot regardless of internal spec
    # representation.
    assert _get(p, "name") == "recommendations"
    assert _get(p, "status") == "ok"


def test_annotate_free_form_narrative_maps_to_ai_summary():
    """AC: a free-form narrative (not a fixed slot) maps to a summaries entry with
    stat 'ai' and the placeholder text carried as the prompt."""
    placeholder_text = "Write a paragraph about progress against targets this quarter"
    proposals = [
        _proposal("narrative", {"prompt": placeholder_text},
                  name="progress_narrative", confidence=_HIGH_CONF,
                  reason=placeholder_text),
    ]
    out = ti.annotate_proposals(proposals, _profile_xtf2())
    p = out[0]
    spec = _get(p, "spec")
    assert spec.get("stat") == "ai", f"expected stat 'ai', got spec={spec}"
    assert placeholder_text in (spec.get("prompt") or ""), spec


# --------------------------------------------------------------------------- #
# MNT-19 -- annotate_proposals can accept a bullet_list proposal when there is
# no categorical column but there IS a usable (e.g. qualitative free-text)
# column, instead of being stuck with the always-failing 'table' fallback.
# --------------------------------------------------------------------------- #
def _profile_no_categorical():
    """A profile with NO categorical column at all -- only quantitative +
    qualitative -- mirroring the motivating bug: a free-text list placeholder
    (e.g. French 'actions_prioritaires') whose only underlying column is
    qualitative. 'table' can never validate against this profile."""
    return {
        "main": {
            "name": "main", "rows": 3,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "Age", "role": "quantitative", "distinct": 3, "missing_pct": 0.0,
                 "min": 10.0, "max": 30.0, "mean": 20.0, "median": 20.0},
                {"name": "actions_prioritaires", "role": "qualitative", "distinct": 3,
                 "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        }
    }


def test_annotate_bullet_list_proposal_validates_ok():
    """MNT-19 AC: given a placeholder whose underlying data has no categorical
    column but does have >=1 usable column, a bullet_list proposal for it must
    validate 'ok' (not needs_attention) -- proving inference can propose
    bullet_list instead of being forced toward the always-failing 'table'."""
    proposals = [
        _proposal("chart",
                  {"name": "actions_prioritaires", "title": "Actions prioritaires",
                   "type": "bullet_list", "questions": ["actions_prioritaires"]},
                  name="actions_prioritaires", confidence=_HIGH_CONF),
    ]
    out = ti.annotate_proposals(proposals, _profile_no_categorical())
    p = out[0]
    assert _get(p, "status") == "ok", _get(p, "reason")


# --------------------------------------------------------------------------- #
# infer_specs — exactly one batched LLM call
# --------------------------------------------------------------------------- #
def test_infer_specs_makes_one_batched_chat_call(monkeypatch):
    """AC: infer_specs makes a SINGLE batched lf_client.chat call for N tokens,
    via get_prompt('template_inference', …) + chat(trace_name='template_inference',
    json_mode=True). Mock the LLM (suggester convention) and assert call count==1."""
    calls = {"chat": 0, "trace_names": [], "json_modes": []}

    monkeypatch.setattr(
        ti.lf_client, "get_prompt",
        lambda *a, **k: ([{"role": "user", "content": "x"}], {}),
    )

    def _fake_chat(*a, **k):
        calls["chat"] += 1
        calls["trace_names"].append(k.get("trace_name"))
        calls["json_modes"].append(k.get("json_mode"))
        # One Proposal per input token, returned as a JSON string (suggester style).
        return (
            '{"proposals": ['
            '{"token_index": 0, "kind": "chart", "name": "by_region", '
            '"spec": {"name": "by_region", "type": "bar", "questions": ["Region"]}, '
            '"confidence": 0.9, "reason": "bar of region"},'
            '{"token_index": 1, "kind": "indicator", "name": "mean_age", '
            '"spec": {"name": "mean_age", "stat": "mean", "question": "Age"}, '
            '"confidence": 0.8, "reason": "mean age"}'
            ']}'
        )

    monkeypatch.setattr(ti.lf_client, "chat", _fake_chat)

    # Three NL tokens (objects with .inner, like XTF-1 Token); batched into one call.
    nl_tokens = [
        ti.Token(raw="[Region breakdown]", inner="Region breakdown",
                 delimiter="[", kind="nl", location=ti.Location()),
        ti.Token(raw="[average age]", inner="average age",
                 delimiter="[", kind="nl", location=ti.Location()),
        ti.Token(raw="[total income]", inner="total income",
                 delimiter="[", kind="nl", location=ti.Location()),
    ]
    catalog = ask_engine.build_catalog(_profile_xtf2())
    ai_cfg = {"provider": "openai", "model": "gpt-x", "api_key": "sk-test"}

    out = ti.infer_specs(nl_tokens, catalog, ai_cfg)

    assert calls["chat"] == 1, f"expected exactly one batched chat call, got {calls['chat']}"
    assert calls["trace_names"] == ["template_inference"]
    assert calls["json_modes"] == [True]
    assert isinstance(out, list) and out, "infer_specs returned no proposals"


# MNT-7 — infer_specs raises RuntimeError on malformed LLM responses, not silent []
# --------------------------------------------------------------------------- #
import pytest

@pytest.mark.parametrize("bad_response", [
    "",                         # empty string
    "null",                     # JSON null
    "{}",                       # empty object — no "proposals" key
    '{"error": "oops"}',        # error-shaped — no "proposals" key
    '{"proposals": null}',      # proposals key present but null, not a list
    '{"proposals": "nope"}',    # proposals is a string, not a list
    "not json at all",          # completely unparseable
])
def test_infer_specs_raises_on_malformed_llm_response(monkeypatch, bad_response):
    """MNT-7: infer_specs must raise RuntimeError (not return []) when the LLM
    response cannot be parsed as {proposals: list}. The endpoint's except Exception
    then surfaces HTTP 500 instead of a silent empty list."""
    monkeypatch.setattr(
        ti.lf_client, "get_prompt",
        lambda *a, **k: ([{"role": "user", "content": "x"}], {}),
    )
    monkeypatch.setattr(ti.lf_client, "chat", lambda *a, **k: bad_response)

    nl_tokens = [
        ti.Token(raw="[Total beneficiaries]", inner="Total beneficiaries",
                 delimiter="[", kind="nl", location=ti.Location()),
    ]
    catalog = ask_engine.build_catalog(_profile_xtf2())
    ai_cfg = {"provider": "openai", "model": "gpt-x", "api_key": "sk-test"}

    with pytest.raises(RuntimeError, match="proposals"):
        ti.infer_specs(nl_tokens, catalog, ai_cfg)


# --------------------------------------------------------------------------- #
# MNT-7 named unit tests (individual scenarios, card §Unit tests)
# --------------------------------------------------------------------------- #

def _nl_token():
    """A single NL token for infer_specs, matching the XTF-2 Token shape."""
    return ti.Token(
        raw="[Total beneficiaries]",
        inner="Total beneficiaries",
        delimiter="[",
        kind="nl",
        location=ti.Location(),
    )


def _ai_cfg_mnt7():
    return {"provider": "openai", "model": "gpt-x", "api_key": "sk-test"}


def test_infer_specs_raises_on_malformed_json(monkeypatch):
    """MNT-7 AC: `infer_specs` raises RuntimeError (not `return []`) when
    `lf_client.chat` returns a non-JSON string that `_loads_lenient` cannot
    parse as a dict with a 'proposals' list.

    Buggy behaviour: the function silently returns [], letting the endpoint
    respond HTTP 200 with {"proposals": [], "message": null} and the frontend
    renders the empty-placeholder state. Fixed behaviour: RuntimeError is raised
    so the endpoint's `except Exception` returns HTTP 500 instead.
    """
    monkeypatch.setattr(
        ti.lf_client, "get_prompt",
        lambda *a, **k: ([{"role": "user", "content": "x"}], {}),
    )
    # Return a non-JSON string — _loads_lenient returns None, proposals is absent.
    monkeypatch.setattr(ti.lf_client, "chat", lambda *a, **k: "not valid json at all")

    with pytest.raises(RuntimeError, match="proposals"):
        ti.infer_specs([_nl_token()], ask_engine.build_catalog(_profile_xtf2()), _ai_cfg_mnt7())


def test_infer_specs_raises_on_missing_proposals_key(monkeypatch):
    """MNT-7 AC: `infer_specs` raises RuntimeError (not `return []`) when
    `_loads_lenient` succeeds (valid JSON) but the 'proposals' key is absent.

    This is the boundary case: `{"result": []}` parses cleanly to a dict, so
    `_loads_lenient` returns it. The bug was that `(data or {}).get("proposals")`
    returned `None`, which is not a list, and the old code `return []`-ed silently.
    The fix replaces that with `raise RuntimeError` when `items` is not a list.
    """
    monkeypatch.setattr(
        ti.lf_client, "get_prompt",
        lambda *a, **k: ([{"role": "user", "content": "x"}], {}),
    )
    # Valid JSON, but the 'proposals' key is absent — _loads_lenient returns {"result": []}
    monkeypatch.setattr(ti.lf_client, "chat", lambda *a, **k: '{"result": []}')

    with pytest.raises(RuntimeError, match="proposals"):
        ti.infer_specs([_nl_token()], ask_engine.build_catalog(_profile_xtf2()), _ai_cfg_mnt7())


# =========================================================================== #
# XTF-3 — Apply: persist config + resolve template (apply_inference)
# =========================================================================== #
# These tests are the spec for ``apply_inference`` appended to
# ``src/reports/template_inference.py`` (XTF-3). Derived strictly from the XTF-3
# acceptance criteria and design spec §4.4. Written before the implementation
# lands; expected RED (AttributeError: module has no attribute 'apply_inference')
# until it exists.
#
# Contract committed to here (AC-derived):
#
# ``apply_inference(approved, cfg, template_path) -> (cfg, resolved_template_path)``
#
#   * ``approved`` is the list of approved Proposal dicts (the same shape
#     ``annotate_proposals`` returns: keys ``token_index``, ``kind``, ``spec``,
#     ``name``, ``confidence``, ``reason``, ``status``). Only ``status == "ok"``
#     proposals are expected to be passed in (the CLI/web layer drops flagged
#     ones before calling); these tests pass ``status="ok"``.
#
#   * To know WHERE each approved proposal's token lives in the .docx,
#     ``apply_inference`` re-runs ``extract_placeholders(template_path)``
#     internally and matches approved proposals to extracted tokens by
#     ``token_index`` (index into the extracted NL-token list). So a test:
#       1. builds a .docx with known NL placeholders,
#       2. calls ``extract_placeholders`` to learn token indices,
#       3. builds approved Proposal dicts referencing those ``token_index`` +
#          a ``spec`` / ``name`` / ``kind`` (status "ok"),
#       4. calls ``apply_inference``.
#
#   * Config: each approved spec is appended/merged into the section for its
#     kind — chart -> ``cfg["charts"]``, indicator -> ``cfg["indicators"]``,
#     summary -> ``cfg["summaries"]``, table -> ``cfg["tables"]`` — using the
#     established list-of-dicts shape, where the entry's ``name`` is the canonical
#     slug (e.g. ``by_region``). Existing user-authored entries are NEVER
#     clobbered; a colliding name is given a numeric suffix.
#
#   * Template resolution: the token's run span is replaced by a SINGLE clean run
#     whose text is the canonical ``{{ <prefix>_<slug> }}`` placeholder
#     (chart -> ``chart_``, indicator -> ``ind_``, summary -> ``summary_``,
#     table -> ``table_``); the other runs in the span are cleared. So the chart
#     placeholder is exactly ONE unbroken XML run. The resolved .docx is saved as
#     a NEW file (the original upload is preserved). The resolved path is returned.

from pathlib import Path


# Canonical placeholder prefix per kind (the {{ }} text the builder fills).
_KIND_PREFIX = {
    "chart": "chart_",
    "indicator": "ind_",
    "summary": "summary_",
    "table": "table_",
}


def _approved(kind, spec, name, token_index, status="ok", confidence=_HIGH_CONF):
    """Build an approved Proposal dict (annotate_proposals output shape)."""
    return {
        "token_index": token_index,
        "kind": kind,
        "spec": dict(spec),
        "name": name,
        "confidence": confidence,
        "reason": "approved",
        "status": status,
    }


def _docx_with_nl_placeholders(tmp_path, texts, name="upload.docx"):
    """Build a .docx with one NL placeholder per paragraph (single run each)."""
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    path = tmp_path / name
    doc.save(str(path))
    return str(path)


def _section_entry(cfg, section, slug):
    """Find the config entry in ``cfg[section]`` whose name matches ``slug`` (or a
    suffixed variant beginning with ``slug``). Returns the entry dict or None."""
    for e in cfg.get(section, []) or []:
        if e.get("name") == slug:
            return e
    return None


# --------------------------------------------------------------------------- #
# AC: chart proposal -> cfg["charts"], indicator -> cfg["indicators"] (shapes)
# --------------------------------------------------------------------------- #
def test_apply_writes_chart_and_indicator_into_their_sections(tmp_path):
    template = _docx_with_nl_placeholders(
        tmp_path, ["[Region breakdown]", "[Average age]"]
    )
    # token_index is the index into extract_placeholders' returned list.
    tokens = ti.extract_placeholders(template)
    assert len(tokens) == 2

    approved = [
        _approved("chart",
                  {"name": "by_region", "title": "By region",
                   "type": "bar", "questions": ["Region"]},
                  name="by_region", token_index=0),
        _approved("indicator",
                  {"name": "mean_age", "stat": "mean", "question": "Age"},
                  name="mean_age", token_index=1),
    ]
    cfg = {"api": {}, "form": {}}

    cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    chart = _section_entry(cfg_out, "charts", "by_region")
    assert chart is not None, f"chart not written: {cfg_out.get('charts')}"
    assert chart["type"] == "bar"
    assert chart["questions"] == ["Region"]

    ind = _section_entry(cfg_out, "indicators", "mean_age")
    assert ind is not None, f"indicator not written: {cfg_out.get('indicators')}"
    assert ind["stat"] == "mean"
    assert ind["question"] == "Age"


# --------------------------------------------------------------------------- #
# AC: never clobber existing user-authored entries; new entry appended
# --------------------------------------------------------------------------- #
def test_apply_preserves_existing_user_chart_and_appends_new(tmp_path):
    template = _docx_with_nl_placeholders(tmp_path, ["[Region breakdown]"])
    approved = [
        _approved("chart",
                  {"name": "by_region", "title": "By region",
                   "type": "bar", "questions": ["Region"]},
                  name="by_region", token_index=0),
    ]
    # Pre-seed a user-authored chart that must survive untouched.
    cfg = {
        "api": {}, "form": {},
        "charts": [
            {"name": "chart_existing", "title": "User chart",
             "type": "pie", "questions": ["Region"]},
        ],
    }

    cfg_out, _resolved = ti.apply_inference(approved, cfg, template)

    names = [c.get("name") for c in cfg_out.get("charts", [])]
    # The user's chart survives verbatim.
    existing = _section_entry(cfg_out, "charts", "chart_existing")
    assert existing is not None, f"user chart clobbered: {names}"
    assert existing["type"] == "pie"
    # The new chart is appended alongside it.
    assert _section_entry(cfg_out, "charts", "by_region") is not None, names
    assert len(cfg_out["charts"]) == 2, names


# --------------------------------------------------------------------------- #
# AC: two approved specs with the same base slug -> distinct suffixed names
# --------------------------------------------------------------------------- #
def test_apply_dedupes_colliding_base_slugs_with_suffix(tmp_path):
    template = _docx_with_nl_placeholders(
        tmp_path, ["[Region breakdown]", "[Region split]"]
    )
    approved = [
        _approved("chart",
                  {"name": "by_region", "title": "By region",
                   "type": "bar", "questions": ["Region"]},
                  name="by_region", token_index=0),
        _approved("chart",
                  {"name": "by_region", "title": "By region (2)",
                   "type": "bar", "questions": ["Region"]},
                  name="by_region", token_index=1),
    ]
    cfg = {"api": {}, "form": {}}

    cfg_out, _resolved = ti.apply_inference(approved, cfg, template)

    names = [c.get("name") for c in cfg_out.get("charts", [])]
    assert len(names) == 2, names
    # Two distinct names; both derived from the base slug.
    assert len(set(names)) == 2, f"slugs not deduped: {names}"
    assert "by_region" in names
    assert all(str(n).startswith("by_region") for n in names), names


# --------------------------------------------------------------------------- #
# AC: resolved chart placeholder occupies exactly ONE run with {{ chart_<slug> }}
# --------------------------------------------------------------------------- #
def test_apply_resolves_chart_placeholder_to_single_run(tmp_path):
    """The chart placeholder must be exactly one unbroken XML run so docxtpl can
    render it. Build the placeholder split across several runs (as Word does),
    apply, then assert the resolved paragraph holds the canonical placeholder in a
    single non-empty run."""
    doc = Document()
    para = doc.add_paragraph()
    for chunk in ["[Reg", "ion break", "down]"]:
        para.add_run(chunk)
    template = str(tmp_path / "upload.docx")
    doc.save(template)

    tokens = ti.extract_placeholders(template)
    assert len(tokens) == 1

    approved = [
        _approved("chart",
                  {"name": "by_region", "title": "By region",
                   "type": "bar", "questions": ["Region"]},
                  name="by_region", token_index=0),
    ]
    cfg = {"api": {}, "form": {}}

    _cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    expected = "{{ chart_by_region }}"
    reopened = Document(str(resolved))
    target = None
    for p in reopened.paragraphs:
        if expected in "".join(r.text for r in p.runs):
            target = p
            break
    assert target is not None, "resolved placeholder paragraph not found"

    # Exactly ONE run carries the placeholder text; the other runs in the span
    # are cleared (empty). So the placeholder is one unbroken run.
    nonempty = [r for r in target.runs if r.text]
    assert len(nonempty) == 1, (
        f"chart placeholder must be exactly one run, got {[r.text for r in target.runs]}"
    )
    assert nonempty[0].text == expected, nonempty[0].text


# --------------------------------------------------------------------------- #
# MNT-19 -- an approved bullet_list proposal (kind "chart", spec type
# "bullet_list") must resolve to {{ list_<name> }}, not {{ chart_<name> }},
# matching builder.py's list_<name> context key (the same convention
# template_generator.py already uses for a manually-added bullet_list
# placeholder).
# --------------------------------------------------------------------------- #
def test_apply_inference_bullet_list_uses_list_prefix(tmp_path):
    template = _docx_with_nl_placeholders(tmp_path, ["[Actions prioritaires]"])
    tokens = ti.extract_placeholders(template)
    assert len(tokens) == 1

    approved = [
        _approved("chart",
                  {"name": "actions_prioritaires", "title": "Actions prioritaires",
                   "type": "bullet_list", "questions": ["actions_prioritaires"]},
                  name="actions_prioritaires", token_index=0),
    ]
    cfg = {"api": {}, "form": {}}

    _cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    expected = "{{ list_actions_prioritaires }}"
    unexpected = "{{ chart_actions_prioritaires }}"
    reopened = Document(str(resolved))
    blob = "\n".join("".join(r.text for r in p.runs) for p in reopened.paragraphs)
    assert expected in blob, blob
    assert unexpected not in blob, blob


# --------------------------------------------------------------------------- #
# MNT-23 -- a first-class kind="list" proposal must persist into cfg["lists"]
# (its OWN section, not cfg["charts"]/cfg["tables"]) and its token must resolve
# to {{ list_<slug> }}, mirroring exactly how kind="table" -> cfg["tables"] /
# {{ table_<slug> }} works today.
# --------------------------------------------------------------------------- #
def test_apply_writes_list_proposal_into_lists_section(tmp_path):
    template = _docx_with_nl_placeholders(tmp_path, ["[Success stories]"])
    tokens = ti.extract_placeholders(template)
    assert len(tokens) == 1

    approved = [
        _approved("list",
                  {"name": "stories", "title": "Success stories", "question": "Story"},
                  name="stories", token_index=0),
    ]
    cfg = {"api": {}, "form": {}}

    cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    entry = _section_entry(cfg_out, "lists", "stories")
    assert entry is not None, f"list proposal not written into cfg['lists']: {cfg_out}"
    assert entry.get("question") == "Story"

    # Must NOT have landed in charts or tables.
    assert _section_entry(cfg_out, "charts", "stories") is None, cfg_out.get("charts")
    assert _section_entry(cfg_out, "tables", "stories") is None, cfg_out.get("tables")

    expected = "{{ list_stories }}"
    reopened = Document(str(resolved))
    blob = "\n".join("".join(r.text for r in p.runs) for p in reopened.paragraphs)
    assert expected in blob, blob


# --------------------------------------------------------------------------- #
# AC: the original uploaded .docx is preserved (resolved saved as new file)
# --------------------------------------------------------------------------- #
def test_apply_preserves_original_upload(tmp_path):
    template = _docx_with_nl_placeholders(tmp_path, ["[Region breakdown]"])
    approved = [
        _approved("chart",
                  {"name": "by_region", "title": "By region",
                   "type": "bar", "questions": ["Region"]},
                  name="by_region", token_index=0),
    ]
    cfg = {"api": {}, "form": {}}

    _cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    # The original upload still exists.
    assert Path(template).exists(), "original upload was not preserved"
    # The resolved template is a distinct, existing file.
    assert Path(str(resolved)).exists(), "resolved template was not written"
    assert Path(str(resolved)).resolve() != Path(template).resolve(), (
        "resolved template must be a new file, not the original upload"
    )


# =========================================================================== #
# XTF-4 — CLI commands (infer-template, apply-template)
# =========================================================================== #
# These tests are the spec for the two new Click commands appended to
# ``src/data/make.py`` (and their registration in ``web.main.ALLOWED_COMMANDS``)
# for XTF-4. They are derived strictly from the XTF-4 acceptance criteria and
# design spec §5. They are written before the commands exist and are expected
# to be RED until the commands land — RED for the RIGHT reason: Click reports
# "No such command 'infer-template' / 'apply-template'" (nonzero exit), and the
# ALLOWED_COMMANDS keys are missing — NOT because of fixture / mock / import bugs.
#
# They are selectable with ``-k "cli or command"`` (every test name below
# contains "cli"); the XTF-1/2/3 tests above are NOT matched by that filter, so
# they keep running under ``-k "extract or infer or annotate or apply"``.
#
# --------------------------------------------------------------------------- #
# Contract committed to here (AC-derived; the implementer must satisfy these
# exact flag names + message substrings, kept aligned with the card/spec wording)
# --------------------------------------------------------------------------- #
#
# ``infer-template --template <file> [--out reports/.template_inference.json]``
#   * runs extract_placeholders -> infer_specs -> annotate_proposals;
#   * writes the proposal LIST to the --out JSON (one entry per NON-literal
#     token; known {{ }} literals are passthrough and NOT proposals);
#   * prints a summary table (placeholder -> kind/name/status) and exits 0.
#   * No AI provider/key configured  -> nonzero exit + a message naming the AI
#     provider requirement (assert substring "AI provider", case-insensitive).
#   * No downloaded data             -> nonzero exit + the "run Download first"
#     message (assert substring "download first", case-insensitive).
#   * Zero placeholders found        -> a friendly no-op message + exit 0.
#
# ``apply-template [--from reports/.template_inference.json] [--build]``
#   * reads the (possibly user-edited) proposal list JSON;
#   * DROPS any proposal still flagged (status == "needs_attention") or not
#     approved before calling apply_inference (so a needs_attention row never
#     reaches config / the resolved template);
#   * calls apply_inference -> persists config via write_config + writes the
#     resolved template;
#   * with --build, chains into the build-report command via ctx.invoke
#     (the same _invoke seam run-all uses, so it is mockable).
#
# Both command names are present in ``web.main.ALLOWED_COMMANDS``.
#
# MOCKING SEAMS (chosen to mirror how build-report / download / run-all /
# ask_engine detect their preconditions, NOT invented):
#   * "no AI provider/key": same shape build-report's AI features check — an
#     ``ai`` config whose provider/api_key are absent (or api_key is an unresolved
#     ``env:`` ref) is treated as not-configured. The configured case uses
#     {provider: openai, api_key: sk-test} like the suggester tests.
#   * "no downloaded data": ``src/data/transform.load_processed_data`` raises
#     ``FileNotFoundError("... Run 'download' first.")`` when no session exists —
#     the same seam ask/build-report read through. We monkeypatch it on the
#     ``make`` module (where the command imports it) to simulate present/absent
#     data without writing real CSVs.
#   * profile + catalog: ``src/data/profile.profile_dataset`` and
#     ``ask_engine.build_catalog`` — patched to return a tiny deterministic
#     profile so the command does not depend on a real download.
#   * the LLM: ``template_inference.infer_specs`` is patched to return canned
#     proposals (the suggester-test convention of mocking at the inference
#     boundary rather than the raw lf_client.chat, since the command orchestrates
#     the module function).
#   * build-report chaining: patch ``make._invoke`` (the run-all sequencing
#     seam) and/or the ``apply_inference`` call to spy on the chained build.
#
# If the implementer wires a precondition at a different-but-equivalent seam,
# these monkeypatches still hold as long as the command imports the helpers via
# the ``make`` module namespace (the established pattern in make.py).

import json as _json

import pytest as _pytest
import yaml as _yaml
from click.testing import CliRunner

from src.data import make as _make


_AI_OK = {"provider": "openai", "api_key": "sk-test", "model": "gpt-x", "max_tokens": 1500}


def _profile_xtf4():
    """A minimal profile in the shape build_catalog / validate_recipe expect."""
    return {
        "main": {
            "name": "main", "rows": 3,
            "columns": [
                {"name": "Region", "role": "categorical", "distinct": 2, "missing_pct": 0.0,
                 "top_values": [{"value": "N", "count": 2}, {"value": "S", "count": 1}]},
                {"name": "Age", "role": "quantitative", "distinct": 3, "missing_pct": 0.0,
                 "min": 10.0, "max": 30.0, "mean": 20.0, "median": 20.0},
            ],
            "correlations": [], "duplicates": None,
        }
    }


def _write_cfg_xtf4(tmp_path, *, ai=True):
    """A config.yml with (optionally) a configured AI provider, written to disk
    so the CLI's ``load_config`` reads it the same way every other command does."""
    cfg = {
        "api": {"platform": "kobo", "url": "https://x.example.com/api/v2", "token": "t"},
        "form": {"uid": "aaa", "alias": "survey"},
        "questions": [{"export_label": "Region", "category": "categorical"}],
    }
    if ai:
        cfg["ai"] = dict(_AI_OK)
    p = tmp_path / "config.yml"
    p.write_text(_yaml.safe_dump(cfg))
    return p


def _docx_for_cli(tmp_path, texts, name="upload.docx"):
    """Build a .docx with one placeholder per paragraph (single run each)."""
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    path = tmp_path / name
    doc.save(str(path))
    return str(path)


def _patch_data_present(monkeypatch, profile=None):
    """Simulate a successful prior download: load_processed_data returns a frame +
    profile_dataset / build_catalog return a tiny deterministic profile/catalog.
    Patched on the ``make`` module namespace (where the command imports them)."""
    import pandas as pd
    from src.data import transform as _transform
    from src.data import profile as _profile_mod

    prof = profile or _profile_xtf4()
    df = pd.DataFrame({"Region": ["N", "N", "S"], "Age": [10, 20, 30]})

    # The command may import these via `from ... import` inside the function or as
    # module attributes; patch both the source modules and the make namespace so
    # whichever binding the implementer chooses resolves to our stub.
    for mod, attr, val in [
        (_transform, "load_processed_data", lambda *a, **k: (df, {})),
        (_profile_mod, "profile_dataset", lambda *a, **k: prof),
    ]:
        monkeypatch.setattr(mod, attr, val, raising=False)
    monkeypatch.setattr(
        ti.ask_engine, "build_catalog", lambda *a, **k: {"tables": []}, raising=False
    )


def _patch_no_data(monkeypatch):
    """Simulate no prior download: load_processed_data raises FileNotFoundError
    with the canonical "Run 'download' first." message (the real seam's wording)."""
    import pandas as pd  # noqa: F401  (ensures pandas importable in this env)
    from src.data import transform as _transform

    def _boom(*a, **k):
        raise FileNotFoundError(
            "No data matching data/processed/survey_data*.csv. Run 'download' first."
        )

    monkeypatch.setattr(_transform, "load_processed_data", _boom, raising=False)


# --------------------------------------------------------------------------- #
# infer-template — writes the --out JSON, one entry per non-literal token, exit 0
# --------------------------------------------------------------------------- #
def test_cli_infer_template_writes_proposal_json_and_exits_zero(tmp_path, monkeypatch):
    """AC: infer-template runs extract -> infer -> annotate, writes the proposal
    list to --out, and exits 0. The template has two NON-literal placeholders and
    one known {{ }} literal (passthrough); the JSON must hold one entry per
    non-literal token (the literal is NOT a proposal)."""
    cfg_path = _write_cfg_xtf4(tmp_path, ai=True)
    template = _docx_for_cli(
        tmp_path,
        ["[Region breakdown]", "[Average age]", "{{ report_title }}"],
    )
    out_json = tmp_path / "proposals.json"

    _patch_data_present(monkeypatch)

    def _fake_infer(nl_tokens, catalog, ai_cfg):
        # One proposal per NON-literal token passed in (the literal is filtered out
        # by the command before infer_specs is called).
        return [
            {"token_index": i, "kind": "chart",
             "spec": {"name": f"c{i}", "type": "bar", "questions": ["Region"]},
             "name": f"c{i}", "confidence": 0.9, "reason": "ok"}
            for i, _t in enumerate(nl_tokens)
        ]

    monkeypatch.setattr(ti, "infer_specs", _fake_infer, raising=False)

    res = CliRunner().invoke(
        _make.cli,
        ["--config", str(cfg_path), "infer-template",
         "--template", template, "--out", str(out_json)],
    )

    assert res.exit_code == 0, res.output
    assert out_json.exists(), f"--out JSON not written. output:\n{res.output}"
    data = _json.loads(out_json.read_text())
    proposals = data["proposals"] if isinstance(data, dict) else data
    assert isinstance(proposals, list)
    # Exactly the two non-literal tokens -> two proposals (the literal is excluded).
    assert len(proposals) == 2, proposals


# --------------------------------------------------------------------------- #
# infer-template — no AI provider/key configured -> nonzero + AI-provider message
# --------------------------------------------------------------------------- #
def test_cli_infer_template_errors_without_ai_provider(tmp_path, monkeypatch):
    """AC: infer-template errors clearly when no AI provider/key is configured,
    with a message naming the AI-provider requirement (the feature requires AI —
    it cannot degrade to seeds)."""
    cfg_path = _write_cfg_xtf4(tmp_path, ai=False)
    template = _docx_for_cli(tmp_path, ["[Region breakdown]"])
    out_json = tmp_path / "proposals.json"

    # Data is present so the ONLY failing precondition is the missing AI provider.
    _patch_data_present(monkeypatch)

    res = CliRunner().invoke(
        _make.cli,
        ["--config", str(cfg_path), "infer-template",
         "--template", template, "--out", str(out_json)],
    )

    assert res.exit_code != 0, res.output
    assert "ai provider" in res.output.lower(), res.output


# --------------------------------------------------------------------------- #
# infer-template — no downloaded data -> nonzero + "run Download first" message
# --------------------------------------------------------------------------- #
def test_cli_infer_template_errors_without_downloaded_data(tmp_path, monkeypatch):
    """AC: infer-template errors when no data has been downloaded (local
    validation needs real columns), with the "run Download first" message."""
    cfg_path = _write_cfg_xtf4(tmp_path, ai=True)
    template = _docx_for_cli(tmp_path, ["[Region breakdown]"])
    out_json = tmp_path / "proposals.json"

    _patch_no_data(monkeypatch)

    res = CliRunner().invoke(
        _make.cli,
        ["--config", str(cfg_path), "infer-template",
         "--template", template, "--out", str(out_json)],
    )

    assert res.exit_code != 0, res.output
    assert "download first" in res.output.lower(), res.output


# --------------------------------------------------------------------------- #
# infer-template — zero placeholders -> friendly no-op message, exit 0
# --------------------------------------------------------------------------- #
def test_cli_infer_template_zero_placeholders_is_friendly_noop(tmp_path, monkeypatch):
    """AC: zero placeholders found -> a friendly no-op message, non-error exit."""
    cfg_path = _write_cfg_xtf4(tmp_path, ai=True)
    template = _docx_for_cli(tmp_path, ["Just prose, no placeholders here."])
    out_json = tmp_path / "proposals.json"

    _patch_data_present(monkeypatch)
    # infer_specs must NOT be needed when there is nothing to infer; if it is
    # called it would still no-op on an empty list, so leave it real.

    res = CliRunner().invoke(
        _make.cli,
        ["--config", str(cfg_path), "infer-template",
         "--template", template, "--out", str(out_json)],
    )

    assert res.exit_code == 0, res.output
    # A friendly message mentioning that no placeholders were found.
    assert "no placeholder" in res.output.lower(), res.output


# --------------------------------------------------------------------------- #
# apply-template — writes config + resolved template and DROPS a needs_attention
# proposal that was not approved
# --------------------------------------------------------------------------- #
def test_cli_applytmpl_drops_needs_attention_and_writes_config(tmp_path, monkeypatch):
    """AC: apply-template reads the proposals, drops any still flagged/unapproved,
    runs apply_inference (writes config + resolved template). Here the proposal
    list has one ``ok`` chart and one ``needs_attention`` chart; only the ``ok``
    one may reach apply_inference (and thus config)."""
    cfg_path = _write_cfg_xtf4(tmp_path, ai=True)
    template = _docx_for_cli(tmp_path, ["[Region breakdown]", "[Mystery thing]"])

    proposals = [
        {"token_index": 0, "kind": "chart",
         "spec": {"name": "by_region", "type": "bar", "questions": ["Region"]},
         "name": "by_region", "confidence": 0.9, "reason": "ok", "status": "ok"},
        {"token_index": 1, "kind": "chart",
         "spec": {"name": "mystery", "type": "bar", "questions": ["NotAColumn"]},
         "name": "mystery", "confidence": 0.2, "reason": "missing column",
         "status": "needs_attention"},
    ]
    from_json = tmp_path / "proposals.json"
    from_json.write_text(_json.dumps({"proposals": proposals, "template": template}))

    # Spy on apply_inference to capture which proposals survived the drop, and to
    # avoid depending on its real config/docx side effects here.
    seen = {"approved": None}

    def _spy_apply(approved, cfg, template_path):
        seen["approved"] = list(approved)
        # Mimic the real return: (cfg with the spec written, resolved path).
        cfg.setdefault("charts", [])
        for p in approved:
            cfg["charts"].append(dict(p.get("spec") or {}, name=p.get("name")))
        resolved = str(tmp_path / "upload.resolved.docx")
        Document().save(resolved)
        return cfg, resolved

    monkeypatch.setattr(ti, "apply_inference", _spy_apply, raising=False)

    res = CliRunner().invoke(
        _make.cli,
        ["--config", str(cfg_path), "apply-template", "--from", str(from_json)],
    )

    assert res.exit_code == 0, res.output
    # The needs_attention proposal was dropped: only the ok one reached apply.
    assert seen["approved"] is not None, "apply_inference was never called"
    names = [p.get("name") for p in seen["approved"]]
    assert names == ["by_region"], f"needs_attention not dropped: {names}"

    # Config was persisted with the approved chart.
    saved = _yaml.safe_load(cfg_path.read_text())
    chart_names = [c.get("name") for c in (saved.get("charts") or [])]
    assert "by_region" in chart_names, saved
    assert "mystery" not in chart_names, saved


# --------------------------------------------------------------------------- #
# apply-template --build — chains into the build-report path
# --------------------------------------------------------------------------- #
def test_cli_applytmpl_build_chains_into_build_report(tmp_path, monkeypatch):
    """AC: with --build, apply-template chains into build-report. We assert the
    chained call via the same _invoke seam run-all uses (monkeypatched to record
    the command name), so no real report is built."""
    cfg_path = _write_cfg_xtf4(tmp_path, ai=True)
    template = _docx_for_cli(tmp_path, ["[Region breakdown]"])

    proposals = [
        {"token_index": 0, "kind": "chart",
         "spec": {"name": "by_region", "type": "bar", "questions": ["Region"]},
         "name": "by_region", "confidence": 0.9, "reason": "ok", "status": "ok"},
    ]
    from_json = tmp_path / "proposals.json"
    from_json.write_text(_json.dumps({"proposals": proposals, "template": template}))

    def _spy_apply(approved, cfg, template_path):
        cfg.setdefault("charts", []).append(
            {"name": "by_region", "type": "bar", "questions": ["Region"]}
        )
        resolved = str(tmp_path / "upload.resolved.docx")
        Document().save(resolved)
        return cfg, resolved

    monkeypatch.setattr(ti, "apply_inference", _spy_apply, raising=False)

    invoked = []
    monkeypatch.setattr(
        _make, "_invoke",
        lambda ctx, command, **kw: invoked.append(command.name),
        raising=False,
    )

    res = CliRunner().invoke(
        _make.cli,
        ["--config", str(cfg_path), "apply-template",
         "--from", str(from_json), "--build"],
    )

    assert res.exit_code == 0, res.output
    assert "build-report" in invoked, (
        f"--build did not chain into build-report; invoked={invoked}\n{res.output}"
    )


# --------------------------------------------------------------------------- #
# Both command names registered in web.main.ALLOWED_COMMANDS
# --------------------------------------------------------------------------- #
def test_cli_names_registered_in_allowed_commands():
    """AC: both commands added to ALLOWED_COMMANDS in web/main.py (so they are
    runnable via the SSE run endpoint)."""
    from web import main as web_main
    assert "infer-template" in web_main.ALLOWED_COMMANDS
    assert "apply-template" in web_main.ALLOWED_COMMANDS


# --------------------------------------------------------------------------- #
# The commands actually exist on the CLI group (guards against the RED being a
# generic Click usage error rather than a missing command — once implemented,
# `--help` must succeed for each).
# --------------------------------------------------------------------------- #
@_pytest.mark.parametrize("command", ["infer-template", "apply-template"])
def test_cli_name_is_registered_on_cli_group(command):
    """AC: infer-template / apply-template are real Click commands on the group."""
    res = CliRunner().invoke(_make.cli, [command, "--help"])
    assert res.exit_code == 0, res.output
    assert "No such command" not in res.output


# =========================================================================== #
# XTF-22 — Deterministic auto-modeling resolver for cross-table columns
# =========================================================================== #
# These tests are the spec for a NEW deterministic pass appended to
# ``src/reports/template_inference.py`` (XTF-22):
#
#   ``resolve_sources(proposals, profile) -> List[Dict]``
#
# It runs AFTER ``infer_specs`` and BEFORE ``annotate_proposals`` (no LLM call),
# and for each DATA proposal (kind ∈ {chart, indicator, summary, table}) it
# resolves the ``source`` deterministically from the profile's tables. Derived
# strictly from the XTF-22 acceptance criteria + design spec §④. Written before
# the resolver exists; expected RED (``AttributeError: module 'template_inference'
# has no attribute 'resolve_sources'``) until it lands — RED for the RIGHT reason
# (the function is missing), not a fixture/import bug.
#
# --------------------------------------------------------------------------- #
# ASSUMED CONTRACT for ``resolve_sources(proposals, profile)`` (the implementer
# must satisfy this; chosen as the most natural design consistent with the card,
# the spec, and the surrounding code):
#
#   Signature:  resolve_sources(proposals: List[Proposal], profile: Dict) -> List[Dict]
#
#   Referenced columns of a data proposal are collected from the spec's
#   ``questions`` (list) PLUS ``group_by`` (str) PLUS ``question`` (str, the
#   indicator single-column field) when present — mirroring how
#   ``ask_engine._validate_chart`` / ``_validate_indicator`` read columns.
#
#   For each data proposal, by where its referenced columns live:
#     * ALL in ``main``                  → leave the spec's ``source`` as-is
#                                          (absent or "main"); no view synthesized.
#     * ALL in ONE non-main table        → stamp ``spec["source"] = <table>``.
#                                          When a column appears in several tables,
#                                          pick the table that contains the MOST of
#                                          the spec's columns (the builder._pick_df
#                                          "most-columns-match" heuristic).
#     * SPAN a repeat table + ``main``   → synthesize a view dict
#                                          ``{name, source:<repeat_table>,
#                                            join_parent:[<main cols referenced>]}``
#                                          (carry ``group_by``/``question``/``agg``
#                                          only when the chart is inherently
#                                          aggregated) and point ``spec["source"]``
#                                          at the new view ``name``. The view dict
#                                          is appended to the RETURNED list (the
#                                          "pending views" collection that
#                                          /api/template/apply persists).
#     * STUCK (a column in NO table, or a genuine multi-table tie with no
#       disambiguating majority) → flag the proposal:
#                                          ``proposal["status"] = "needs_attention"``
#                                          and ``proposal["reason"]`` names the
#                                          candidate tables (or states no table
#                                          contains the column). No view synthesized.
#
#   View names are DETERMINISTIC + collision-safe (e.g.
#   ``auto_<repeat_leaf>__<joincols>``), de-duped against existing ``profile``
#   views AND against views already synthesized in the same call, so re-running
#   the resolver on the same proposals/profile yields the SAME name(s) and no
#   duplicates (idempotent).
#
#   ``resolve_sources`` MUTATES proposals in place (stamping ``spec["source"]`` /
#   ``status`` / ``reason``) AND returns the list of synthesized view dicts.
#
# The tests below read the proposal back via ``_get`` (dict-or-object tolerant)
# and inspect the proposal object the resolver mutated in place; the synthesized
# views are read from the resolver's return value.
# --------------------------------------------------------------------------- #


def _profile_xtf22():
    """A profile with ``main`` + one repeat-group base table, shaped exactly like
    ``profile_dataset`` returns (keyed by table name; each table
    ``{name, rows, columns:[{name, role, …}]}``; repeat rows carry linkage cols).

    The repeat table key uses the full slash-replaced repeat-path convention
    (``health_facilities`` here standing in for a leaf repeat group)."""
    return {
        "main": {
            "name": "main", "rows": 3,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "Commune", "role": "categorical", "distinct": 2, "missing_pct": 0.0,
                 "top_values": [{"value": "A", "count": 2}, {"value": "B", "count": 1}]},
                {"name": "Region", "role": "categorical", "distinct": 2, "missing_pct": 0.0,
                 "top_values": [{"value": "N", "count": 2}, {"value": "S", "count": 1}]},
            ],
            "correlations": [], "duplicates": None,
        },
        "health_facilities": {
            "name": "health_facilities", "rows": 5,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "_root_id", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "USB operational", "role": "categorical", "distinct": 2,
                 "missing_pct": 0.0,
                 "top_values": [{"value": "Oui", "count": 3}, {"value": "Non", "count": 2}]},
                {"name": "Beds", "role": "quantitative", "distinct": 4, "missing_pct": 0.0,
                 "min": 1.0, "max": 40.0, "mean": 12.0, "median": 8.0},
            ],
            "correlations": [], "duplicates": None,
        },
    }


def _data_proposal(kind, spec, name, token_index=0, confidence=_HIGH_CONF,
                   reason="proposed"):
    """A data Proposal in the shape ``infer_specs`` returns (no ``status`` yet —
    resolve_sources runs BEFORE annotate_proposals stamps status)."""
    return {
        "token_index": token_index,
        "kind": kind,
        "spec": dict(spec),
        "name": name,
        "confidence": confidence,
        "reason": reason,
    }


# --------------------------------------------------------------------------- #
# (1) single repeat-group column → source stamped to that table; validates clean
# --------------------------------------------------------------------------- #
def test_resolve_single_repeat_column():
    """AC: a proposal whose referenced columns all live in a single repeat-group
    table gets ``source`` stamped to that table and validates clean (no
    needs_attention) through the existing annotate/validate path.

    The chart references one column ("USB operational") that lives ONLY in the
    ``health_facilities`` repeat table. Before the resolver, validation defaults
    source to "main" and rejects it; after the resolver stamps
    ``source: health_facilities`` it validates clean."""
    profile = _profile_xtf22()
    proposals = [
        _data_proposal(
            "chart",
            {"name": "usb_ops", "title": "USB operational",
             "type": "bar", "questions": ["USB operational"]},
            name="usb_ops",
        ),
    ]

    views = ti.resolve_sources(proposals, profile)

    # No view is synthesized for the pure single-table case.
    assert views == [] or all(v.get("source") != "health_facilities" or v.get("join_parent")
                              for v in views), views
    # The spec's source is stamped to the repeat table.
    spec = _get(proposals[0], "spec")
    assert spec.get("source") == "health_facilities", spec

    # And the proposal now validates clean through the existing annotate path.
    out = ti.annotate_proposals(proposals, profile)
    assert _get(out[0], "status") == "ok", _get(out[0], "reason")


# --------------------------------------------------------------------------- #
# (2) repeat column + main column → synthesized view; spec sources the view
# --------------------------------------------------------------------------- #
def test_resolve_join_synthesizes_view():
    """AC: a proposal referencing a repeat-group column + a ``main`` column yields
    a synthesized view ``{source:<repeat_table>, join_parent:[<main col>]}`` and
    the spec's ``source`` points at the new view name.

    The chart groups the repeat column "Beds" by the main column "Commune", so it
    spans ``health_facilities`` (repeat) + ``main`` — the join case."""
    profile = _profile_xtf22()
    proposals = [
        _data_proposal(
            "chart",
            {"name": "beds_by_commune", "title": "Beds by commune",
             "type": "bar", "questions": ["Beds"], "group_by": "Commune"},
            name="beds_by_commune",
        ),
    ]

    views = ti.resolve_sources(proposals, profile)

    # Exactly one synthesized view is returned for apply to persist.
    assert isinstance(views, list) and len(views) == 1, views
    view = views[0]
    assert view.get("source") == "health_facilities", view
    assert view.get("join_parent") == ["Commune"], view
    assert view.get("name"), "synthesized view must carry a name"

    # The spec's source now points at the synthesized view (not "main").
    spec = _get(proposals[0], "spec")
    assert spec.get("source") == view["name"], spec


# --------------------------------------------------------------------------- #
# (3) idempotent — re-running yields the same view name(s), no duplicates
# --------------------------------------------------------------------------- #
def test_resolve_idempotent():
    """AC: synthesized view names are deterministic + collision-safe, so running
    the resolver on the same proposals/profile produces the SAME view name(s) with
    NO duplicates (re-running Infer is idempotent)."""
    profile = _profile_xtf22()

    def _fresh():
        return [
            _data_proposal(
                "chart",
                {"name": "beds_by_commune", "title": "Beds by commune",
                 "type": "bar", "questions": ["Beds"], "group_by": "Commune"},
                name="beds_by_commune",
            ),
        ]

    views_a = ti.resolve_sources(_fresh(), profile)
    views_b = ti.resolve_sources(_fresh(), profile)

    names_a = [v["name"] for v in views_a]
    names_b = [v["name"] for v in views_b]
    # Deterministic: same name(s) on both runs.
    assert names_a == names_b, (names_a, names_b)
    # No duplicate names within a single run.
    assert len(names_a) == len(set(names_a)), names_a

    # And persisting the run-1 views into the profile must NOT create a duplicate
    # on a re-run: a profile that already contains the synthesized view (by name)
    # is reused rather than a new differently-named view minted.
    profile_with_view = dict(profile)
    # Mirror how build_views reads views: cfg["views"] is a list of view dicts;
    # but the resolver de-dupes against EXISTING view names. Represent the already
    # persisted view both as a profile table (so its columns resolve) keyed by the
    # view name — the minimal stand-in for "this view already exists".
    profile_with_view[names_a[0]] = {
        "name": names_a[0], "rows": 5,
        "columns": profile["health_facilities"]["columns"]
        + [{"name": "Commune", "role": "categorical", "distinct": 2, "missing_pct": 0.0}],
        "correlations": [], "duplicates": None,
    }
    views_c = ti.resolve_sources(_fresh(), profile_with_view)
    names_c = [v["name"] for v in views_c]
    # The same deterministic name is reused (no `_2` suffix duplicate minted).
    assert names_c == names_a, (names_c, names_a)


# --------------------------------------------------------------------------- #
# (4) column in NO table → needs_attention, reason says no table contains it
# --------------------------------------------------------------------------- #
def test_resolve_unknown_column_flagged():
    """AC: a column present in NO table stays needs_attention with a reason saying
    no table contains it.

    The chart references "Ghost" which exists in neither ``main`` nor
    ``health_facilities``; the resolver flags the proposal and the reason names the
    offending column."""
    profile = _profile_xtf22()
    proposals = [
        _data_proposal(
            "chart",
            {"name": "ghost", "title": "Ghost", "type": "bar",
             "questions": ["Ghost"]},
            name="ghost",
        ),
    ]

    ti.resolve_sources(proposals, profile)

    assert _get(proposals[0], "status") == "needs_attention", proposals[0]
    reason = _get(proposals[0], "reason")
    assert "Ghost" in reason, reason

    # The existing annotate path keeps it flagged too (does not silently pass).
    out = ti.annotate_proposals(proposals, profile)
    assert _get(out[0], "status") == "needs_attention", _get(out[0], "reason")


# --------------------------------------------------------------------------- #
# (5) genuine multi-table tie → needs_attention naming both candidate tables
# --------------------------------------------------------------------------- #
def test_resolve_tie_flagged():
    """AC: a genuine multi-table tie (a single referenced column present in 2+
    tables with no disambiguating majority) stays needs_attention with a reason
    naming BOTH candidate tables.

    "Shared" lives in two repeat tables and nowhere else; a single-column chart on
    it has no majority to break the tie — the resolver must flag it and name both
    candidate tables."""
    profile = {
        "main": {
            "name": "main", "rows": 3,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        },
        "facilities": {
            "name": "facilities", "rows": 4,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "Shared", "role": "categorical", "distinct": 2, "missing_pct": 0.0,
                 "top_values": [{"value": "x", "count": 2}, {"value": "y", "count": 2}]},
            ],
            "correlations": [], "duplicates": None,
        },
        "staff": {
            "name": "staff", "rows": 6,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 3, "missing_pct": 0.0},
                {"name": "Shared", "role": "categorical", "distinct": 2, "missing_pct": 0.0,
                 "top_values": [{"value": "x", "count": 3}, {"value": "y", "count": 3}]},
            ],
            "correlations": [], "duplicates": None,
        },
    }
    proposals = [
        _data_proposal(
            "chart",
            {"name": "shared", "title": "Shared", "type": "bar",
             "questions": ["Shared"]},
            name="shared",
        ),
    ]

    ti.resolve_sources(proposals, profile)

    assert _get(proposals[0], "status") == "needs_attention", proposals[0]
    reason = _get(proposals[0], "reason")
    assert "facilities" in reason and "staff" in reason, reason


# =========================================================================== #
# XTF-25 — extractor must read Word content controls (w:sdt)
# =========================================================================== #
# ``paragraph.runs`` only yields top-level ``w:r`` elements; text wrapped in a
# Word gray-shaded content control lives at ``w:sdt/w:sdtContent/w:r/w:t`` and
# is therefore invisible to the current extractor.  Fix: walk
# ``paragraph._p.iter()`` for ALL ``w:t`` descendants.
# These four tests are derived from the XTF-25 Acceptance criteria and are
# expected to be RED until the fix lands.
# =========================================================================== #

from lxml import etree as _etree
from src.reports.template_inference import _tokens_in_paragraph


_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _sdt_paragraph(text: str):
    """Return a minimal python-docx-compatible paragraph object whose text is
    wrapped inside a ``w:sdt`` content control rather than a bare ``w:r``."""
    doc = Document()
    para = doc.add_paragraph()

    # Build:  <w:sdt><w:sdtContent><w:r><w:t>text</w:t></w:r></w:sdtContent></w:sdt>
    sdt = _etree.SubElement(para._p, f"{{{_W}}}sdt")
    sdt_content = _etree.SubElement(sdt, f"{{{_W}}}sdtContent")
    wr = _etree.SubElement(sdt_content, f"{{{_W}}}r")
    wt = _etree.SubElement(wr, f"{{{_W}}}t")
    wt.text = text

    return para


def _mixed_paragraph(plain_text: str, sdt_text: str):
    """Return a paragraph that has BOTH a plain ``w:r`` run AND a ``w:sdt`` run."""
    doc = Document()
    para = doc.add_paragraph(plain_text)  # creates a plain w:r run

    sdt = _etree.SubElement(para._p, f"{{{_W}}}sdt")
    sdt_content = _etree.SubElement(sdt, f"{{{_W}}}sdtContent")
    wr = _etree.SubElement(sdt_content, f"{{{_W}}}r")
    wt = _etree.SubElement(wr, f"{{{_W}}}t")
    wt.text = sdt_text

    return para


# AC1 — sdt-wrapped placeholder is detected
def test_xtf25_sdt_token_detected():
    """_tokens_in_paragraph must return the token inside a w:sdt content control."""
    para = _sdt_paragraph("[[TOKEN_IN_SDT]]")
    tokens = _tokens_in_paragraph(para)
    inners = [t.inner for t in tokens]
    assert "TOKEN_IN_SDT" in inners, (
        f"expected 'TOKEN_IN_SDT' but got {inners!r} — "
        "paragraph.runs misses w:sdt content"
    )


# AC2 — plain-paragraph regression: existing behaviour must be preserved
def test_xtf25_plain_run_still_detected(tmp_path):
    """_tokens_in_paragraph must still find tokens in a plain w:r paragraph."""
    doc = Document()
    para = doc.add_paragraph("[[TOKEN_PLAIN]]")
    tokens = _tokens_in_paragraph(para)
    inners = [t.inner for t in tokens]
    assert "TOKEN_PLAIN" in inners, (
        f"expected 'TOKEN_PLAIN' but got {inners!r}"
    )


# AC3 — paragraph with BOTH plain run and sdt run returns both tokens
def test_xtf25_mixed_paragraph_returns_both_tokens():
    """A paragraph containing a plain w:r AND a w:sdt run must yield both tokens."""
    para = _mixed_paragraph("[[PLAIN_TOKEN]] ", "[[SDT_TOKEN]]")
    tokens = _tokens_in_paragraph(para)
    inners = [t.inner for t in tokens]
    assert "PLAIN_TOKEN" in inners, f"plain token missing from {inners!r}"
    assert "SDT_TOKEN" in inners, f"sdt token missing from {inners!r}"


# AC4 — extract_placeholders end-to-end: non-empty on a doc with sdt placeholders
def test_xtf25_extract_placeholders_returns_nonempty_for_sdt_doc(tmp_path):
    """extract_placeholders must return a non-empty list when the doc's tokens
    are inside w:sdt content controls."""
    doc = Document()
    para = doc.add_paragraph()
    # Inject sdt wrapping [[REPORT_TITLE]] — mirrors the real-world failure.
    sdt = _etree.SubElement(para._p, f"{{{_W}}}sdt")
    sdt_content = _etree.SubElement(sdt, f"{{{_W}}}sdtContent")
    wr = _etree.SubElement(sdt_content, f"{{{_W}}}r")
    wt = _etree.SubElement(wr, f"{{{_W}}}t")
    wt.text = "[[REPORT_TITLE]]"

    path = tmp_path / "sdt_doc.docx"
    doc.save(str(path))

    tokens = extract_placeholders(str(path))
    assert len(tokens) > 0, (
        "extract_placeholders returned [] for a doc whose only placeholder is "
        "inside a w:sdt content control"
    )


# =========================================================================== #
# XTF-26 — annotate_proposals auto-resolves repeat-table columns
# =========================================================================== #
# Bug: when a proposal's target column is absent from ``main`` but lives in a
# repeat-group table, ``annotate_proposals`` sets ``status: needs_attention``
# with no ``source`` suggestion, leaving the placeholder blank.
#
# Fix (derived strictly from the XTF-26 Acceptance criteria):
#   1. Column absent from ``main`` but present in EXACTLY ONE repeat table
#      → set ``source`` to that table name and ``status`` to ``"ok"``.
#   2. Column present in MULTIPLE repeat tables
#      → set ``source`` to the table with the most rows, ``status`` to
#        ``"review"``, and include a ``note`` listing the alternative table names.
#   3. Column not found in ``main`` or any repeat table
#      → ``status`` remains ``"needs_attention"``.
#
# These tests call ``annotate_proposals`` directly (no intermediate
# ``resolve_sources``) and are expected to be RED until the fix lands.
# =========================================================================== #


def _profile_xtf26_single_repeat():
    """Profile: ``main`` has Region/Age; ``demographics`` repeat holds
    ``nombre_menages`` and ``nombre_habitants``; ``collaborations`` repeat
    holds ``organisation``."""
    return {
        "main": {
            "name": "main", "rows": 10,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 10, "missing_pct": 0.0},
                {"name": "Region", "role": "categorical", "distinct": 3,
                 "missing_pct": 0.0,
                 "top_values": [{"value": "Nord", "count": 5},
                                {"value": "Sud", "count": 5}]},
                {"name": "Age", "role": "quantitative", "distinct": 10,
                 "missing_pct": 0.0, "min": 18.0, "max": 65.0,
                 "mean": 35.0, "median": 33.0},
            ],
            "correlations": [], "duplicates": None,
        },
        "demographics": {
            "name": "demographics", "rows": 25,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 10,
                 "missing_pct": 0.0},
                {"name": "_root_id", "role": "linkage", "distinct": 10,
                 "missing_pct": 0.0},
                {"name": "nombre_menages", "role": "quantitative", "distinct": 10,
                 "missing_pct": 0.0, "min": 1.0, "max": 15.0,
                 "mean": 5.0, "median": 4.0},
                {"name": "nombre_habitants", "role": "quantitative", "distinct": 10,
                 "missing_pct": 0.0, "min": 1.0, "max": 50.0,
                 "mean": 12.0, "median": 10.0},
            ],
            "correlations": [], "duplicates": None,
        },
        "collaborations": {
            "name": "collaborations", "rows": 8,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 8,
                 "missing_pct": 0.0},
                {"name": "_root_id", "role": "linkage", "distinct": 8,
                 "missing_pct": 0.0},
                {"name": "organisation", "role": "categorical", "distinct": 6,
                 "missing_pct": 0.0,
                 "top_values": [{"value": "ONG A", "count": 3},
                                {"value": "ONG B", "count": 3},
                                {"value": "Autre", "count": 2}]},
            ],
            "correlations": [], "duplicates": None,
        },
    }


def _profile_xtf26_two_repeats_same_column():
    """Profile: ``main`` has no user columns; two repeat tables (``facilities``
    with 20 rows, ``staff`` with 5 rows) both carry ``groupe_socioeconomique``.
    Used to test the ambiguous multi-repeat case (most-rows wins)."""
    return {
        "main": {
            "name": "main", "rows": 5,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 5, "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        },
        "facilities": {
            "name": "facilities", "rows": 20,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 5,
                 "missing_pct": 0.0},
                {"name": "groupe_socioeconomique", "role": "categorical",
                 "distinct": 3, "missing_pct": 0.0,
                 "top_values": [{"value": "A", "count": 10}, {"value": "B", "count": 10}]},
            ],
            "correlations": [], "duplicates": None,
        },
        "staff": {
            "name": "staff", "rows": 5,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 5,
                 "missing_pct": 0.0},
                {"name": "groupe_socioeconomique", "role": "categorical",
                 "distinct": 3, "missing_pct": 0.0,
                 "top_values": [{"value": "A", "count": 3}, {"value": "B", "count": 2}]},
            ],
            "correlations": [], "duplicates": None,
        },
    }


def _proposal_xtf26(kind, spec, name, confidence=_HIGH_CONF, token_index=0):
    """A Proposal in the shape ``infer_specs`` returns (no status yet)."""
    return {
        "token_index": token_index,
        "kind": kind,
        "spec": dict(spec),
        "name": name,
        "confidence": confidence,
        "reason": "proposed",
    }


# --------------------------------------------------------------------------- #
# AC1: column absent from main, present in exactly one repeat table
#      → source set to that repeat table name, status "ok"
# --------------------------------------------------------------------------- #

def test_annotate_sets_source_from_single_repeat_table():
    """XTF-26 AC1: when a placeholder's target column is absent from ``main``
    but present in exactly one repeat table, ``annotate_proposals`` must set
    ``source`` to that repeat table name and ``status`` to ``"ok"``.

    ``nombre_menages`` lives only in the ``demographics`` repeat table; the
    proposal's spec has no ``source``, so the current code defaults to ``"main"``
    and rejects it with ``needs_attention``. The fix must auto-resolve it.
    """
    profile = _profile_xtf26_single_repeat()
    proposals = [
        _proposal_xtf26(
            "chart",
            {"name": "menages_chart", "title": "Nombre de ménages",
             "type": "bar", "questions": ["nombre_menages"]},
            name="menages_chart",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "ok", (
        f"Expected status 'ok' but got '{_get(out[0], 'status')}'; "
        f"reason: {_get(out[0], 'reason')!r}. "
        "annotate_proposals must auto-resolve single-repeat-table columns."
    )
    source = _get(out[0], "spec").get("source")
    assert source == "demographics", (
        f"Expected source 'demographics' but got {source!r}. "
        "annotate_proposals must set source to the repeat table name."
    )


# --------------------------------------------------------------------------- #
# AC2: column present in multiple repeat tables
#      → source set to the table with most rows, status "review",
#        note mentions both alternatives
# --------------------------------------------------------------------------- #

def test_annotate_sets_source_review_for_ambiguous_repeat():
    """XTF-26 AC2: when a column is present in multiple repeat tables,
    ``annotate_proposals`` must set ``source`` to the repeat table with the most
    rows, ``status`` to ``"review"``, and include a note listing the alternative
    table names.

    ``groupe_socioeconomique`` is in ``facilities`` (20 rows) and ``staff``
    (5 rows). The fix must pick ``facilities`` as source, set ``status: review``,
    and mention both table names in the note.
    """
    profile = _profile_xtf26_two_repeats_same_column()
    proposals = [
        _proposal_xtf26(
            "chart",
            {"name": "socio_chart", "title": "Groupe socio-économique",
             "type": "bar", "questions": ["groupe_socioeconomique"]},
            name="socio_chart",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "review", (
        f"Expected status 'review' for ambiguous multi-repeat column but got "
        f"'{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}."
    )
    source = _get(out[0], "spec").get("source")
    assert source == "facilities", (
        f"Expected source 'facilities' (most rows=20) but got {source!r}. "
        "annotate_proposals must pick the repeat table with the most rows."
    )
    # The note/reason must mention both candidate table names.
    note = _get(out[0], "reason") or ""
    assert "facilities" in note and "staff" in note, (
        f"Expected reason to mention both 'facilities' and 'staff' but got: {note!r}"
    )


# --------------------------------------------------------------------------- #
# AC3: column absent from main and all repeat tables → status remains needs_attention
# --------------------------------------------------------------------------- #

def test_annotate_keeps_needs_attention_when_column_nowhere():
    """XTF-26 AC3: when a proposal's column is not found in ``main`` or any
    repeat table, ``status`` must remain ``"needs_attention"``.

    ``ghost_column`` does not exist in any table of the profile; the proposal
    must stay flagged after the fix (the fix must not swallow genuine misses).
    """
    profile = _profile_xtf26_single_repeat()
    proposals = [
        _proposal_xtf26(
            "chart",
            {"name": "ghost_chart", "title": "Ghost",
             "type": "bar", "questions": ["ghost_column"]},
            name="ghost_chart",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "needs_attention", (
        f"Expected status 'needs_attention' for a column absent from all tables "
        f"but got '{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}."
    )


# --------------------------------------------------------------------------- #
# AC4 (fixture): fixture profile with demographic + socioeconomic repeat columns
#   → each proposal gets correct source and status "ok"
# --------------------------------------------------------------------------- #

def test_annotate_resolves_known_repeat_columns_from_fixture_profile():
    """XTF-26 AC4: fixture profile contains demographic columns
    (``nombre_menages``, ``nombre_habitants``) in one repeat table and an
    organisation column (``organisation``) in a second repeat table. Each
    proposal must get the correct ``source`` and ``status == 'ok'``.

    This mirrors the real-world case described in the bug report: columns from
    different repeat groups all resolve cleanly to their respective tables.
    """
    profile = _profile_xtf26_single_repeat()

    proposals = [
        _proposal_xtf26(
            "chart",
            {"name": "menages", "title": "Nombre de ménages",
             "type": "bar", "questions": ["nombre_menages"]},
            name="menages", token_index=0,
        ),
        _proposal_xtf26(
            "chart",
            {"name": "habitants", "title": "Nombre d'habitants",
             "type": "histogram", "questions": ["nombre_habitants"]},
            name="habitants", token_index=1,
        ),
        _proposal_xtf26(
            "chart",
            {"name": "organisations", "title": "Organisations",
             "type": "bar", "questions": ["organisation"]},
            name="organisations", token_index=2,
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    # nombre_menages and nombre_habitants → demographics
    assert _get(out[0], "status") == "ok", (
        f"nombre_menages: expected ok, got '{_get(out[0], 'status')}' "
        f"reason={_get(out[0], 'reason')!r}"
    )
    assert _get(out[0], "spec").get("source") == "demographics", (
        f"nombre_menages: expected source='demographics', "
        f"got {_get(out[0], 'spec').get('source')!r}"
    )

    assert _get(out[1], "status") == "ok", (
        f"nombre_habitants: expected ok, got '{_get(out[1], 'status')}' "
        f"reason={_get(out[1], 'reason')!r}"
    )
    assert _get(out[1], "spec").get("source") == "demographics", (
        f"nombre_habitants: expected source='demographics', "
        f"got {_get(out[1], 'spec').get('source')!r}"
    )

    # organisation → collaborations
    assert _get(out[2], "status") == "ok", (
        f"organisation: expected ok, got '{_get(out[2], 'status')}' "
        f"reason={_get(out[2], 'reason')!r}"
    )
    assert _get(out[2], "spec").get("source") == "collaborations", (
        f"organisation: expected source='collaborations', "
        f"got {_get(out[2], 'spec').get('source')!r}"
    )


# =========================================================================== #
# XTF-28 — split_value placeholder inference
# =========================================================================== #
# Fix: infer_specs must recognize a "split_value" kind, informed by the config's
# split_by dimension, so a short-label token that clearly refers to the unit of
# analysis (e.g. "[[NOM]]", "[[Commune]]") is proposed as kind="split_value"
# instead of a low-confidence indicator/needs_attention placeholder.
#
# AC1: infer_specs proposes kind == "split_value" for a NOM-like token when
#      config has split_by: Commune.
# AC2: annotate_proposals marks a split_value proposal "ok" when config split_by
#      is set, "needs_attention" when it is not.
# AC3: apply_inference (called by apply-template) writes the literal
#      "{{ split_value }}" placeholder for an accepted split_value proposal.

def test_infer_specs_proposes_split_value_for_nom_token(monkeypatch):
    """AC1: with split_by: Commune configured and a [[NOM]] placeholder token,
    infer_specs returns at least one proposal with kind == "split_value".

    This pins the fix's actual mechanism (card body): "When the config has a
    split_by dimension, include it in the LLM prompt so the model can propose
    kind: split_value". So the prompt variables built by infer_specs must
    surface the split_by dimension (e.g. under a "split_by" variable key) when
    ai_cfg carries one — a real LLM can only propose "split_value" if it is
    told what the split dimension is. Without that wiring the split_by value
    never reaches get_prompt's variables, which this test catches directly.
    """
    captured_variables = {}

    def _fake_get_prompt(name, variables):
        captured_variables.update(variables or {})
        return ([{"role": "user", "content": "x"}], {})

    monkeypatch.setattr(ti.lf_client, "get_prompt", _fake_get_prompt)

    def _fake_chat(*a, **k):
        # Simulates the LLM correctly using the split_by context supplied in
        # the prompt variables to propose kind: split_value for the NOM token.
        return (
            '{"proposals": ['
            '{"token_index": 0, "kind": "split_value", "name": "split_value", '
            '"spec": {}, "confidence": 0.95, "reason": "NOM maps to split_by Commune"}'
            ']}'
        )

    monkeypatch.setattr(ti.lf_client, "chat", _fake_chat)

    nl_tokens = [
        ti.Token(raw="[[NOM]]", inner="NOM", delimiter="[[", kind="nl",
                  location=ti.Location()),
    ]
    catalog = ask_engine.build_catalog(_profile_xtf2())
    ai_cfg = {
        "provider": "openai", "model": "gpt-x", "api_key": "sk-test",
        "split_by": "Commune",
    }

    out = ti.infer_specs(nl_tokens, catalog, ai_cfg)

    # The split_by dimension configured on ai_cfg must reach the prompt so the
    # LLM has the context needed to propose kind: split_value.
    joined_vars = " ".join(str(v) for v in captured_variables.values())
    assert "Commune" in joined_vars, (
        f"split_by dimension 'Commune' was not included in the prompt variables "
        f"passed to get_prompt; variables={captured_variables}"
    )

    kinds = [_get(p, "kind") for p in out]
    assert "split_value" in kinds, (
        f"expected a 'split_value' proposal for the NOM token, got kinds={kinds}"
    )


def test_annotate_split_value_ok_when_split_by_set():
    """AC2: a proposal with kind: split_value gets status "ok" when the config
    has split_by set, and "needs_attention" when split_by is NOT set.

    annotate_proposals takes a profile dict as its second argument; the
    split_by flag must be reachable from that same call so the implementation
    can validate against it (e.g. via a "split_by" key in the profile/context
    passed to annotate_proposals, per the card's fix description).
    """
    proposal_ok = [
        _proposal("split_value", {}, name="split_value", confidence=_HIGH_CONF),
    ]
    profile_with_split_by = dict(_profile_xtf2())
    profile_with_split_by["split_by"] = "Commune"

    out_ok = ti.annotate_proposals(proposal_ok, profile_with_split_by)
    assert _get(out_ok[0], "status") == "ok", (
        f"expected 'ok' when split_by is configured, got "
        f"{_get(out_ok[0], 'status')!r} reason={_get(out_ok[0], 'reason')!r}"
    )

    proposal_missing = [
        _proposal("split_value", {}, name="split_value", confidence=_HIGH_CONF),
    ]
    profile_without_split_by = _profile_xtf2()  # no split_by key at all

    out_missing = ti.annotate_proposals(proposal_missing, profile_without_split_by)
    assert _get(out_missing[0], "status") == "needs_attention", (
        f"expected 'needs_attention' when split_by is NOT configured, got "
        f"{_get(out_missing[0], 'status')!r}"
    )


def test_apply_template_writes_split_value_placeholder(tmp_path):
    """AC3: an accepted split_value proposal causes apply_inference to write the
    literal "{{ split_value }}" placeholder into the resolved .docx (so
    build-report's existing split_value substitution fills it with the real
    commune name instead of leaving the literal "NOM" text)."""
    template = _docx_with_nl_placeholders(tmp_path, ["[[NOM]]"])
    tokens = ti.extract_placeholders(template)
    assert len(tokens) == 1

    approved = [
        _approved("split_value", {}, name="split_value", token_index=0),
    ]
    cfg = {"api": {}, "form": {}, "report": {"split_by": "Commune"}}

    _cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    expected = "{{ split_value }}"
    reopened = Document(str(resolved))
    found = any(
        expected in "".join(r.text for r in p.runs) for p in reopened.paragraphs
    )
    assert found, (
        "resolved template does not contain the literal '{{ split_value }}' "
        "placeholder for the accepted split_value proposal"
    )


# =========================================================================== #
# MNT-26 — split_by placeholder inference (mirrors XTF-28's split_value case)
# =========================================================================== #
# `{{ split_value }}` (the split dimension's VALUE) is already recognized by
# Express Fill's inference; `{{ split_by }}` (the split dimension's NAME,
# e.g. "Region") must be recognized the same way.
#
# AC: Express Fill's inference recognizes a placeholder naming the split
#     dimension (e.g. "Split by", "Grouping") as the split_by literal kind
#     and resolves it to canonical {{ split_by }}, warning if no split_by
#     dimension is configured — mirroring split_value's behavior exactly.

def test_infer_specs_proposes_split_by_for_grouping_token(monkeypatch):
    """infer_specs returns at least one proposal with kind == "split_by" for
    a token naming the split dimension itself (e.g. "[[Split by]]"), when
    ai_cfg carries a configured split_by dimension. Mirrors
    test_infer_specs_proposes_split_value_for_nom_token."""
    captured_variables = {}

    def _fake_get_prompt(name, variables):
        captured_variables.update(variables or {})
        return ([{"role": "user", "content": "x"}], {})

    monkeypatch.setattr(ti.lf_client, "get_prompt", _fake_get_prompt)

    def _fake_chat(*a, **k):
        return (
            '{"proposals": ['
            '{"token_index": 0, "kind": "split_by", "name": "split_by", '
            '"spec": {}, "confidence": 0.95, "reason": "labels the split dimension"}'
            ']}'
        )

    monkeypatch.setattr(ti.lf_client, "chat", _fake_chat)

    nl_tokens = [
        ti.Token(raw="[[Split by]]", inner="Split by", delimiter="[[", kind="nl",
                  location=ti.Location()),
    ]
    catalog = ask_engine.build_catalog(_profile_xtf2())
    ai_cfg = {
        "provider": "openai", "model": "gpt-x", "api_key": "sk-test",
        "split_by": "Region",
    }

    out = ti.infer_specs(nl_tokens, catalog, ai_cfg)

    joined_vars = " ".join(str(v) for v in captured_variables.values())
    assert "Region" in joined_vars, (
        f"split_by dimension 'Region' was not included in the prompt variables "
        f"passed to get_prompt; variables={captured_variables}"
    )

    # "split_by" must be one of the valid kinds offered to the LLM (mirroring
    # how "split_value" is already an enumerated kind) — otherwise the model
    # is never told this is a legal choice, regardless of what a stubbed
    # response happens to contain.
    assert "split_by" in str(captured_variables.get("kinds", "")), (
        "'split_by' is not among the valid kinds surfaced to the LLM in the "
        f"prompt variables; kinds variable={captured_variables.get('kinds')!r}"
    )

    kinds = [_get(p, "kind") for p in out]
    assert "split_by" in kinds, (
        f"expected a 'split_by' proposal for the grouping-label token, got kinds={kinds}"
    )


def test_annotate_split_by_ok_when_split_by_set():
    """A proposal with kind: split_by gets status "ok" when the config has
    split_by set, and "needs_attention" when split_by is NOT set — mirroring
    test_annotate_split_value_ok_when_split_by_set exactly."""
    proposal_ok = [
        _proposal("split_by", {}, name="split_by", confidence=_HIGH_CONF),
    ]
    profile_with_split_by = dict(_profile_xtf2())
    profile_with_split_by["split_by"] = "Region"

    out_ok = ti.annotate_proposals(proposal_ok, profile_with_split_by)
    assert _get(out_ok[0], "status") == "ok", (
        f"expected 'ok' when split_by is configured, got "
        f"{_get(out_ok[0], 'status')!r} reason={_get(out_ok[0], 'reason')!r}"
    )

    proposal_missing = [
        _proposal("split_by", {}, name="split_by", confidence=_HIGH_CONF),
    ]
    profile_without_split_by = _profile_xtf2()  # no split_by key at all

    out_missing = ti.annotate_proposals(proposal_missing, profile_without_split_by)
    assert _get(out_missing[0], "status") == "needs_attention", (
        f"expected 'needs_attention' when split_by is NOT configured, got "
        f"{_get(out_missing[0], 'status')!r}"
    )


def test_apply_template_writes_split_by_placeholder(tmp_path):
    """An accepted split_by proposal causes apply_inference to write the
    literal "{{ split_by }}" placeholder into the resolved .docx — mirroring
    test_apply_template_writes_split_value_placeholder exactly."""
    template = _docx_with_nl_placeholders(tmp_path, ["[[Split by]]"])
    tokens = ti.extract_placeholders(template)
    assert len(tokens) == 1

    approved = [
        _approved("split_by", {}, name="split_by", token_index=0),
    ]
    cfg = {"api": {}, "form": {}, "report": {"split_by": "Region"}}

    _cfg_out, resolved = ti.apply_inference(approved, cfg, template)

    expected = "{{ split_by }}"
    reopened = Document(str(resolved))
    found = any(
        expected in "".join(r.text for r in p.runs) for p in reopened.paragraphs
    )
    assert found, (
        "resolved template does not contain the literal '{{ split_by }}' "
        "placeholder for the accepted split_by proposal"
    )


# =========================================================================== #
# MNT-32 — Express-inferred `list` on repeat-group data is auto-resolved
# =========================================================================== #
# Bug: `_DATA_KINDS` in template_inference excludes "list", so
# `_autoresolve_repeat_source` bails for a `list` proposal whose column lives in
# a repeat-group base table. `annotate_proposals` then validates the list against
# `main` only, the column isn't there, and it is flagged `needs_attention` and
# dropped — even though the IDENTICAL column validates clean as a `table`.
#
# Fix (derived strictly from the MNT-32 Acceptance criteria): a `list` proposal
# must auto-resolve its repeat-group `source` EXACTLY like a `table` proposal:
#   * column absent from `main`, present in exactly ONE repeat table
#       -> status "ok", spec["source"] stamped to that repeat table.
#   * column present in MULTIPLE repeat tables
#       -> status "review", spec["source"] = largest by rows, note lists alts.
#   * column absent from every table
#       -> status remains "needs_attention" (no false-positive resolution).
#   * column already in `main`
#       -> unchanged (regression-pinned).
#
# These call `annotate_proposals` directly, mirroring the XTF-26 tests but with
# kind="list". Expected RED until "list" is added to `_DATA_KINDS`.
# =========================================================================== #


def _profile_mnt32_single_repeat():
    """Profile: ``main`` has Region; a single ``hh_members`` repeat table holds
    ``Village`` (the natural list case -- a list of village/member names)."""
    return {
        "main": {
            "name": "main", "rows": 10,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 10, "missing_pct": 0.0},
                {"name": "Region", "role": "categorical", "distinct": 2,
                 "missing_pct": 0.0,
                 "top_values": [{"value": "N", "count": 6},
                                {"value": "S", "count": 4}]},
            ],
            "correlations": [], "duplicates": None,
        },
        "hh_members": {
            "name": "hh_members", "rows": 30,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 10,
                 "missing_pct": 0.0},
                {"name": "_root_id", "role": "linkage", "distinct": 10,
                 "missing_pct": 0.0},
                {"name": "Village", "role": "qualitative", "distinct": 12,
                 "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        },
    }


def _profile_mnt32_two_repeats_same_column():
    """Profile: ``main`` has no user columns; two repeat tables (``hh_members``
    with 30 rows, ``visitors`` with 6 rows) both carry ``Village``. The largest
    (``hh_members``) must win with status "review"."""
    return {
        "main": {
            "name": "main", "rows": 5,
            "columns": [
                {"name": "_id", "role": "linkage", "distinct": 5, "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        },
        "hh_members": {
            "name": "hh_members", "rows": 30,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 5,
                 "missing_pct": 0.0},
                {"name": "Village", "role": "qualitative", "distinct": 12,
                 "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        },
        "visitors": {
            "name": "visitors", "rows": 6,
            "columns": [
                {"name": "_parent_index", "role": "linkage", "distinct": 5,
                 "missing_pct": 0.0},
                {"name": "Village", "role": "qualitative", "distinct": 4,
                 "missing_pct": 0.0},
            ],
            "correlations": [], "duplicates": None,
        },
    }


# --------------------------------------------------------------------------- #
# AC1: a `list` proposal whose only referenced column lives in a single
#      repeat-group base table (not in `main`) is auto-resolved:
#      status "ok" with the repeat table stamped onto spec["source"].
# --------------------------------------------------------------------------- #
def test_annotate_list_autoresolves_single_repeat_source():
    """MNT-32 AC1: a `list` proposal on a repeat-only column must resolve to
    status "ok" with spec["source"] stamped to the repeat table -- not be dropped
    as needs_attention. ``Village`` lives only in the ``hh_members`` repeat table.
    """
    profile = _profile_mnt32_single_repeat()
    proposals = [
        _proposal_xtf26(
            "list",
            {"name": "villages", "title": "Villages", "question": "Village"},
            name="villages",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "ok", (
        f"Expected 'ok' for a list on a single-repeat-table column but got "
        f"'{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}. "
        "A list must auto-resolve its repeat-group source exactly like a table."
    )
    assert _get(out[0], "spec").get("source") == "hh_members", (
        f"Expected source 'hh_members' stamped onto the list spec, got "
        f"{_get(out[0], 'spec').get('source')!r}"
    )


def test_annotate_list_autoresolves_repeat_source_from_questions_plural():
    """MNT-32 (Files note): `_referenced_columns` must read a list spec's column
    whether it is stored as ``question`` (singular) or ``questions`` (plural). A
    list spec that carries only ``questions: ["Village"]`` must resolve the same
    repeat-group source as the singular form."""
    profile = _profile_mnt32_single_repeat()
    proposals = [
        _proposal_xtf26(
            "list",
            {"name": "villages", "title": "Villages", "questions": ["Village"]},
            name="villages",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "ok", (
        f"Expected 'ok' for a list spec using the plural 'questions' key, got "
        f"'{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}"
    )
    assert _get(out[0], "spec").get("source") == "hh_members", (
        f"Expected source 'hh_members', got "
        f"{_get(out[0], 'spec').get('source')!r}"
    )


# --------------------------------------------------------------------------- #
# AC (matches the table kind): the IDENTICAL `table` and `list` proposals on the
#      same repeat-only column resolve to the same source and status.
# --------------------------------------------------------------------------- #
def test_annotate_list_and_table_resolve_identically():
    """MNT-32 AC: the identical column validates clean as a `table` and must now
    validate identically as a `list` -- same status "ok", same stamped source.
    This is the crux of the bug (the user was funnelled onto the table path
    because only `table` auto-resolved the repeat source)."""
    profile = _profile_mnt32_single_repeat()

    table_out = ti.annotate_proposals(
        [_proposal_xtf26(
            "table",
            {"name": "villages", "title": "Villages", "questions": ["Village"]},
            name="villages")],
        profile,
    )
    list_out = ti.annotate_proposals(
        [_proposal_xtf26(
            "list",
            {"name": "villages", "title": "Villages", "question": "Village"},
            name="villages")],
        profile,
    )

    assert _get(table_out[0], "status") == "ok", (
        f"precondition: the table proposal should validate ok, got "
        f"{_get(table_out[0], 'status')!r} reason={_get(table_out[0], 'reason')!r}"
    )
    assert _get(list_out[0], "status") == _get(table_out[0], "status"), (
        f"list status {_get(list_out[0], 'status')!r} must match table status "
        f"{_get(table_out[0], 'status')!r}"
    )
    assert _get(list_out[0], "spec").get("source") == _get(table_out[0], "spec").get("source"), (
        f"list source {_get(list_out[0], 'spec').get('source')!r} must match "
        f"table source {_get(table_out[0], 'spec').get('source')!r}"
    )
    assert _get(list_out[0], "spec").get("source") == "hh_members"


# --------------------------------------------------------------------------- #
# AC2: multiple repeat tables hold the column -> status "review", source = the
#      largest by row count, note listing the alternatives (like the table kind).
# --------------------------------------------------------------------------- #
def test_annotate_list_multi_repeat_resolves_to_largest_with_review():
    """MNT-32 AC2: when multiple repeat tables hold the list column, the `list`
    proposal must resolve to the largest by rows with status "review" and a note
    listing the alternative table names -- identical to the table kind's
    multi-table behavior. ``Village`` is in ``hh_members`` (30) and ``visitors``
    (6); ``hh_members`` must win."""
    profile = _profile_mnt32_two_repeats_same_column()
    proposals = [
        _proposal_xtf26(
            "list",
            {"name": "villages", "title": "Villages", "question": "Village"},
            name="villages",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "review", (
        f"Expected 'review' for an ambiguous multi-repeat list column, got "
        f"'{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}"
    )
    assert _get(out[0], "spec").get("source") == "hh_members", (
        f"Expected source 'hh_members' (most rows=30), got "
        f"{_get(out[0], 'spec').get('source')!r}"
    )
    note = _get(out[0], "reason") or ""
    assert "hh_members" in note and "visitors" in note, (
        f"Expected the note to mention both 'hh_members' and 'visitors', got: {note!r}"
    )


# --------------------------------------------------------------------------- #
# AC3: column genuinely absent from every table -> still needs_attention.
# --------------------------------------------------------------------------- #
def test_annotate_list_keeps_needs_attention_when_column_nowhere():
    """MNT-32 AC3: a `list` proposal whose column is absent from `main` and every
    repeat table must still flag `needs_attention` -- the fix must not swallow a
    genuine miss (no false-positive resolution)."""
    profile = _profile_mnt32_single_repeat()
    proposals = [
        _proposal_xtf26(
            "list",
            {"name": "ghost", "title": "Ghost", "question": "NotAColumn"},
            name="ghost",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "needs_attention", (
        f"Expected 'needs_attention' for a truly-absent list column, got "
        f"'{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}"
    )


# --------------------------------------------------------------------------- #
# AC5 (regression): the `main`-table list case (column already in main) is
#      unchanged -- status "ok", source stays "main" (never stamped to a repeat).
# --------------------------------------------------------------------------- #
def test_annotate_list_on_main_column_unchanged():
    """MNT-32 AC5: a `list` whose column already lives in `main` must remain
    status "ok" and must NOT be re-sourced to a repeat table (regression pin)."""
    profile = _profile_mnt32_single_repeat()
    proposals = [
        _proposal_xtf26(
            "list",
            {"name": "regions", "title": "Regions", "question": "Region"},
            name="regions",
        ),
    ]

    out = ti.annotate_proposals(proposals, profile)

    assert _get(out[0], "status") == "ok", (
        f"Expected 'ok' for a list on a main-table column, got "
        f"'{_get(out[0], 'status')}'; reason: {_get(out[0], 'reason')!r}"
    )
    source = _get(out[0], "spec").get("source")
    assert source in (None, "main"), (
        f"A main-table list must not be re-sourced to a repeat table; got "
        f"source={source!r}"
    )
