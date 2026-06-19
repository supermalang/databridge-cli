"""API tests for the Express Template Fill web surface (XTF-5).

These exercise the two endpoints the Templates review/approve panel calls:
``POST /api/template/infer`` and ``POST /api/template/apply``. They mirror the
conventions in ``tests/test_ask_api.py`` — a synchronous ``TestClient`` over the
FastAPI app, auth disabled by default (the dev user has an active project via the
session-scoped conftest fixtures), and the LLM / inference seam monkeypatched so
no network or real provider is touched.

Every assertion is derived from the XTF-5 acceptance criteria + spec §6:
- infer returns the no-AI message payload when no provider/key is configured;
- infer returns the "run Download first" payload when there is no downloaded data;
- infer returns ``{proposals: [...]}`` (inference mocked) when AI + data are present;
- infer resolves an existing-template ref (not just a multipart upload);
- apply runs ``apply_inference``, writes config, and returns
  ``{ok, template, n_written}`` with the resolved template path.
"""
import io

import pandas as pd
import pytest
from docx import Document
from fastapi.testclient import TestClient

import web.main as wm


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def _df():
    return pd.DataFrame({"_id": [1, 2, 3], "Region": ["N", "E", "E"]})


def _ai_cfg():
    # provider + resolved api_key → AI is "configured" per spec §6 / ai_status.
    return {"ai": {"provider": "openai", "api_key": "sk-x"},
            "questions": [{"export_label": "Region", "category": "categorical"}]}


@pytest.fixture
def client():
    return TestClient(wm.app)


@pytest.fixture
def dev_active_project():
    """Guarantee the dev user has a real active project for the duration of the test.

    The XTF-8 apply/delete tests exercise the org/project resolution in require_role
    so apply can push the resolved .docx to durable storage. They must not depend on
    suite order: an earlier module (test_projects_api -> delete_project) sets
    User.active_project_id = None, tearing down the session-shared active project.

    This fixture re-establishes (or provisions) an active project per-test by setting
    User.active_project_id via the repository, restoring whatever it was afterwards.
    """
    with wm.db_session.SessionLocal() as db:
        dev = wm.db_repo.get_user_by_sub(db, "dev-local")
        if dev is None:
            dev = wm.db_provision.ensure_dev_user(db)
        prev = dev.active_project_id
        if dev.active_project_id is None:
            projects = wm.db_repo.list_projects_for_user(db, dev)
            if not projects:
                # No project at all (another module deleted it) — provision a fresh one.
                from web.db import bootstrap as _bootstrap
                proj = _bootstrap.import_legacy_config(db, owner=dev)
                if proj is None:
                    proj = wm.db_repo.create_project(db, user=dev, name="Test active project")
                projects = [proj]
            wm.db_repo.set_active_project(db, dev, projects[0].id)
        active_id = dev.active_project_id
    yield active_id
    # Restore the prior active_project_id so we don't perturb later tests.
    with wm.db_session.SessionLocal() as db:
        dev = wm.db_repo.get_user_by_sub(db, "dev-local")
        if dev is not None:
            dev.active_project_id = prev
            db.commit()


# --------------------------------------------------------------------------- #
# Precondition payloads (AC: friendly messages, never a crash)
# --------------------------------------------------------------------------- #
def test_infer_no_ai_provider_returns_configure_message(monkeypatch, client):
    """No AI provider/key configured → friendly "Configure an AI provider" payload.

    AC: "no AI provider/key → 'Configure an AI provider to use Express fill.'"
    """
    cfg = {"questions": []}  # no ai section at all
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "load_processed_data", lambda *a, **k: (_df(), {}))

    resp = client.post("/api/template/infer", json={"template": "report.docx"})
    assert resp.status_code == 200
    body = resp.json()
    assert body.get("proposals") in (None, [])
    assert "Configure an AI provider" in (body.get("message") or "")


def test_infer_no_data_returns_download_first_message(monkeypatch, client):
    """No downloaded data → friendly "run Download first" payload.

    AC: "no downloaded data → 'No data yet — run Download first.'"
    """
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: _ai_cfg())

    def _raise(*a, **k):
        raise FileNotFoundError("no data")

    monkeypatch.setattr(wm, "load_processed_data", _raise)

    resp = client.post("/api/template/infer", json={"template": "report.docx"})
    assert resp.status_code == 200
    body = resp.json()
    assert body.get("proposals") in (None, [])
    assert "Download" in (body.get("message") or "")


# --------------------------------------------------------------------------- #
# Happy path (AC: parse → infer → annotate → {proposals})
# --------------------------------------------------------------------------- #
def test_infer_returns_proposals_when_ai_and_data_present(monkeypatch, client, tmp_path):
    """AI configured + data present → {proposals: [...]} from parse→infer→annotate.

    AC: "returns {proposals: [...]} (LLM mocked) when AI + data are present."
    The inference seam (extract/infer/annotate in template_inference) is mocked so
    no network/provider is touched, mirroring the suggester-style mocks.
    """
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: _ai_cfg())
    monkeypatch.setattr(wm, "load_processed_data", lambda *a, **k: (_df(), {}))

    proposals = [
        {"token_index": 0, "kind": "chart", "name": "by_region",
         "spec": {"name": "by_region", "type": "bar", "questions": ["Region"]},
         "confidence": 0.9, "reason": "", "status": "ok"},
        {"token_index": 1, "kind": "indicator", "name": "n_rows",
         "spec": {"name": "n_rows", "stat": "count"},
         "confidence": 0.3, "reason": "low confidence", "status": "needs_attention"},
    ]

    import src.reports.template_inference as ti
    # Make a real (but trivial) template path resolvable; extract is mocked anyway.
    tpl = tmp_path / "report.docx"
    tpl.write_bytes(b"PK")  # contents irrelevant — extract is patched

    monkeypatch.setattr(ti, "extract_placeholders",
                        lambda *a, **k: [object(), object()])
    monkeypatch.setattr(ti, "infer_specs", lambda *a, **k: proposals)
    monkeypatch.setattr(ti, "annotate_proposals", lambda props, *a, **k: props)

    resp = client.post("/api/template/infer", json={"template": "report.docx"})
    assert resp.status_code == 200
    body = resp.json()
    names = [p["name"] for p in body["proposals"]]
    assert names == ["by_region", "n_rows"]
    assert body["proposals"][1]["status"] == "needs_attention"


def test_infer_resolves_existing_template_ref(monkeypatch, client, tmp_path):
    """An existing-template *ref* (not a multipart upload) resolves to the stored file.

    AC / Unit tests: "resolves an existing-template ref (not just a multipart upload)
    to the correct stored template." The endpoint must pass the resolved path of the
    named stored template into extract_placeholders.
    """
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: _ai_cfg())
    monkeypatch.setattr(wm, "load_processed_data", lambda *a, **k: (_df(), {}))

    # Stage a stored template under a throwaway templates dir the API serves from
    # (monkeypatched so the test never pollutes the repo's templates/ directory).
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(wm, "TEMPLATES_DIR", templates_dir)
    stored = templates_dir / "annual_report.docx"
    stored.write_bytes(b"PK")

    seen = {}
    import src.reports.template_inference as ti

    def _extract(path, *a, **k):
        seen["path"] = str(path)
        return []

    monkeypatch.setattr(ti, "extract_placeholders", _extract)
    monkeypatch.setattr(ti, "infer_specs", lambda *a, **k: [])
    monkeypatch.setattr(ti, "annotate_proposals", lambda props, *a, **k: props)

    resp = client.post("/api/template/infer", json={"template": "annual_report.docx"})
    assert resp.status_code == 200
    # The endpoint must have resolved the ref to the actual stored file path.
    assert seen.get("path") == str(stored)


# --------------------------------------------------------------------------- #
# Apply (AC: writes config + returns {ok, template, n_written})
# --------------------------------------------------------------------------- #
def test_apply_writes_config_and_returns_resolved_template(monkeypatch, client, tmp_path):
    """Approved proposals → apply_inference writes config and returns the resolved path.

    AC: "POST /api/template/apply {proposals} runs apply_inference and returns
    {ok, template, n_written}; ... (resolved template path)."
    """
    cfg = {"charts": []}
    saved = {}
    # /api/template/apply is a config mutation and gates on _require("editor") like
    # every other config-writing endpoint. RBAC enforcement is covered by the shared
    # require_role / config tests; here we isolate the endpoint's behavior from the
    # (order-dependent) active-project setup so this unit test is deterministic.
    monkeypatch.setattr(wm, "_require", lambda *a, **k: None)
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "write_config", lambda c, p: saved.update(c))

    resolved = str(tmp_path / "report.resolved.docx")
    approved = [
        {"token_index": 0, "kind": "chart", "name": "by_region",
         "spec": {"name": "by_region", "type": "bar", "questions": ["Region"]},
         "status": "ok"},
    ]

    import src.reports.template_inference as ti

    def _apply(approved_props, cfg_in, template_path):
        cfg_in.setdefault("charts", []).append(approved_props[0]["spec"])
        return cfg_in, resolved

    monkeypatch.setattr(ti, "apply_inference", _apply)

    resp = client.post("/api/template/apply",
                       json={"proposals": approved, "template": "report.docx"})
    assert resp.status_code == 200
    body = resp.json()
    assert body["ok"] is True
    assert body["template"] == resolved
    assert body["n_written"] == 1
    # Config was persisted with the approved chart spec.
    assert saved["charts"][0]["name"] == "by_region"


# --------------------------------------------------------------------------- #
# XTF-6 — Persist the uploaded template across infer → apply (the bug).
#
# These are deliberately UN-mocked at the inference layer: only the LLM seam
# (`infer_specs`) is stubbed with a deterministic proposal. `extract_placeholders`,
# the upload-persistence path, and `apply_inference` all run for real, so they
# reproduce the production bug — an uploaded .docx that infer never persists nor
# returns a ref for, so apply cannot resolve it.
# --------------------------------------------------------------------------- #
def _real_docx_bytes(placeholder="[bar chart of Region]"):
    """A minimal real .docx (python-docx) carrying one NL placeholder."""
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    doc.add_paragraph(placeholder)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_infer_apply_roundtrip_real(monkeypatch, api_client, tmp_path, dev_active_project):
    """Real infer→apply round-trip: the upload must survive into apply.

    AC (XTF-6, contract updated by XTF-8):
    - infer persists the uploaded .docx and returns a resolvable `template` ref;
    - the same ref is carried into apply, which resolves the persisted file and
      runs the REAL `apply_inference` against it;
    - apply returns a RELATIVE templates/<name>.resolved.docx ref, the physical
      .docx exists at TEMPLATES_DIR, config gained the chart spec, and the
      response is {ok, template, n_written}.

    Only the LLM seam (`infer_specs`) is mocked; `extract_placeholders` and
    `apply_inference` run for real. The dev user has a real active project (the
    `dev_active_project` fixture) — NOT a mocked-away `_require` — because the
    relative-ref / storage-push branch in apply only fires once require_role
    resolves a real org/project; mocking `_require` would silently keep the old
    absolute-path behavior and make the test order-dependent.
    """
    client = api_client
    import src.reports.template_inference as ti

    # Isolate template storage so we never touch the repo's templates/ dir.
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(wm, "TEMPLATES_DIR", templates_dir)

    # Shared config object: infer reads it; apply reads + writes it.
    cfg = {"charts": [], "questions": [{"export_label": "Region", "category": "categorical"}],
           "ai": {"provider": "openai", "api_key": "sk-x"}}
    saved = {}
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "write_config", lambda c, p: saved.update({"cfg": c}))
    monkeypatch.setattr(wm, "load_processed_data", lambda *a, **k: (_df(), {}))
    # Keep profiling/catalog cheap + deterministic (they feed infer_specs, which is mocked).
    monkeypatch.setattr(wm, "profile_dataset", lambda *a, **k: {})
    monkeypatch.setattr(wm.ask_engine, "build_catalog", lambda *a, **k: {})

    # The ONLY mocked seam: the batched LLM call. One real chart proposal for the
    # single NL placeholder (token_index 0) in the uploaded docx.
    proposal = {
        "token_index": 0, "kind": "chart", "name": "by_region",
        "spec": {"name": "by_region", "title": "By region", "type": "bar",
                 "questions": ["Region"]},
        "confidence": 0.95, "reason": "ok", "status": "ok",
    }
    monkeypatch.setattr(ti, "infer_specs", lambda *a, **k: [dict(proposal)])
    # annotate runs for real (validate_recipe is reused) — but with an empty profile
    # it would flag the chart; stamp it ok directly so the proposal is approvable.
    monkeypatch.setattr(ti, "annotate_proposals",
                        lambda props, *a, **k: [{**p, "status": "ok", "reason": "ok"} for p in props])

    # 1. POST the real .docx as a multipart upload to infer.
    files = {"file": ("fresh_upload.docx", _real_docx_bytes(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    infer_resp = client.post("/api/template/infer", files=files)
    assert infer_resp.status_code == 200, infer_resp.text
    infer_body = infer_resp.json()
    assert infer_body.get("proposals"), "infer should return the mocked proposal"

    # AC: infer must return a resolvable template ref alongside the proposals.
    ref = infer_body.get("template")
    assert ref, "infer must return a 'template' ref so apply can resolve the upload"

    # 2. Carry that ref (NOT a bare client file.name) into the REAL apply.
    approved = [dict(p) for p in infer_body["proposals"]]
    apply_resp = client.post("/api/template/apply",
                             json={"proposals": approved, "template": ref})
    assert apply_resp.status_code == 200, apply_resp.text
    apply_body = apply_resp.json()

    # AC: response shape {ok, template, n_written}.
    assert apply_body.get("ok") is True
    assert apply_body.get("n_written") == 1
    resolved = apply_body.get("template")
    assert resolved, "apply must return the resolved template path"

    # AC (updated by XTF-8): apply returns a RELATIVE templates/<name>.resolved.docx
    # ref (not an absolute host-mirror path), and the physical file lives at
    # TEMPLATES_DIR / <name>.resolved.docx (pushed to storage for later runs).
    import os
    from pathlib import Path as _Path
    assert not os.path.isabs(resolved), \
        f"resolved template must be relative, got absolute: {resolved!r}"
    assert _Path(resolved).parts[0] == "templates", \
        f"resolved template must live under templates/, got: {resolved!r}"
    assert ".." not in _Path(resolved).parts, \
        f"resolved template must not contain '..': {resolved!r}"
    on_disk = templates_dir / _Path(resolved).name
    assert on_disk.is_file(), \
        f"resolved template should exist on disk at TEMPLATES_DIR: {on_disk}"

    # AC: config gained the chart spec.
    written = saved.get("cfg") or cfg
    chart_names = [c.get("name") for c in (written.get("charts") or [])]
    assert "by_region" in chart_names, "apply should have written the chart spec into config"


def test_apply_unresolvable_ref_returns_clear_error(monkeypatch, client, tmp_path):
    """Apply with a ref that resolves to no stored file → a clear error, not a 500.

    AC (XTF-6): "if the ref cannot be resolved it returns a clear error (no
    traceback / no silent wrong-path)." The REAL apply_inference is used; the bug
    today is that an unresolvable basename is passed straight through to
    apply_inference, which then fails opening a non-existent path with a 500.
    """
    import src.reports.template_inference as ti  # noqa: F401  (kept un-mocked)

    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(wm, "TEMPLATES_DIR", templates_dir)

    cfg = {"charts": []}
    monkeypatch.setattr(wm, "_require", lambda *a, **k: None)
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "write_config", lambda c, p: None)

    approved = [
        {"token_index": 0, "kind": "chart", "name": "by_region",
         "spec": {"name": "by_region", "type": "bar", "questions": ["Region"]},
         "status": "ok"},
    ]
    resp = client.post("/api/template/apply",
                       json={"proposals": approved, "template": "does_not_exist.docx"})

    # A clear, client-facing error — not an unhandled 500 traceback.
    assert resp.status_code in (400, 404, 422), \
        f"expected a clear client error, got {resp.status_code}: {resp.text}"
    body = resp.json()
    detail = (body.get("detail") or body.get("message") or "")
    assert detail, "the error response should carry a human-readable message"


# --------------------------------------------------------------------------- #
# XTF-8 — Apply must persist the resolved template to durable storage AND write
# a RELATIVE report.template ref (the shape set_active_template already writes),
# so a subsequent web run's hydrate_run_dir pulls the resolved .docx from Minio
# into its isolated tempdir instead of reading a stale absolute host-mirror path.
#
# These run apply for real (only the LLM `infer_specs` seam is mocked) against a
# real uploaded/stored .docx, and they do NOT mock `_require` / `require_role`:
# the dev user has a real active project (conftest), so apply can resolve the
# org/project it needs to push the resolved file to storage — exactly the seam
# under test.
# --------------------------------------------------------------------------- #
import os  # noqa: E402
from pathlib import Path  # noqa: E402

from web.storage import workspace as _ws  # noqa: E402


def test_apply_persists_relative_template(monkeypatch, api_client, tmp_path, dev_active_project):
    """Apply writes a RELATIVE templates/… ref AND pushes the resolved .docx to storage.

    AC (XTF-8):
    - after /api/template/apply, cfg["report"]["template"] is a RELATIVE path of the
      form templates/<name>.resolved.docx (no absolute path, no "..");
    - the resolved .docx is pushed to durable storage via
      put_project_file(... "templates" ...) so a later run's hydrate_run_dir pulls it.

    Today this is RED: apply sets report.template to the ABSOLUTE resolved path
    returned by apply_inference and never calls put_project_file.
    """
    import src.reports.template_inference as ti

    # Isolate template storage so we never touch the repo's templates/ dir.
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(wm, "TEMPLATES_DIR", templates_dir)

    cfg = {"charts": [],
           "questions": [{"export_label": "Region", "category": "categorical"}],
           "ai": {"provider": "openai", "api_key": "sk-x"}}
    saved = {}
    # Real active project (no _require / require_role mocking) so apply can resolve
    # the org/project it pushes under. Only config IO is stubbed for determinism.
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "write_config", lambda c, p: saved.update({"cfg": c}))
    monkeypatch.setattr(wm, "load_processed_data", lambda *a, **k: (_df(), {}))
    monkeypatch.setattr(wm, "profile_dataset", lambda *a, **k: {})
    monkeypatch.setattr(wm.ask_engine, "build_catalog", lambda *a, **k: {})

    # Spy on the durable-storage push. Wrap the real function so the file still lands
    # in the local backend, but record every (org, project, category, path) call.
    pushes = []
    _real_put = _ws.put_project_file

    def _spy_put(org_id, project_id, category, local_path):
        pushes.append({"org_id": org_id, "project_id": project_id,
                       "category": category, "path": str(local_path)})
        return _real_put(org_id, project_id, category, local_path)

    monkeypatch.setattr(wm.storage_workspace, "put_project_file", _spy_put)

    # One real chart proposal for the single NL placeholder; LLM seam mocked only.
    proposal = {
        "token_index": 0, "kind": "chart", "name": "by_region",
        "spec": {"name": "by_region", "title": "By region", "type": "bar",
                 "questions": ["Region"]},
        "confidence": 0.95, "reason": "ok", "status": "ok",
    }
    monkeypatch.setattr(ti, "infer_specs", lambda *a, **k: [dict(proposal)])
    monkeypatch.setattr(ti, "annotate_proposals",
                        lambda props, *a, **k: [{**p, "status": "ok", "reason": "ok"} for p in props])

    # Infer the real uploaded .docx so apply resolves a genuinely-stored file.
    files = {"file": ("fresh_upload.docx", _real_docx_bytes(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    infer_resp = api_client.post("/api/template/infer", files=files)
    assert infer_resp.status_code == 200, infer_resp.text
    infer_body = infer_resp.json()
    ref = infer_body.get("template")
    assert ref, "infer must return a resolvable template ref"
    approved = [dict(p) for p in infer_body["proposals"]]

    # Drop the infer-time uploads so we only inspect the apply-time push below.
    pushes.clear()

    apply_resp = api_client.post("/api/template/apply",
                             json={"proposals": approved, "template": ref})
    assert apply_resp.status_code == 200, apply_resp.text
    apply_body = apply_resp.json()
    assert apply_body.get("ok") is True
    assert apply_body.get("n_written") == 1

    # (a) report.template in the WRITTEN config is a RELATIVE templates/… ref.
    written = saved.get("cfg") or cfg
    tref = (written.get("report") or {}).get("template")
    assert tref, "apply must write report.template"
    assert not os.path.isabs(tref), f"report.template must be relative, got absolute: {tref!r}"
    assert ".." not in Path(tref).parts, f"report.template must not contain '..': {tref!r}"
    assert Path(tref).parts[0] == "templates", \
        f"report.template must live under templates/, got: {tref!r}"
    assert Path(tref).name.endswith(".resolved.docx"), \
        f"report.template must point at the resolved .docx, got: {tref!r}"
    # The response 'template' is the same relative ref.
    assert apply_body.get("template") == tref

    # (b) the resolved .docx was pushed to durable storage under category 'templates'.
    template_pushes = [p for p in pushes if p["category"] == "templates"]
    assert template_pushes, \
        "apply must put_project_file the resolved .docx under category 'templates'"
    pushed_names = [Path(p["path"]).name for p in template_pushes]
    assert any(n.endswith(".resolved.docx") for n in pushed_names), \
        f"the pushed file must be the resolved .docx, pushed: {pushed_names}"
    # The pushed filename matches the relative ref written into config.
    assert Path(tref).name in pushed_names, \
        f"pushed file {pushed_names} must match report.template {tref!r}"


def test_sanitize_run_config_keeps_relative_template():
    """sanitize_run_config leaves a relative report.template intact.

    AC (XTF-8): sanitize_run_config does not blank or absolutize the relative
    report.template, so a hydrated run resolves the same file build-report loads.

    NOTE: this is a REGRESSION GUARD — sanitize_run_config currently does not
    touch report.template, so this may already pass today. It locks the AC that
    sanitize must keep the relative ref intact when XTF-8 changes that function.
    """
    cfg = {"report": {"template": "templates/x.resolved.docx", "output_dir": "x"},
           "export": {"output_dir": "y"}}
    out = _ws.sanitize_run_config(cfg)
    assert out["report"]["template"] == "templates/x.resolved.docx", \
        "sanitize must keep the relative report.template unchanged"
    # And it must not mutate the caller's dict.
    assert cfg["report"]["template"] == "templates/x.resolved.docx"


def test_delete_active_template_clears_ref(monkeypatch, api_client, tmp_path, dev_active_project):
    """Deleting the template referenced by report.template clears/repoints the ref.

    AC (XTF-8): delete_template on the file currently referenced by report.template
    clears or repoints the ref — no dangling path left in config.

    Today this is RED: delete_template unlinks the file and removes it from storage
    but never touches report.template, leaving a dangling reference.
    """
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(wm, "TEMPLATES_DIR", templates_dir)
    target = templates_dir / "foo.docx"
    target.write_bytes(b"PK")

    # Active config references the template we're about to delete.
    cfg = {"report": {"template": "templates/foo.docx"}}
    saved = {}

    # delete_template reads/writes the file at wm.CONFIG_PATH (async aiofiles). Point
    # it at a throwaway config so we observe the post-delete report.template.
    cfg_path = tmp_path / "config.yml"
    import yaml as _yaml
    cfg_path.write_text(_yaml.safe_dump(cfg), encoding="utf-8")
    monkeypatch.setattr(wm, "CONFIG_PATH", cfg_path)
    monkeypatch.setattr(wm, "_sync_active_project_from_file", lambda *a, **k: None)
    # Don't depend on the file-vs-DB config sync; just capture any write helper used.
    if hasattr(wm, "load_config"):
        monkeypatch.setattr(wm, "load_config", lambda *a, **k: dict(cfg))
    if hasattr(wm, "write_config"):
        monkeypatch.setattr(wm, "write_config", lambda c, p=None: saved.update({"cfg": c}))

    resp = api_client.delete("/api/templates/foo.docx")
    assert resp.status_code == 200, resp.text

    # Read back whatever the endpoint persisted (file and/or via write_config).
    final = {}
    try:
        final = _yaml.safe_load(cfg_path.read_text(encoding="utf-8")) or {}
    except FileNotFoundError:
        final = {}
    written = saved.get("cfg") or final
    tref = (written.get("report") or {}).get("template")

    # The reference must NOT still point at the deleted file.
    assert tref in (None, "", ) or Path(str(tref)).name != "foo.docx", \
        f"report.template still points at the deleted template: {tref!r}"


# --------------------------------------------------------------------------- #
# XTF-22 — /api/template/apply persists synthesized auto-modeling views.
#
# The deterministic resolver (resolve_sources, tested in
# tests/test_template_inference.py) synthesizes a persisted view when a placeholder
# spans a repeat table + main. On INFER those pending views travel back to the
# panel; on APPLY they must be written into config ``views:`` (appended, de-duped
# by name) so a re-run of build-report resolves the join.
#
# ASSUMED CONTRACT (chosen as the most natural design given the existing apply
# wiring, which carries the full proposal dicts from infer → apply):
#   * Each proposal that needed a synthesized view carries that view dict under a
#     ``view`` key on the proposal (the same proposal the panel approves), e.g.
#       {"token_index": 0, "kind": "chart", "name": "...", "status": "ok",
#        "spec": {..., "source": "auto_health_facilities__commune"},
#        "view": {"name": "auto_health_facilities__commune",
#                 "source": "health_facilities", "join_parent": ["Commune"]}}
#   * /api/template/apply collects every approved proposal's ``view`` and APPENDS
#     it into ``cfg["views"]`` (creating the list if absent), de-duped by ``name``
#     (a view whose name already exists in cfg["views"] is NOT appended twice).
#
# If the implementer instead threads pending views via a top-level ``views`` field
# on the apply payload, that is an equivalent design; this test pins the
# proposal-carried shape because it matches how the existing apply endpoint already
# receives proposals and how infer returns them, and is the minimal wiring change.
# --------------------------------------------------------------------------- #
def test_apply_persists_synthesized_view(monkeypatch, client, tmp_path):
    """AC: /api/template/apply persists a synthesized view into config ``views:``
    (appended, de-duped).

    A chart proposal whose spec sources a synthesized join-view carries the view
    dict on the proposal. After apply, cfg["views"] contains that view exactly
    once with source = the repeat table and join_parent = [the main col]."""
    cfg = {"charts": [], "views": []}
    saved = {}
    monkeypatch.setattr(wm, "_require", lambda *a, **k: None)
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "write_config", lambda c, p: saved.update({"cfg": c}))

    resolved = str(tmp_path / "report.resolved.docx")
    view = {"name": "auto_health_facilities__commune",
            "source": "health_facilities", "join_parent": ["Commune"]}
    approved = [
        {"token_index": 0, "kind": "chart", "name": "beds_by_commune",
         "spec": {"name": "beds_by_commune", "type": "bar",
                  "questions": ["Beds"], "group_by": "Commune",
                  "source": "auto_health_facilities__commune"},
         "status": "ok",
         "view": dict(view)},
    ]

    import src.reports.template_inference as ti

    def _apply(approved_props, cfg_in, template_path):
        # apply_inference writes the chart spec into config (its real job); the
        # view-persistence is the endpoint's responsibility under test here.
        cfg_in.setdefault("charts", []).append(approved_props[0]["spec"])
        return cfg_in, resolved

    monkeypatch.setattr(ti, "apply_inference", _apply)

    resp = client.post("/api/template/apply",
                       json={"proposals": approved, "template": "report.docx"})
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["ok"] is True

    written = saved.get("cfg") or cfg
    views = written.get("views") or []
    matching = [v for v in views if v.get("name") == "auto_health_facilities__commune"]
    assert len(matching) == 1, f"synthesized view not persisted exactly once: {views}"
    assert matching[0].get("source") == "health_facilities", matching[0]
    assert matching[0].get("join_parent") == ["Commune"], matching[0]


def test_apply_dedupes_existing_synthesized_view(monkeypatch, client, tmp_path):
    """AC: views are de-duped by name — applying TWO approved proposals that carry
    the SAME synthesized view name persists it exactly once (no duplicate appended).

    This is the API-level de-dupe guard. It is RED until the endpoint appends each
    proposal's carried ``view`` into cfg["views"] AND de-dupes by name: an
    unimplemented endpoint leaves cfg["views"] empty (0 != 1)."""
    view = {"name": "auto_health_facilities__commune",
            "source": "health_facilities", "join_parent": ["Commune"]}
    cfg = {"charts": [], "views": []}
    saved = {}
    monkeypatch.setattr(wm, "_require", lambda *a, **k: None)
    monkeypatch.setattr(wm, "load_config", lambda *a, **k: cfg)
    monkeypatch.setattr(wm, "write_config", lambda c, p: saved.update({"cfg": c}))

    resolved = str(tmp_path / "report.resolved.docx")
    # Two approved proposals whose join resolves to the SAME synthesized view.
    approved = [
        {"token_index": 0, "kind": "chart", "name": "beds_by_commune",
         "spec": {"name": "beds_by_commune", "type": "bar",
                  "questions": ["Beds"], "group_by": "Commune",
                  "source": "auto_health_facilities__commune"},
         "status": "ok", "view": dict(view)},
        {"token_index": 1, "kind": "chart", "name": "beds_by_commune_2",
         "spec": {"name": "beds_by_commune_2", "type": "horizontal_bar",
                  "questions": ["Beds"], "group_by": "Commune",
                  "source": "auto_health_facilities__commune"},
         "status": "ok", "view": dict(view)},
    ]

    import src.reports.template_inference as ti
    monkeypatch.setattr(ti, "apply_inference",
                        lambda ap, c, t: (c, resolved))

    resp = client.post("/api/template/apply",
                       json={"proposals": approved, "template": "report.docx"})
    assert resp.status_code == 200, resp.text

    written = saved.get("cfg") or cfg
    views = written.get("views") or []
    matching = [v for v in views if v.get("name") == "auto_health_facilities__commune"]
    assert len(matching) == 1, f"synthesized view not persisted exactly once: {views}"
