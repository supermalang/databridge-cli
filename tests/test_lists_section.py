"""MNT-23 -- tests for the new first-class `list` kind's report-generation
side: ReportBuilder._generate_lists() (cfg['lists'] -> {{ list_<name> }}
context entries, rendered via the existing build_bullet_list_text) and
template_generator.py emitting {{ list_<name> }} placeholders for cfg['lists']
entries.

Mirrors the precedent set by tests/test_tables_section.py for cfg['tables'].
"""
from pathlib import Path

import pandas as pd
from docx import Document

from src.reports.builder import ReportBuilder
from src.reports.charts import build_bullet_list_text
from src.reports.template_generator import generate_template


def _doc_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


# ── ReportBuilder._generate_lists() ─────────────────────────────────────────

def test_generate_lists_produces_list_keys():
    """Each cfg['lists'] entry yields a list_<name> context key, whose value is
    the raw bullet-prefixed text produced by build_bullet_list_text -- not an
    image."""
    cfg = {
        "questions": [],
        "lists": [{"name": "stories", "title": "Stories", "question": "Story"}],
    }
    df = pd.DataFrame({"Story": ["Alpha", "Beta", "Gamma"]})

    rb = ReportBuilder(cfg)

    class _FakeTpl:
        pass

    images = rb._generate_lists(_FakeTpl(), df, {})
    assert "list_stories" in images, f"expected 'list_stories' key, got: {list(images.keys())}"
    value = images["list_stories"]
    assert isinstance(value, str), f"list_stories must be a plain string, got {type(value)}"
    for name in ("Alpha", "Beta", "Gamma"):
        assert name in value, f"expected '{name}' in rendered list: {value!r}"


def test_generate_lists_empty_when_no_lists():
    cfg = {"questions": []}
    rb = ReportBuilder(cfg)
    assert rb._generate_lists(object(), pd.DataFrame(), {}) == {}


def test_generate_lists_renders_identically_to_build_bullet_list_text():
    """AC: _generate_lists() renders via build_bullet_list_text identically to
    how the old bullet_list-as-chart-type path renders today -- i.e. the
    context value for a cfg['lists'] entry must equal calling
    build_bullet_list_text directly on the same column."""
    cfg = {
        "questions": [],
        "lists": [{"name": "stories", "title": "Stories", "question": "Story"}],
    }
    df = pd.DataFrame({"Story": ["Alpha", "", "Gamma", None]})

    rb = ReportBuilder(cfg)
    images = rb._generate_lists(object(), df, {})

    expected = build_bullet_list_text(df, ["Story"], {})
    assert images["list_stories"] == expected, (
        f"expected identical output to build_bullet_list_text: {expected!r}, "
        f"got: {images.get('list_stories')!r}"
    )


def test_generate_lists_applies_optional_filter():
    """Config/schema impact: cfg['lists'] entries support an optional `filter`
    field -- only rows matching the filter expression are listed."""
    cfg = {
        "questions": [],
        "lists": [{
            "name": "stories", "title": "Stories",
            "question": "Story", "filter": "Region == 'N'",
        }],
    }
    df = pd.DataFrame({
        "Story": ["Alpha", "Beta", "Gamma"],
        "Region": ["N", "N", "S"],
    })

    rb = ReportBuilder(cfg)
    images = rb._generate_lists(object(), df, {})
    value = images["list_stories"]
    assert "Alpha" in value and "Beta" in value, value
    assert "Gamma" not in value, (
        f"row with Region='S' must be excluded by the filter, got: {value!r}"
    )


# ── template_generator.py: {{ list_<name> }} placeholder emission ──────────

def test_generate_template_emits_list_placeholder_for_lists_section(tmp_path):
    """AC: template_generator.py emits {{ list_<name> }} placeholders for
    cfg['lists'] entries when auto-building a Word template."""
    cfg = {
        "questions": [],
        "lists": [{"name": "stories", "title": "Success stories", "question": "Story"}],
    }
    out = tmp_path / "tpl.docx"

    generate_template(cfg, out)

    text = _doc_text(out)
    assert "{{ list_stories }}" in text, (
        f"generate_template must emit '{{{{ list_stories }}}}' for a cfg['lists'] "
        f"entry named 'stories'. Document text:\n{text}"
    )
    assert "{{ chart_stories }}" not in text
    assert "{{ table_stories }}" not in text


# =========================================================================== #
# MNT-32 — a resolved repeat-group `list` lands in cfg["lists"] and renders as
# non-empty text (no PNG / InlineImage) from its repeat-group source table.
# =========================================================================== #
# Card AC: "After apply_inference, the resolved list spec lands in cfg['lists']
# (never cfg['charts'] or cfg['tables']) and a build renders {{ list_<name> }} as
# a text run (a w:t, no InlineImage/embedded PNG) containing the repeat-group row
# values." These tests are the report-generation half; the inference/auto-resolve
# half is pinned in tests/test_template_inference.py::test_annotate_list_*.
from docxtpl import InlineImage as _InlineImage

from src.reports import template_inference as _ti


def _docx_one_placeholder(tmp_path, text, name="upload.docx"):
    doc = Document()
    doc.add_paragraph(text)
    p = tmp_path / name
    doc.save(str(p))
    return str(p)


def test_apply_inference_resolved_repeat_list_lands_in_lists_section(tmp_path):
    """MNT-32 AC: after apply_inference, a resolved (source-stamped) repeat-group
    list spec is written to cfg['lists'] -- never cfg['charts'] or cfg['tables']
    -- carrying the key _generate_lists reads so its rows actually render."""
    template = _docx_one_placeholder(tmp_path, "[Villages]")
    tokens = _ti.extract_placeholders(template)
    assert len(tokens) == 1

    approved = [{
        "token_index": 0,
        "kind": "list",
        "spec": {"name": "villages", "title": "Villages",
                 "question": "Village", "source": "hh_members"},
        "name": "villages",
        "confidence": 0.95,
        "reason": "approved",
        "status": "ok",
    }]
    cfg = {"api": {}, "form": {}}

    cfg_out, _resolved = _ti.apply_inference(approved, cfg, template)

    entries = cfg_out.get("lists") or []
    entry = next((e for e in entries if e.get("name") == "villages"), None)
    assert entry is not None, f"resolved list not written to cfg['lists']: {cfg_out}"
    # The auto-stamped repeat-group source survives onto the persisted spec.
    assert entry.get("source") == "hh_members", entry
    # _generate_lists reads l.get("question") (singular) -- the persisted entry
    # must carry that key so the list renders its rows (not empty).
    assert entry.get("question") == "Village", (
        f"persisted list entry must carry the 'question' key _generate_lists "
        f"reads; got: {entry}"
    )

    # Must NOT have landed in charts or tables.
    assert not any((c.get("name") == "villages") for c in (cfg_out.get("charts") or []))
    assert not any((t.get("name") == "villages") for t in (cfg_out.get("tables") or []))


def test_generate_lists_renders_repeat_group_list_as_nonempty_text(tmp_path):
    """MNT-32 AC: _generate_lists renders a repeat-group-sourced list from its
    repeat table (not main) as a NON-EMPTY plain string containing the repeat
    row values -- and never an InlineImage / embedded PNG."""
    cfg = {
        "questions": [],
        "lists": [{
            "name": "villages", "title": "Villages",
            "question": "Village", "source": "hh_members",
        }],
    }
    # main has NO Village column; the repeat table carries it.
    main_df = pd.DataFrame({"Region": ["N", "S"]})
    repeat_tables = {
        "hh_members": pd.DataFrame(
            {"Village": ["Keur Massar", "Mbao", "Rufisque"]}
        ),
    }

    rb = ReportBuilder(cfg)
    images = rb._generate_lists(object(), main_df, repeat_tables)

    assert "list_villages" in images, (
        f"expected 'list_villages' context key, got: {list(images.keys())}"
    )
    value = images["list_villages"]
    assert not isinstance(value, _InlineImage), (
        "a list must render as text, never an InlineImage/PNG"
    )
    assert isinstance(value, str) and value.strip(), (
        f"list_villages must be a non-empty string, got {value!r}"
    )
    for village in ("Keur Massar", "Mbao", "Rufisque"):
        assert village in value, (
            f"expected repeat-group value '{village}' in rendered list: {value!r}"
        )
