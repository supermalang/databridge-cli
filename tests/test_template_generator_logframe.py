from docx import Document
from src.reports.template_generator import generate_template


def _doc_text(path):
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _cfg():
    return {
        "charts": [],
        "indicators": [
            {"name": "doses", "label": "Doses", "framework_ref": "OP1.1", "primary": True, "target": 100},
            {"name": "by_sex", "label": "By sex", "disaggregate_by": ["Sex"]},
        ],
        "framework": {
            "goal": {"id": "GOAL", "label": "G"},
            "outcomes": [{"id": "OC1", "label": "O", "parent": "GOAL"}],
            "outputs": [{"id": "OP1.1", "label": "Out", "parent": "OC1"}],
        },
    }


def test_template_renders_node_achievement_and_targets(tmp_path):
    out = tmp_path / "tpl.docx"
    generate_template(_cfg(), out)
    text = _doc_text(out)
    # node-level achievement + per-indicator target/% rendered in the logframe loop
    assert "row.node_pct_achievement" in text
    assert "ind.target" in text
    assert "ind.pct_achievement" in text


def test_template_emits_breakdown_placeholder_for_disaggregated_indicator(tmp_path):
    out = tmp_path / "tpl.docx"
    generate_template(_cfg(), out)
    text = _doc_text(out)
    assert "{{ ind_doses }}" in text                 # scalar still present
    assert "{{ ind_by_sex_table }}" in text          # breakdown table for the disaggregated one
    assert "By sex — breakdown" in text


def test_template_renders_data_quality_section(tmp_path):
    out = tmp_path / "tpl.docx"
    generate_template(_cfg(), out)
    text = _doc_text(out)
    assert "data_quality.has_data" in text
    assert "data_quality.rows" in text
    assert "row.completeness" in text and "row.outlier_rate" in text and "row.duplicate_rate" in text
    assert "data_quality.tables" in text and "t.rows" in text
