"""MNT-8 — Strip residual [[...]] delimiters from built report output.

After `docxtpl` renders the Word template, any `[[...]]` tokens that were not
resolved by the Express Fill pipeline remain in the output .docx with their raw
brackets (e.g. ``[[NOM]]`` renders literally).  The fix is a post-render pass
that replaces ``[[<inner>]]`` with ``<inner>`` in every run.

Acceptance criteria tested here:
  AC1 — After build_report, no paragraph or table cell contains ``[[`` or ``]]``.
  AC2 — Inner text is preserved: ``[[NOM]]`` → ``NOM``,
         ``[[LISTE DES PARTENAIRES]]`` → ``LISTE DES PARTENAIRES``.
  AC3 — Existing ``{{ }}`` Jinja2 placeholders that were properly filled are unaffected.
"""
import zipfile
from pathlib import Path
from unittest.mock import patch

import pandas as pd
import pytest
import yaml
from docx import Document
from docx.shared import Pt

from src.reports.builder import ReportBuilder, _strip_residual_brackets
from src.reports.template_generator import generate_template
from src.utils.config import get_palette


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------

def _docx_full_text(path: Path) -> str:
    """Return all raw XML text from word/document.xml (includes run text nodes)."""
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Shared fixture: minimal workspace with a template that contains [[tokens]]
# ---------------------------------------------------------------------------

@pytest.fixture
def workspace_with_tokens(tmp_path, monkeypatch):
    """Workspace whose Word template contains a ``[[CUSTOM_TOKEN]]`` run.

    We first generate the standard template (so all standard placeholders are
    present and docxtpl renders cleanly), then open it with python-docx and
    append a paragraph whose single run is the literal text ``[[CUSTOM_TOKEN]]``,
    then save it back.  This simulates what happens when a user has an Express
    Fill token that does not match any filled variable.
    """
    ws = tmp_path / "ws"
    (ws / "data" / "processed").mkdir(parents=True)
    (ws / "templates").mkdir()
    (ws / "reports").mkdir()

    # Minimal data file — filename follows the {alias}_data_{ts}.csv convention.
    csv_path = ws / "data" / "processed" / "mnt8_data_20260101_120000.csv"
    pd.DataFrame({"Region": ["A", "B", "A"], "Age": [10, 20, 30]}).to_csv(
        csv_path, index=False
    )

    template_path = ws / "templates" / "t.docx"

    cfg = {
        "api": {
            "url": "https://kf.kobotoolbox.org/api/v2",
            "token": "dummy",
            "platform": "kobo",
        },
        "form": {"alias": "mnt8", "uid": "x"},
        "questions": [
            {
                "kobo_key": "Region",
                "label": "Region",
                "type": "select_one",
                "category": "categorical",
                "group": "",
                "export_label": "Region",
            },
            {
                "kobo_key": "Age",
                "label": "Age",
                "type": "integer",
                "category": "quantitative",
                "group": "",
                "export_label": "Age",
            },
        ],
        "filters": [],
        "charts": [
            {
                "name": "age_dist",
                "title": "Age Distribution",
                "type": "histogram",
                "questions": ["Age"],
                "options": {},
            }
        ],
        "report": {
            "template": str(template_path),
            "output_dir": str(ws / "reports"),
            "title": "MNT-8 Smoke",
            "period": "Q2 2026",
        },
        "export": {
            "format": "csv",
            "output_dir": str(ws / "data" / "processed"),
        },
    }

    # Generate a standard template that renders cleanly.
    generate_template(cfg, template_path)

    # Inject a [[CUSTOM_TOKEN]] paragraph into the template.
    doc = Document(str(template_path))
    p = doc.add_paragraph()
    run = p.add_run("[[CUSTOM_TOKEN]]")
    run.font.size = Pt(11)
    doc.save(str(template_path))

    (ws / "config.yml").write_text(yaml.dump(cfg, allow_unicode=True))
    monkeypatch.chdir(ws)
    return {"ws": ws, "cfg": cfg, "out_dir": ws / "reports"}


@pytest.fixture
def workspace_with_multiword_token(tmp_path, monkeypatch):
    """Same as workspace_with_tokens but injects ``[[LISTE DES PARTENAIRES]]``."""
    ws = tmp_path / "ws2"
    (ws / "data" / "processed").mkdir(parents=True)
    (ws / "templates").mkdir()
    (ws / "reports").mkdir()

    csv_path = ws / "data" / "processed" / "mnt8b_data_20260101_120000.csv"
    pd.DataFrame({"Region": ["X", "Y"], "Age": [5, 15]}).to_csv(
        csv_path, index=False
    )

    template_path = ws / "templates" / "t.docx"

    cfg = {
        "api": {
            "url": "https://kf.kobotoolbox.org/api/v2",
            "token": "dummy",
            "platform": "kobo",
        },
        "form": {"alias": "mnt8b", "uid": "x"},
        "questions": [
            {
                "kobo_key": "Region",
                "label": "Region",
                "type": "select_one",
                "category": "categorical",
                "group": "",
                "export_label": "Region",
            },
        ],
        "filters": [],
        "charts": [
            {
                "name": "region_bar",
                "title": "Region",
                "type": "bar",
                "questions": ["Region"],
                "options": {},
            }
        ],
        "report": {
            "template": str(template_path),
            "output_dir": str(ws / "reports"),
            "title": "MNT-8b Smoke",
            "period": "Q2 2026",
        },
        "export": {
            "format": "csv",
            "output_dir": str(ws / "data" / "processed"),
        },
    }

    generate_template(cfg, template_path)

    doc = Document(str(template_path))
    # Two separate tokens to cover both AC1 and AC2 scenarios.
    p = doc.add_paragraph()
    p.add_run("[[NOM]]")
    p2 = doc.add_paragraph()
    p2.add_run("[[LISTE DES PARTENAIRES]]")
    doc.save(str(template_path))

    (ws / "config.yml").write_text(yaml.dump(cfg, allow_unicode=True))
    monkeypatch.chdir(ws)
    return {"ws": ws, "cfg": cfg, "out_dir": ws / "reports"}


# ---------------------------------------------------------------------------
# AC1 — output docx must NOT contain the literal substrings "[[" or "]]"
# ---------------------------------------------------------------------------

def test_no_double_bracket_open_in_output(workspace_with_tokens):
    """AC1a: After build, output docx must not contain '[[' anywhere."""
    ReportBuilder(workspace_with_tokens["cfg"]).build()
    docs = list(workspace_with_tokens["out_dir"].glob("*.docx"))
    assert docs, "build() produced no .docx output"
    xml = _docx_full_text(docs[0])
    assert "[[" not in xml, (
        "Output docx still contains raw '[['  — the post-render stripping pass is missing.\n"
        f"Sample of XML containing '[[': "
        f"{xml[max(0, xml.index('[[') - 60):xml.index('[[') + 80]!r}"
    )


def test_no_double_bracket_close_in_output(workspace_with_tokens):
    """AC1b: After build, output docx must not contain ']]' anywhere."""
    ReportBuilder(workspace_with_tokens["cfg"]).build()
    docs = list(workspace_with_tokens["out_dir"].glob("*.docx"))
    assert docs, "build() produced no .docx output"
    xml = _docx_full_text(docs[0])
    assert "]]" not in xml, (
        "Output docx still contains raw ']]'  — the post-render stripping pass is missing.\n"
        f"Sample of XML containing ']]': "
        f"{xml[max(0, xml.index(']]') - 60):xml.index(']]') + 80]!r}"
    )


# ---------------------------------------------------------------------------
# AC2 — inner text is preserved after stripping the delimiters
# ---------------------------------------------------------------------------

def test_custom_token_inner_text_preserved(workspace_with_tokens):
    """AC2a: [[CUSTOM_TOKEN]] → CUSTOM_TOKEN (inner text survives stripping)."""
    ReportBuilder(workspace_with_tokens["cfg"]).build()
    docs = list(workspace_with_tokens["out_dir"].glob("*.docx"))
    assert docs, "build() produced no .docx output"
    xml = _docx_full_text(docs[0])
    assert "CUSTOM_TOKEN" in xml, (
        "Inner text 'CUSTOM_TOKEN' was not found in the output docx. "
        "The stripping pass must preserve the content between [[ ]] delimiters."
    )


def test_multiword_token_inner_text_preserved(workspace_with_multiword_token):
    """AC2b: [[LISTE DES PARTENAIRES]] → LISTE DES PARTENAIRES (spaces preserved)."""
    ReportBuilder(workspace_with_multiword_token["cfg"]).build()
    docs = list(workspace_with_multiword_token["out_dir"].glob("*.docx"))
    assert docs, "build() produced no .docx output"
    xml = _docx_full_text(docs[0])
    assert "LISTE DES PARTENAIRES" in xml, (
        "Inner text 'LISTE DES PARTENAIRES' was not found in the output docx. "
        "Multi-word token content must be preserved after stripping delimiters."
    )


def test_nom_token_inner_text_preserved(workspace_with_multiword_token):
    """AC2c: [[NOM]] → NOM (short token inner text survives)."""
    ReportBuilder(workspace_with_multiword_token["cfg"]).build()
    docs = list(workspace_with_multiword_token["out_dir"].glob("*.docx"))
    assert docs, "build() produced no .docx output"
    xml = _docx_full_text(docs[0])
    assert "NOM" in xml, (
        "Inner text 'NOM' was not found in the output docx. "
        "Short token content must be preserved after stripping delimiters."
    )


def test_strip_token_with_delimiter_char_split_across_runs(tmp_path):
    """MNT-14 AC: a `[[`/`]]` delimiter *character itself* broken mid-token
    across two or more separate runs must still be fully stripped, with the
    inner text preserved.

    MNT-8's ``test_strip_token_split_across_runs`` covers a delimiter split
    *between* runs where each run still holds a complete, intact delimiter
    (e.g. "[[" | "NOM]]"). This test covers the stricter case flagged by
    MNT-14: the delimiter character itself is broken mid-character across
    runs, e.g. "[" | "[NOM]" | "]" as three separate python-docx Run objects
    in the same paragraph — so no single run contains a complete "[[" or "]]"
    pair. Per MNT-14's own probe-test finding, this case leaves "[[NOM]]"
    unstripped under the current per-run-only replace.
    """
    doc_path = tmp_path / "cross_run_char_split.docx"

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("[")
    p.add_run("[NOM]")
    p.add_run("]")
    assert len(p.runs) == 3, "fixture setup error: expected three distinct runs"
    doc.save(str(doc_path))

    _strip_residual_brackets(doc_path)

    reloaded = Document(str(doc_path))
    full_text = "".join(run.text for para in reloaded.paragraphs for run in para.runs)

    assert "[[" not in full_text, (
        "Output still contains raw '[[' after stripping a token whose "
        "delimiter character itself was split across separate runs. "
        f"Reconstructed paragraph text: {full_text!r}"
    )
    assert "]]" not in full_text, (
        "Output still contains raw ']]' after stripping a token whose "
        "delimiter character itself was split across separate runs. "
        f"Reconstructed paragraph text: {full_text!r}"
    )
    assert "NOM" in full_text, (
        "Inner text 'NOM' was lost when stripping a token whose delimiter "
        f"character was split across separate runs. Reconstructed paragraph "
        f"text: {full_text!r}"
    )


def test_strip_token_split_across_runs(tmp_path):
    """AC1/AC2: [[ and NOM]] split across two separate runs of the same
    paragraph must still be stripped by the post-render pass.

    Word frequently splits a single visual token across multiple <w:r> runs
    (e.g. because of spell-check boundaries, mid-edit formatting changes, or
    copy/paste). This builds a paragraph with two genuinely separate
    python-docx Run objects — one containing only "[[" and the next
    containing only "NOM]]" — so the delimiters and inner text do not share
    a single run's text attribute. It then invokes the same
    ``_strip_residual_brackets`` post-render pass build() uses and asserts
    the saved-and-reloaded document no longer contains the literal
    substrings '[[' or ']]' anywhere in the paragraph.
    """
    doc_path = tmp_path / "cross_run.docx"

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("[[")
    p.add_run("NOM]]")
    assert len(p.runs) == 2, "fixture setup error: expected two distinct runs"
    doc.save(str(doc_path))

    _strip_residual_brackets(doc_path)

    reloaded = Document(str(doc_path))
    full_text = "".join(run.text for para in reloaded.paragraphs for run in para.runs)

    assert "[[" not in full_text, (
        "Output still contains raw '[[' after stripping a token split across "
        f"separate runs. Reconstructed paragraph text: {full_text!r}"
    )
    assert "]]" not in full_text, (
        "Output still contains raw ']]' after stripping a token split across "
        f"separate runs. Reconstructed paragraph text: {full_text!r}"
    )
    assert "NOM" in full_text, (
        "Inner text 'NOM' was lost when stripping a token split across "
        f"separate runs. Reconstructed paragraph text: {full_text!r}"
    )


# ---------------------------------------------------------------------------
# AC3 — standard Jinja2-filled values are unaffected by the stripping pass
# ---------------------------------------------------------------------------

def test_jinja2_filled_values_unaffected(workspace_with_tokens):
    """AC3: Standard Jinja2 context values (period, title, etc.) survive intact.

    The report title 'MNT-8 Smoke', the period 'Q2 2026' and the submission
    count are rendered by docxtpl from {{ }} placeholders; they must be present
    and correct in the output after the stripping pass runs.
    """
    ReportBuilder(workspace_with_tokens["cfg"]).build()
    docs = list(workspace_with_tokens["out_dir"].glob("*.docx"))
    assert docs, "build() produced no .docx output"
    xml = _docx_full_text(docs[0])
    assert "MNT-8 Smoke" in xml, (
        "Report title 'MNT-8 Smoke' missing — the stripping pass must not corrupt "
        "text that was properly rendered by docxtpl."
    )
    assert "Q2 2026" in xml, (
        "Period 'Q2 2026' missing — the stripping pass must not corrupt "
        "text that was properly rendered by docxtpl."
    )


# ---------------------------------------------------------------------------
# MNT-11 — builder-level palette wiring: ReportBuilder._generate_charts must
# resolve brand.palette via get_palette(cfg) and pass it into generate_chart().
# ---------------------------------------------------------------------------

@pytest.fixture
def workspace_with_brand_palette(tmp_path, monkeypatch):
    """Minimal workspace identical in shape to workspace_with_tokens, but with
    `brand.palette: "teal"` set in config.yml and a categorical bar chart."""
    ws = tmp_path / "ws_palette"
    (ws / "data" / "processed").mkdir(parents=True)
    (ws / "templates").mkdir()
    (ws / "reports").mkdir()

    csv_path = ws / "data" / "processed" / "mnt11_data_20260101_120000.csv"
    pd.DataFrame({"Region": ["North", "South", "North", "East", "South"]}).to_csv(
        csv_path, index=False
    )

    template_path = ws / "templates" / "t.docx"

    cfg = {
        "api": {
            "url": "https://kf.kobotoolbox.org/api/v2",
            "token": "dummy",
            "platform": "kobo",
        },
        "form": {"alias": "mnt11", "uid": "x"},
        "questions": [
            {
                "kobo_key": "Region",
                "label": "Region",
                "type": "select_one",
                "category": "categorical",
                "group": "",
                "export_label": "Region",
            },
        ],
        "filters": [],
        "brand": {"palette": "teal"},
        "charts": [
            {
                "name": "region_bar",
                "title": "Region",
                "type": "bar",
                "questions": ["Region"],
                "options": {},
            }
        ],
        "report": {
            "template": str(template_path),
            "output_dir": str(ws / "reports"),
            "title": "MNT-11 Smoke",
            "period": "Q2 2026",
        },
        "export": {
            "format": "csv",
            "output_dir": str(ws / "data" / "processed"),
        },
    }

    generate_template(cfg, template_path)

    (ws / "config.yml").write_text(yaml.dump(cfg, allow_unicode=True))
    monkeypatch.chdir(ws)
    return {"ws": ws, "cfg": cfg, "out_dir": ws / "reports"}


def test_generate_charts_passes_resolved_brand_palette_to_generate_chart(
    workspace_with_brand_palette,
):
    """MNT-11 AC: builder.py passes the resolved palette into the chart
    dispatch so all charts in a report share the same colour sequence.

    Patches src.reports.builder.generate_chart (the exact call site inside
    ReportBuilder._generate_charts) and asserts it is invoked with a
    `palette` kwarg equal to get_palette(cfg)'s resolved "teal" sequence —
    not None, not the default "slate" sequence, and not merely "truthy".
    """
    cfg = workspace_with_brand_palette["cfg"]
    expected_palette = get_palette(cfg)
    assert expected_palette != get_palette({}), (
        "test fixture sanity check: 'teal' must differ from the default "
        "'slate' palette, otherwise this test could pass vacuously"
    )

    with patch(
        "src.reports.builder.generate_chart", wraps=None, return_value=None
    ) as mock_generate_chart:
        ReportBuilder(cfg).build()

    assert mock_generate_chart.called, (
        "generate_chart was never called — _generate_charts must invoke it "
        "for each configured chart"
    )
    for call in mock_generate_chart.call_args_list:
        _, kwargs = call
        assert "palette" in kwargs, (
            "generate_chart was called without a 'palette' kwarg — "
            "_generate_charts must pass palette=get_palette(self.cfg)"
        )
        assert kwargs["palette"] == expected_palette, (
            f"generate_chart was called with palette={kwargs['palette']!r}, "
            f"expected the resolved 'teal' palette {expected_palette!r} "
            "from get_palette(cfg). The builder must wire brand.palette "
            "through to every chart it generates."
        )
