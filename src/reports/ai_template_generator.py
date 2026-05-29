"""
ai_template_generator.py — LLM-powered Word template generation.

Asks the LLM to design a report layout (JSON spec), then renders it into a
.docx using python-docx helpers from template_generator.py.

Called by the ai-generate-template CLI command.
"""
import json
import logging
import re
from pathlib import Path
from typing import Dict, List

from docx import Document

from src.reports.template_generator import (
    _chart_ph, _descriptor, _divider, _editable,
    _heading, _margins, _meta, _note, _ref_table,
)

log = logging.getLogger(__name__)


# ── public entry point ────────────────────────────────────────────────────────

def ai_generate_template(
    cfg: Dict,
    out_path: Path,
    description: str,
    pages: int = 10,
    language: str = "English",
    summary_prompt: str = None,
) -> Path:
    """Generate a Word template from an LLM layout spec and save to out_path."""
    ai_cfg = cfg.get("ai")
    if not ai_cfg:
        raise ValueError("No ai: section in config.yml. Configure AI first.")

    log.info("Requesting layout from LLM…")
    spec = _get_layout_spec(
        ai_cfg, cfg, description, pages, language,
        summary_prompt=summary_prompt,
    )
    log.info(f"Layout received: {len(spec.get('sections', []))} sections")

    doc = Document()
    _margins(doc)

    # Fixed header — always present
    _heading(doc, "{{ report_title }}", 0)
    _meta(doc, "Period", "{{ period }}")
    _meta(doc, "Submissions", "{{ n_submissions }}")
    _meta(doc, "Generated", "{{ generated_at }}")
    _divider(doc)

    # Build indicator lookup
    ind_map = {
        ind["name"]: ind.get("label", ind["name"])
        for ind in cfg.get("indicators", [])
    }

    for section in spec.get("sections", []):
        _heading(doc, section.get("heading", ""), section.get("level", 1))
        for item in section.get("content", []):
            _render_item(doc, item, ind_map, cfg)
        _divider(doc)

    # Always append placeholder reference
    _heading(doc, "Placeholder Reference", 1)
    _note(doc, "Delete this section before sharing.")
    _ref_table(doc, cfg)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    log.info(f"AI template saved → {out_path}")
    return out_path


# ── item renderer ─────────────────────────────────────────────────────────────

def _render_item(doc, item: Dict, ind_map: Dict, cfg: Dict):
    t = item.get("type", "")

    if t == "editable":
        ph = item.get("placeholder", "")
        hint = item.get("hint", "")
        _editable(doc, f"{{{{ {ph} }}}}", hint)

    elif t == "chart":
        name = item.get("name", "")
        # Find chart title for descriptor
        title = next(
            (c.get("title", name) for c in cfg.get("charts", []) if c.get("name") == name),
            name,
        )
        chart_type = next(
            (c.get("type", "") for c in cfg.get("charts", []) if c.get("name") == name),
            "",
        )
        if chart_type:
            _descriptor(doc, f"{chart_type.replace('_', ' ').title()}")
        _chart_ph(doc, f"{{{{ chart_{name} }}}}")
        doc.add_paragraph()

    elif t == "indicator":
        name = item.get("name", "")
        label = ind_map.get(name, name.replace("_", " ").title())
        _meta(doc, label, f"{{{{ ind_{name} }}}}")

    elif t == "summary":
        name = item.get("name", "")
        # Find label from config summaries
        summary_cfg = cfg.get("summaries", [])
        label = next(
            (s.get("label", name) for s in summary_cfg if s.get("name") == name),
            name.replace("_", " ").title(),
        )
        _descriptor(doc, label)
        _editable(doc, f"{{{{ summary_{name} }}}}", hint="")

    elif t == "text":
        text = item.get("text", "")
        if text:
            doc.add_paragraph(text)

    elif t == "divider":
        _divider(doc)

    elif t == "stats_table":
        from docx.shared import Pt, RGBColor
        quant = [q for q in cfg.get("questions", []) if q.get("category") == "quantitative"]
        if quant:
            for text, sz, color, italic in [
                ("{%p for row in stats_table %}", 9, RGBColor(0x1D, 0x9E, 0x75), True),
                ("{{ row.label }}: n={{ row.n }}, mean={{ row.mean }}, median={{ row.median }}", 9, RGBColor(0x1D, 0x9E, 0x75), True),
                ("{%p endfor %}", 9, RGBColor(0x1D, 0x9E, 0x75), True),
            ]:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.size = Pt(sz)
                r.font.color.rgb = color
                r.font.italic = italic


# ── LLM interaction ───────────────────────────────────────────────────────────

def _get_layout_spec(
    ai_cfg: Dict, cfg: Dict, description: str, pages: int, language: str,
    summary_prompt: str = None,
) -> Dict:
    from src.utils import lf_client

    provider   = ai_cfg.get("provider", "openai").lower()
    api_key    = ai_cfg.get("api_key", "")
    model      = ai_cfg.get("model", "gpt-4o")
    max_tokens = max(int(ai_cfg.get("max_tokens", 1500)), 2000)  # layout needs more tokens

    if not api_key or str(api_key).startswith("env:"):
        raise ValueError("AI api_key not resolved. Set the environment variable.")

    variables = _build_variables(cfg, description, pages, language, summary_prompt=summary_prompt)
    messages, config = lf_client.get_prompt("template_generator", variables)
    raw = lf_client.chat(
        messages,
        model=model,
        provider=provider,
        api_key=api_key,
        max_tokens=max_tokens,
        trace_name="template_generator",
        base_url=ai_cfg.get("base_url"),
        json_mode=(provider != "anthropic"),
        output_schema=config.get("output_schema"),
    )

    return _parse_spec(raw)


def _build_variables(
    cfg: Dict, description: str, pages: int, language: str,
    summary_prompt: str = None,
) -> Dict:
    summary_prompt_line = (
        f"Executive summary guidance (use as hint for the summary_text editable): {summary_prompt}"
        if summary_prompt else ""
    )

    # Charts block — pass chart questions so LLM can group charts into logical sections
    charts = cfg.get("charts", [])
    charts_block = ""
    if charts:
        items = []
        for c in charts:
            questions_str = ", ".join(c.get("questions", []))
            detail = f"{c.get('title', c['name'])}"
            if questions_str:
                detail += f" — columns: {questions_str}"
            items.append(f"  - {c['name']} ({c.get('type', '')}): {detail}")
        charts_block = "Available charts:\n" + "\n".join(items) + "\n\n"

    # Indicators block
    indicators = cfg.get("indicators", [])
    indicators_block = ""
    if indicators:
        items = [
            f"  - {ind['name']}: {ind.get('label', ind['name'])} ({ind.get('stat', '')})"
            for ind in indicators
        ]
        indicators_block = "Available indicators:\n" + "\n".join(items) + "\n\n"

    # Summaries block
    summaries = cfg.get("summaries", [])
    summaries_block = ""
    if summaries:
        items = []
        for s in summaries:
            stat = s.get("stat", "")
            questions_str = ", ".join(s.get("questions", []))
            detail = s.get("label", s["name"])
            if questions_str:
                detail += f" — {questions_str}"
            items.append(f"  - {s['name']} ({stat}): {detail}")
        summaries_block = "Available summaries (computed text paragraphs):\n" + "\n".join(items) + "\n\n"

    # Views block — inform LLM of pre-built data tables so it can note them in chart hints
    views = cfg.get("views", [])
    views_block = ""
    if views:
        items = []
        for v in views:
            name = v.get("name", "")
            gb   = v.get("group_by", "")
            q_v  = v.get("question", "")
            desc = f"source: {v.get('source', 'main')}"
            if v.get("join_parent"): desc += f", joined with {', '.join(v['join_parent'])}"
            if gb and q_v:           desc += f", {v.get('agg','sum')}({q_v}) by {gb}"
            items.append(f"  - {name}: {desc}")
        views_block = "Named data views (virtual tables used as chart/summary sources):\n" + "\n".join(items) + "\n\n"

    # Questions block — actual question labels grouped by category
    questions = cfg.get("questions", [])
    questions_block = ""
    if questions:
        items = []
        for cat in ("categorical", "quantitative", "qualitative", "date", "geographical", "undefined"):
            qs = [
                q.get("export_label") or q.get("label") or q["kobo_key"]
                for q in questions if q.get("category") == cat
            ]
            if qs:
                # Cap to 10 per category to keep prompt bounded
                sample = qs[:10]
                suffix = f" (+{len(qs)-10} more)" if len(qs) > 10 else ""
                items.append(f"  {cat}: {', '.join(sample)}{suffix}")
        if items:
            questions_block = "Survey questions by category:\n" + "\n".join(items) + "\n\n"

    return {
        "description": description,
        "pages": pages,
        "language": language,
        "summary_prompt_line": summary_prompt_line,
        "charts_block": charts_block,
        "indicators_block": indicators_block,
        "summaries_block": summaries_block,
        "views_block": views_block,
        "questions_block": questions_block,
        # The template_generator system prompt documents these as docxtpl
        # placeholders the LLM may emit in generated text. They must survive
        # {{mustache}} compilation verbatim, so map each to its literal form.
        "period": "{{ period }}",
        "n_submissions": "{{ n_submissions }}",
        "generated_at": "{{ generated_at }}",
    }


# ── parser ────────────────────────────────────────────────────────────────────

def _parse_spec(raw: str) -> Dict:
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                pass
    log.warning("Could not parse LLM layout spec — using minimal fallback.")
    return {"sections": [
        {"heading": "Executive Summary", "level": 1, "content": [
            {"type": "editable", "placeholder": "summary_text",
             "hint": "Summarize the key findings, coverage rates, and any critical gaps or trends observed during the reporting period."}
        ]},
        {"heading": "Key Findings", "level": 1, "content": [
            {"type": "text", "text": "This section presents the main findings from the {{ period }} data collection ({{ n_submissions }} submissions)."}
        ]},
        {"heading": "Observations", "level": 1, "content": [
            {"type": "editable", "placeholder": "observations",
             "hint": "Document factual observations from the data — patterns, outliers, and comparisons against targets or baselines."}
        ]},
        {"heading": "Recommendations", "level": 1, "content": [
            {"type": "editable", "placeholder": "recommendations",
             "hint": "List actionable recommendations for programme teams, with priority level and responsible party where possible."}
        ]},
    ]}
