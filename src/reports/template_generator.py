import logging
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

log = logging.getLogger(__name__)

def generate_template(cfg: Dict, out_path: Path, context: str = None, summary_prompt: str = None) -> Path:
    doc = Document(); charts: List[Dict] = cfg.get("charts",[])
    _margins(doc)
    _heading(doc,"{{ report_title }}",0)
    _meta(doc,"Period","{{ period }}")
    _meta(doc,"Submissions","{{ n_submissions }}")
    _meta(doc,"Generated","{{ generated_at }}")
    if context:
        _divider(doc)
        _heading(doc,"Background & Context",1)
        _editable(doc,"",context)
    _divider(doc)
    _heading(doc,"Executive Summary",1)
    _editable(doc,"{{ summary_text }}",summary_prompt or "Write your executive summary here.")
    _divider(doc)
    if charts:
        _heading(doc,"Data & Visualizations",1)
        for c in charts:
            name=c.get("name",""); _heading(doc,c.get("title",name),2)
            _descriptor(doc,f"{c.get('type','').replace('_',' ').title()} — {', '.join(c.get('questions',[]))}")
            _chart_ph(doc,f"{{{{ chart_{name} }}}}"); doc.add_paragraph()
    else:
        _note(doc,"No charts in config.yml. Add charts and re-run generate-template.")
    _divider(doc)
    _heading(doc,"Observations",1); _editable(doc,"{{ observations }}","Write your observations here.")
    _heading(doc,"Recommendations",1); _editable(doc,"{{ recommendations }}","Write your recommendations here.")
    quant=[q for q in cfg.get("questions",[]) if q.get("category")=="quantitative"]
    if quant:
        _divider(doc); _heading(doc,"Numeric Summary",1)
        for text,sz,color,italic in [
            ("{%p for row in stats_table %}",9,RGBColor(0x1D,0x9E,0x75),True),
            ("{{ row.label }}: n={{ row.n }}, mean={{ row.mean }}, median={{ row.median }}",9,RGBColor(0x1D,0x9E,0x75),True),
            ("{%p endfor %}",9,RGBColor(0x1D,0x9E,0x75),True),
        ]:
            p=doc.add_paragraph(); r=p.add_run(text)
            r.font.size=Pt(sz); r.font.color.rgb=color; r.font.italic=italic
    indicators=cfg.get("indicators",[])
    if indicators:
        _divider(doc); _heading(doc,"Key Indicators",1)
        for ind in indicators:
            name=ind.get("name",""); label=ind.get("label",name)
            _meta(doc,label,f"{{{{ ind_{name} }}}}")
    summaries=cfg.get("summaries",[])
    if summaries:
        _divider(doc); _heading(doc,"Data Summaries",1)
        for s in summaries:
            name=s.get("name",""); label=s.get("label",name)
            _heading(doc,label,2)
            _descriptor(doc,f"{s.get('stat','distribution')} — {', '.join(s.get('questions',[]))}")
            _editable(doc,f"{{{{ summary_{name} }}}}")
    # Provenance footer — single Jinja line; ReportBuilder fills it in.
    # Must live BEFORE the deletable "Placeholder Reference" section so that
    # deleting that section does not remove the footer.
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run("{{ provenance.footer }}")
    run.italic = True
    run.font.size = Pt(8)
    _divider(doc); _heading(doc,"Placeholder Reference",1)
    _note(doc,"Delete this section before sharing."); _ref_table(doc,cfg)
    out_path.parent.mkdir(parents=True,exist_ok=True)
    doc.save(out_path)
    log.info(f"Template saved → {out_path}")
    for c in charts: log.info(f"  {{{{ chart_{c.get('name')} }}}}")
    return out_path

def _margins(doc):
    for s in doc.sections:
        s.top_margin=Inches(1); s.bottom_margin=Inches(1)
        s.left_margin=Inches(1.2); s.right_margin=Inches(1.2)

def _heading(doc,text,level):
    if level==0:
        p=doc.add_paragraph(); r=p.add_run(text)
        r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(0x0F,0x6E,0x56)
    else: doc.add_heading(text,level=level)

def _meta(doc,label,ph):
    p=doc.add_paragraph()
    r1=p.add_run(f"{label}: "); r1.bold=True; r1.font.size=Pt(10)
    r2=p.add_run(ph); r2.font.size=Pt(10); r2.font.color.rgb=RGBColor(0x37,0x8A,0xDD)
    p.paragraph_format.space_after=Pt(2)

def _divider(doc):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    b=OxmlElement("w:bottom")
    b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6"); b.set(qn("w:space"),"1"); b.set(qn("w:color"),"CCCCCC")
    pBdr.append(b); pPr.append(pBdr)

def _editable(doc,ph,hint=""):
    p=doc.add_paragraph(); r=p.add_run(ph)
    r.font.color.rgb=RGBColor(0x18,0x5F,0xA5); r.font.italic=True
    if hint:
        hp=doc.add_paragraph(hint)
        hp.runs[0].font.color.rgb=RGBColor(0xAA,0xAA,0xAA)
        hp.runs[0].font.size=Pt(9); hp.runs[0].font.italic=True

def _descriptor(doc,text):
    p=doc.add_paragraph(text); p.runs[0].font.size=Pt(9)
    p.runs[0].font.color.rgb=RGBColor(0x88,0x87,0x80); p.paragraph_format.space_after=Pt(4)

def _chart_ph(doc,ph):
    p=doc.add_paragraph(); r=p.add_run(ph)
    r.font.color.rgb=RGBColor(0x1D,0x9E,0x75); r.font.bold=True; r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)

def _note(doc,text):
    p=doc.add_paragraph(f"i  {text}")
    p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor(0x88,0x87,0x80); p.runs[0].font.italic=True

def _ref_table(doc,cfg):
    rows=[
        ("{{ report_title }}","Report title"),("{{ period }}","Reporting period"),
        ("{{ n_submissions }}","Number of submissions"),("{{ generated_at }}","Generation timestamp"),
        ("{{ summary_text }}","Executive summary — fill in Word"),
        ("{{ observations }}","Observations — fill in Word"),
        ("{{ recommendations }}","Recommendations — fill in Word"),
    ]
    for ind in cfg.get("indicators",[]):
        rows.append((f"{{{{ ind_{ind.get('name','')} }}}}",f"Indicator: {ind.get('label',ind.get('name',''))}"))
    for s in cfg.get("summaries",[]):
        rows.append((f"{{{{ summary_{s.get('name','')} }}}}",f"Summary: {s.get('label',s.get('name',''))}"))
    for c in cfg.get("charts",[]):
        rows.append((f"{{{{ chart_{c.get('name','')} }}}}",f"Chart: {c.get('title','')}"))
    table=doc.add_table(rows=1,cols=2); table.style="Table Grid"
    hdr=table.rows[0].cells
    hdr[0].text="Placeholder"; hdr[1].text="Description"
    for cell in hdr:
        for r in cell.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"0F6E56")
        tcPr.append(shd)
    for ph,desc in rows:
        row=table.add_row().cells
        row[0].text=ph; row[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0x1D,0x9E,0x75)
        row[1].text=desc
