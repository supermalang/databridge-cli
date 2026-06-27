"""Express Template Fill — placeholder extraction (XTF-1).

Pure, offline parsing of an uploaded ``.docx`` into structured :class:`Token`s.
No AI, no network. This is the foundation for the rest of the express path
(inference / validation / resolution live in later cards).

``extract_placeholders(docx_path)`` walks the document body paragraphs, table
cells, headers, and footers. For each paragraph it reconstructs the full text by
concatenating its runs (so a placeholder split across several runs — which
hand-typed placeholders almost always are — is still matched), then matches the
three interchangeable delimiters in precedence order ``[[ … ]]`` → ``[ … ]`` →
``{{ … }}`` (so a ``[[x]]`` token is never double-matched as ``[x]``).

A ``{{ }}`` token whose inner text is a *known literal* placeholder (the set the
existing report builder already understands) is marked ``kind == "literal"`` and
left untouched; everything else is returned as a non-literal natural-language
token for downstream inference.
"""
import json
import logging
import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Tuple

from docx import Document

from src.reports import ask_engine
from src.utils import lf_client

log = logging.getLogger(__name__)

# --------------------------------------------------------------------------- #
# Known {{ }} literal placeholders (today's report-builder contract).
# --------------------------------------------------------------------------- #
# Exact names the builder fills directly. Prefix families (chart_, ind_,
# summary_, table_, data_quality, logframe) are matched via the regexes below.
_LITERAL_EXACT = frozenset({
    "report_title",
    "period",
    "n_submissions",
    "generated_at",
    "summary_text",
    "observations",
    "recommendations",
    "data_quality",
    "logframe",
    "provenance.footer",
})

# ``ind_*`` (incl. ``_table`` / ``_breakdown``), ``chart_*``, ``summary_*``,
# ``table_*``, ``data_quality*``, ``logframe*``.
_LITERAL_PREFIX_RE = re.compile(
    r"^(?:chart_|ind_|summary_|table_|data_quality|logframe)\w*$"
)


def _is_known_literal(inner: str) -> bool:
    """True if a ``{{ }}`` placeholder's inner text is a known literal."""
    return inner in _LITERAL_EXACT or bool(_LITERAL_PREFIX_RE.match(inner))


# --------------------------------------------------------------------------- #
# Delimiter patterns, in precedence order. ``[[ … ]]`` first so it is never
# double-matched as ``[ … ]``. Inner text is non-greedy and forbids the
# delimiter chars so matches don't run past their own close.
# --------------------------------------------------------------------------- #
_PATTERNS = (
    ("[[", re.compile(r"\[\[(.*?)\]\]")),
    ("[", re.compile(r"\[([^\[\]]*?)\]")),
    ("{{", re.compile(r"\{\{(.*?)\}\}")),
)


# --------------------------------------------------------------------------- #
# Token + location types (attribute access is part of the contract).
# --------------------------------------------------------------------------- #
@dataclass
class Location:
    """Where a token lives, with enough detail to rewrite it later.

    ``runs`` is the sequence of integer run indices (into the owning paragraph's
    ``runs``) that the token's ``raw`` text spans. ``paragraph_text`` is the full
    paragraph text reconstructed by concatenating those runs.
    """

    runs: List[int] = field(default_factory=list)
    paragraph_text: str = ""


@dataclass
class Token:
    raw: str            # full delimited string, e.g. "[[Total]]" / "{{ x }}"
    inner: str          # trimmed inner text, e.g. "Total" / "x"
    delimiter: str      # opening delimiter: "[[", "[", or "{{"
    kind: str           # "literal" for known {{ }} literals, else "nl"
    location: Location


def _run_span(run_offsets: List[int], start: int, end: int) -> List[int]:
    """Run indices whose character ranges overlap ``[start, end)``.

    ``run_offsets`` are the cumulative character offsets where each run begins,
    with a final sentinel offset equal to the paragraph length.
    """
    spanned: List[int] = []
    for i in range(len(run_offsets) - 1):
        r_start, r_end = run_offsets[i], run_offsets[i + 1]
        # Overlap (treat zero-length runs at the boundary as non-spanning).
        if r_start < end and r_end > start:
            spanned.append(i)
    return spanned


def _tokens_in_paragraph(paragraph) -> List[Token]:
    """Extract all tokens from a single python-docx paragraph."""
    runs = paragraph.runs
    text = "".join(r.text for r in runs)
    if not text:
        return []

    # Cumulative run start offsets (+ sentinel = total length).
    run_offsets: List[int] = []
    acc = 0
    for r in runs:
        run_offsets.append(acc)
        acc += len(r.text)
    run_offsets.append(acc)

    # Find matches per delimiter, honouring precedence by claiming spans: once a
    # character range is claimed by a higher-precedence delimiter, lower ones
    # cannot match over it (e.g. the "[Total]" inside "[[Total]]").
    claimed = [False] * len(text)
    found = []  # (start, end, delimiter, inner)
    for delim, pat in _PATTERNS:
        for m in pat.finditer(text):
            s, e = m.start(), m.end()
            if any(claimed[s:e]):
                continue
            for i in range(s, e):
                claimed[i] = True
            found.append((s, e, delim, m.group(1)))

    found.sort(key=lambda t: t[0])

    tokens: List[Token] = []
    for s, e, delim, inner_raw in found:
        raw = text[s:e]
        inner = inner_raw.strip()
        kind = "literal" if (delim == "{{" and _is_known_literal(inner)) else "nl"
        tokens.append(
            Token(
                raw=raw,
                inner=inner,
                delimiter=delim,
                kind=kind,
                location=Location(
                    runs=_run_span(run_offsets, s, e),
                    paragraph_text=text,
                ),
            )
        )
    return tokens


def _iter_table_paragraphs(table):
    """Yield every paragraph in a table, recursing into nested tables."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def extract_placeholders(docx_path) -> List[Token]:
    """Parse all placeholder tokens out of a ``.docx`` file.

    Walks body paragraphs and tables, plus each section's header and footer
    (paragraphs and tables). Returns a flat list of :class:`Token` in document
    order.
    """
    doc = Document(str(docx_path))
    tokens: List[Token] = []

    # Body paragraphs + tables.
    for para in doc.paragraphs:
        tokens.extend(_tokens_in_paragraph(para))
    for table in doc.tables:
        for para in _iter_table_paragraphs(table):
            tokens.extend(_tokens_in_paragraph(para))

    # Headers + footers (each section).
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                tokens.extend(_tokens_in_paragraph(para))
            for table in part.tables:
                for para in _iter_table_paragraphs(table):
                    tokens.extend(_tokens_in_paragraph(para))

    return tokens


# =========================================================================== #
# XTF-2 — Batched inference + local validation.
# =========================================================================== #
# A ``Proposal`` is a plain dict (mirrors the Ask engine's recipe shape):
#   {token_index, kind, spec, name, confidence, reason}
# ``infer_specs`` produces them from one batched LLM call; ``annotate_proposals``
# validates them locally (no AI) and stamps ``status`` + a human ``reason``.

Proposal = Dict

# Kinds the inference path understands.
_KINDS = ("chart", "indicator", "summary", "table", "narrative", "metadata")

# Confidence below this is treated as low → needs_attention.
_CONFIDENCE_THRESHOLD = 0.5

# Narrative inner text → fixed report slot. The match is a loose keyword test so
# a placeholder like "Recommendations" or "key recommendations" routes correctly.
_NARRATIVE_SLOTS = ("summary_text", "observations", "recommendations")
_NARRATIVE_SLOT_KEYWORDS = {
    "summary_text": ("summary", "executive summary", "overview"),
    "observations": ("observation", "findings", "finding", "key points"),
    "recommendations": ("recommendation", "recommend", "next steps"),
}

# metadata inner text → canonical report.* field.
_METADATA_FIELDS = {
    "title": "report.title",
    "report title": "report.title",
    "period": "report.period",
    "reporting period": "report.period",
}


def _slugify(text: str) -> str:
    """A snake_case slug from arbitrary placeholder text."""
    slug = re.sub(r"[^0-9a-zA-Z]+", "_", str(text or "")).strip("_").lower()
    return slug or "item"


def infer_specs(nl_tokens: List["Token"], catalog: Dict, ai_cfg: Dict) -> List[Proposal]:
    """One batched LLM call: NL placeholders + data catalog → config-shaped Proposals.

    Builds a single prompt over ALL ``nl_tokens`` (their ``inner`` text) plus the
    Ask-engine ``catalog``, calls ``lf_client.get_prompt("template_inference", …)``
    + ``lf_client.chat(trace_name="template_inference", json_mode=True)`` exactly
    once, and returns one :data:`Proposal` per token. Returns ``[]`` on failure.
    """
    if not nl_tokens:
        return []

    provider = (ai_cfg.get("provider") or "openai").lower()
    placeholders = [
        {"token_index": i, "text": getattr(t, "inner", "")}
        for i, t in enumerate(nl_tokens)
    ]
    variables = {
        "placeholders": json.dumps(placeholders, ensure_ascii=False),
        "catalog": json.dumps(catalog, ensure_ascii=False),
        "kinds": ", ".join(_KINDS),
        "chart_types": ask_engine._CHART_TYPES_BLOCK,
        "indicator_stats": ask_engine._INDICATOR_STATS_BLOCK,
        "language": ai_cfg.get("language") or "English",
    }
    try:
        messages, _config = lf_client.get_prompt("template_inference", variables)
        raw = lf_client.chat(
            messages,
            model=ai_cfg.get("model", "gpt-4o"),
            provider=provider,
            api_key=ai_cfg.get("api_key", ""),
            max_tokens=max(int(ai_cfg.get("max_tokens", 1500)), 2000),
            trace_name="template_inference",
            base_url=ai_cfg.get("base_url"),
            json_mode=True,
            output_schema=_config.get("output_schema"),
        )
    except Exception as e:  # noqa: BLE001
        log.warning(f"template_inference: infer_specs failed: {e}")
        return []

    data = ask_engine._loads_lenient(raw)
    items = (data or {}).get("proposals")
    if not isinstance(items, list):
        return []

    proposals: List[Proposal] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        spec = it.get("spec")
        proposals.append({
            "token_index": it.get("token_index"),
            "kind": it.get("kind", "chart"),
            "spec": dict(spec) if isinstance(spec, dict) else {},
            "name": it.get("name") or _slugify(it.get("kind", "item")),
            "confidence": float(it.get("confidence", 0.0) or 0.0),
            "reason": it.get("reason", ""),
        })
    return proposals


def _validate_data_proposal(kind: str, spec: Dict, profile: Dict) -> "tuple[bool, str]":
    """Local validation for a data proposal (chart/indicator/summary/table).

    Reuses ``ask_engine.validate_recipe`` / ``CHART_REQS`` / ``INDICATOR_STATS``.
    A summary references columns via ``questions`` like a chart/table but is not a
    chart type, so its columns are checked against the profile directly.
    """
    if kind == "summary":
        source = spec.get("source") or "main"
        tp = profile.get(source)
        if tp is None:
            return False, f"unknown source table '{source}'"
        roles = {c["name"] for c in tp.get("columns", [])}
        cols = list(spec.get("questions") or [])
        if not cols:
            return False, "no columns specified"
        for c in cols:
            if c not in roles:
                return False, f"column '{c}' not found in '{source}'"
        return True, ""
    recipe = {**spec, "kind": kind}
    return ask_engine.validate_recipe(recipe, profile)


def _route_narrative(proposal: Proposal) -> Proposal:
    """Map a narrative proposal to a fixed report slot or an AI summary entry."""
    spec = dict(proposal.get("spec") or {})
    text = " ".join(str(x) for x in (proposal.get("name", ""), proposal.get("reason", ""),
                                     spec.get("prompt", ""))).lower()
    for slot in _NARRATIVE_SLOTS:
        keywords = _NARRATIVE_SLOT_KEYWORDS[slot]
        if any(kw in text for kw in keywords):
            proposal["name"] = slot
            proposal["spec"] = {}
            proposal["status"] = "ok"
            proposal["reason"] = f"narrative → {slot}"
            return proposal
    # Free-form narrative → a summaries entry with stat 'ai'; carry the placeholder
    # text as the prompt so the narrator fills it at build time.
    prompt_text = spec.get("prompt") or proposal.get("reason") or proposal.get("name", "")
    proposal["spec"] = {
        "name": proposal.get("name") or "narrative",
        "stat": "ai",
        "prompt": prompt_text,
    }
    proposal["status"] = "ok"
    proposal["reason"] = "free-form narrative → AI summary"
    return proposal


def _route_metadata(proposal: Proposal) -> Proposal:
    """Map a metadata proposal to a canonical ``report.*`` field."""
    key = str(proposal.get("name", "")).lower()
    field_name = _METADATA_FIELDS.get(key, "report.title")
    proposal["name"] = field_name
    proposal["status"] = "ok"
    proposal["reason"] = f"metadata → {field_name}"
    return proposal


def annotate_proposals(proposals: List[Proposal], profile: Dict) -> List[Proposal]:
    """Local, deterministic validation of inferred proposals (no AI).

    Stamps each proposal with ``status`` (``"ok"`` / ``"needs_attention"``) and a
    human-readable ``reason``. ``needs_attention`` is set when confidence is low,
    validation fails, or a referenced column is absent. Narrative and metadata
    kinds are routed to their fixed slots. Canonical ``name``s are deduped with a
    numeric suffix on collision.
    """
    profile = profile or {}
    out: List[Proposal] = []
    for p in proposals:
        ann = dict(p)
        ann.setdefault("spec", {})
        kind = ann.get("kind", "chart")

        if kind == "narrative":
            ann = _route_narrative(ann)
        elif kind == "metadata":
            ann = _route_metadata(ann)
        else:
            ok, reason = _validate_data_proposal(kind, ann.get("spec") or {}, profile)
            if not ok:
                ann["status"] = "needs_attention"
                ann["reason"] = reason
            elif float(ann.get("confidence", 0.0) or 0.0) < _CONFIDENCE_THRESHOLD:
                ann["status"] = "needs_attention"
                ann["reason"] = (
                    f"low confidence ({ann.get('confidence')}); please review"
                )
            else:
                ann["status"] = "ok"
                ann["reason"] = ann.get("reason") or "validated"
        out.append(ann)

    # Dedupe canonical names (suffix on collision).
    seen: set = set()
    for ann in out:
        base = ann.get("name") or ann.get("kind") or "item"
        name = base
        i = 2
        while name in seen:
            name = f"{base}_{i}"
            i += 1
        seen.add(name)
        ann["name"] = name
    return out


# =========================================================================== #
# XTF-22 — Deterministic auto-modeling resolver for cross-table columns.
# =========================================================================== #
# ``resolve_sources(proposals, profile)`` runs AFTER ``infer_specs`` and BEFORE
# ``annotate_proposals`` — pure Python, no LLM. The inference catalog already
# carries every table (main + repeat-group bases), and ``builder._pick_df``
# already auto-selects the right table at build time; but validation defaults
# ``source`` to ``"main"`` and rejects anything not found there. This pass
# stamps the correct ``source`` (single-table case) or synthesizes a join view
# (repeat-table + main case) so the existing validation passes for resolved
# specs, leaving only genuinely ambiguous/unknown columns flagged.

# Data kinds whose columns map onto profile tables.
_DATA_KINDS = ("chart", "indicator", "summary", "table")


def _referenced_columns(spec: Dict) -> List[str]:
    """Columns a data spec references, mirroring ask_engine validation reads.

    ``questions`` (list) + ``group_by`` (str) + ``question`` (str, the indicator
    single-column field). Order-preserving, de-duped, blanks dropped.
    """
    cols: List[str] = []
    for c in (spec.get("questions") or []):
        if c:
            cols.append(c)
    for key in ("group_by", "question"):
        c = spec.get(key)
        if c:
            cols.append(c)
    seen: set = set()
    out: List[str] = []
    for c in cols:
        if c not in seen:
            seen.add(c)
            out.append(c)
    return out


def _table_columns(profile: Dict, table: str) -> set:
    """Set of column names in a profile table (empty when the table is absent)."""
    tp = profile.get(table) or {}
    return {c.get("name") for c in (tp.get("columns") or []) if isinstance(c, dict)}


def _is_auto_view_table(name: str) -> bool:
    """True for a profile table that is a previously synthesized auto-view.

    Synthesized views use the deterministic ``auto_<leaf>__<joincols>`` prefix;
    once persisted into ``config.yml`` ``views:`` they show up as profile tables.
    They must not themselves be candidate join sources (they would otherwise tie
    with the real repeat-group base) — recognising them keeps re-runs idempotent.
    """
    return str(name).startswith("auto_")


def _tables_containing(profile: Dict, column: str) -> List[str]:
    """Real (non auto-view) table names whose columns include ``column``."""
    return [t for t in profile
            if not _is_auto_view_table(t) and column in _table_columns(profile, t)]


def _pick_table(profile: Dict, candidates: set, columns: List[str]) -> "tuple[str | None, list[str]]":
    """Most-columns-match heuristic (mirrors ``builder._pick_df``).

    Among ``candidates`` (table names), return the one containing the MOST of
    ``columns``. Returns ``(table, ties)`` where ``ties`` is the list of tables
    sharing the top score when there is no unique winner (>1 table tied on a
    non-zero score), else an empty list.
    """
    scored = [(t, sum(1 for c in columns if c in _table_columns(profile, t)))
              for t in candidates]
    scored = [(t, n) for t, n in scored if n > 0]
    if not scored:
        return None, []
    best = max(n for _t, n in scored)
    top = sorted(t for t, n in scored if n == best)
    if len(top) > 1:
        return None, top
    return top[0], []


def _is_aggregated_chart(kind: str, spec: Dict, profile: Dict) -> bool:
    """True when the chart inherently aggregates a measure by a category.

    Carry ``group_by``/``question``/``agg`` onto a synthesized view only for
    these: a ``group_by`` over a quantitative ``questions`` column means the
    view should pre-aggregate (matching ``transform.build_views``)."""
    if kind not in ("chart", "table"):
        return False
    group_by = spec.get("group_by")
    questions = list(spec.get("questions") or [])
    if not group_by or not questions:
        return False
    # Treat as aggregated when the measure column is quantitative anywhere.
    measure = questions[0]
    for t in profile:
        for c in (profile.get(t, {}).get("columns") or []):
            if c.get("name") == measure and c.get("role") == "quantitative":
                return True
    return False


def _view_name(repeat_table: str, join_cols: List[str]) -> str:
    """Deterministic, collision-safe synthesized-view name.

    ``auto_<repeat_leaf>__<joincols>`` — leaf of the repeat table key + the
    joined main columns, all slugified, so the same join always yields the same
    name (idempotent re-runs)."""
    leaf = _slugify(str(repeat_table).split("/")[-1].split("__")[-1])
    joins = "_".join(_slugify(c) for c in join_cols) or "main"
    return f"auto_{leaf}__{joins}"


def resolve_sources(proposals: List[Proposal], profile: Dict) -> List[Dict]:
    """Resolve each data proposal's ``source`` deterministically (no LLM).

    Mutates proposals in place — stamping ``spec["source"]`` (single-table /
    view), attaching a synthesized ``view`` dict (join case), or setting
    ``status``/``reason`` (stuck) — and returns the list of synthesized view
    dicts (de-duped by name) for ``/api/template/apply`` to persist.
    """
    profile = profile or {}
    synthesized: List[Dict] = []
    synthesized_by_name: Dict[str, Dict] = {}

    for prop in proposals:
        kind = prop.get("kind", "chart")
        if kind not in _DATA_KINDS:
            continue
        spec = prop.get("spec") or {}
        cols = _referenced_columns(spec)
        if not cols:
            continue

        # Map each referenced column to the tables that contain it.
        located = {c: _tables_containing(profile, c) for c in cols}
        missing = [c for c, tabs in located.items() if not tabs]
        if missing:
            prop["status"] = "needs_attention"
            prop["reason"] = (
                f"column '{missing[0]}' not found in any table "
                f"({', '.join(sorted(profile))})"
            )
            continue

        in_main = [c for c in cols if "main" in located[c]]
        non_main = [c for c in cols if "main" not in located[c]]

        # All referenced columns are in main → leave source as-is.
        if not non_main:
            continue

        # Candidate non-main tables: those holding at least one non-main column.
        candidate_tables = set()
        for c in non_main:
            candidate_tables.update(located[c])
        candidate_tables.discard("main")

        if not in_main:
            # All columns live in non-main tables. Pick the single table holding
            # the most of them; a genuine tie with no majority is flagged.
            table, ties = _pick_table(profile, candidate_tables, cols)
            if table is None:
                prop["status"] = "needs_attention"
                prop["reason"] = (
                    f"column '{non_main[0]}' is ambiguous — present in "
                    f"{', '.join(ties)}; specify a source"
                )
                continue
            # Every referenced column must be reachable from the picked table.
            table_cols = _table_columns(profile, table)
            unreachable = [c for c in cols if c not in table_cols]
            if unreachable:
                prop["status"] = "needs_attention"
                prop["reason"] = (
                    f"columns span multiple tables; '{unreachable[0]}' not in "
                    f"'{table}'"
                )
                continue
            spec["source"] = table
            continue

        # Span a repeat table + main → synthesize a join view. Pick the repeat
        # table holding the most of the non-main columns.
        table, ties = _pick_table(profile, candidate_tables, non_main)
        if table is None:
            prop["status"] = "needs_attention"
            prop["reason"] = (
                f"column '{non_main[0]}' is ambiguous — present in "
                f"{', '.join(ties)}; specify a source"
            )
            continue
        # The repeat table must hold all non-main columns to be the join source.
        table_cols = _table_columns(profile, table)
        unreachable = [c for c in non_main if c not in table_cols]
        if unreachable:
            prop["status"] = "needs_attention"
            prop["reason"] = (
                f"columns span multiple tables; '{unreachable[0]}' not in "
                f"'{table}'"
            )
            continue

        name = _view_name(table, in_main)
        view = synthesized_by_name.get(name)
        if view is None:
            view = {"name": name, "source": table, "join_parent": list(in_main)}
            if _is_aggregated_chart(kind, spec, profile):
                view["group_by"] = spec.get("group_by")
                view["question"] = (spec.get("questions") or [None])[0]
                view["agg"] = spec.get("agg", "sum")
            synthesized_by_name[name] = view
            synthesized.append(view)
        spec["source"] = name
        prop["view"] = dict(view)

    return synthesized


# =========================================================================== #
# XTF-3 — Apply: persist config + resolve template.
# =========================================================================== #
# ``apply_inference(approved, cfg, template_path) -> (cfg, resolved_path)``
# writes each approved Proposal's spec into the right config section (without
# clobbering user entries) and rewrites each token's run span in the .docx to a
# single clean ``{{ canonical }}`` run — critical so chart placeholders are one
# unbroken XML run for docxtpl.

# Approved proposal kind → (config section, canonical {{ }} prefix).
_KIND_SECTION = {
    "chart": ("charts", "chart_"),
    "indicator": ("indicators", "ind_"),
    "summary": ("summaries", "summary_"),
    "table": ("tables", "table_"),
}


def _iter_all_paragraphs(doc):
    """Yield every paragraph in the same document order as ``extract_placeholders``.

    Body paragraphs + tables, then each section's header and footer. Must mirror
    :func:`extract_placeholders` exactly so flat token indices line up.
    """
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for para in _iter_table_paragraphs(table):
            yield para
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                yield para
            for table in part.tables:
                for para in _iter_table_paragraphs(table):
                    yield para


def _unique_name(base: str, taken: set) -> str:
    """A name not already in ``taken``; suffix ``_2``, ``_3`` … on collision."""
    name = base
    i = 2
    while name in taken:
        name = f"{base}_{i}"
        i += 1
    taken.add(name)
    return name


def _write_spec(cfg: Dict, section: str, spec: Dict, slug: str) -> str:
    """Append ``spec`` (with its ``name`` set to a deduped slug) into ``cfg[section]``.

    Never clobbers an existing entry; returns the final (possibly suffixed) name.
    """
    entries = cfg.setdefault(section, [])
    if not isinstance(entries, list):
        entries = []
        cfg[section] = entries
    taken = {e.get("name") for e in entries if isinstance(e, dict)}
    name = _unique_name(slug, taken)
    new_entry = dict(spec)
    new_entry["name"] = name
    entries.append(new_entry)
    return name


def _set_narrative_slot(cfg: Dict, proposal: Proposal):
    """Route a narrative proposal: fixed slot → report.*; AI summary → summaries."""
    name = proposal.get("name", "")
    spec = proposal.get("spec") or {}
    if name in _NARRATIVE_SLOTS:
        report = cfg.setdefault("report", {})
        report[name] = spec.get("prompt") or report.get(name) or ""
        return name
    slug = _slugify(spec.get("name") or name or "narrative")
    return _write_spec(cfg, "summaries", spec, slug)


def _set_metadata(cfg: Dict, proposal: Proposal):
    """Route a metadata proposal to its canonical ``report.*`` field."""
    name = str(proposal.get("name", "") or "report.title")
    section, _, field_name = name.partition(".")
    field_name = field_name or "title"
    target = cfg.setdefault(section or "report", {})
    spec = proposal.get("spec") or {}
    target[field_name] = spec.get("value") or spec.get(field_name) or target.get(field_name) or ""
    return name


def _resolved_path(template_path: str) -> str:
    """A new path next to the upload for the resolved template."""
    base, ext = os.path.splitext(str(template_path))
    return f"{base}.resolved{ext or '.docx'}"


def apply_inference(
    approved: List[Proposal], cfg: Dict, template_path
) -> Tuple[Dict, str]:
    """Persist approved specs into ``cfg`` and resolve the template.

    For each approved :data:`Proposal` (matched to an extracted token by
    ``token_index``), the spec is appended to its config section using the
    established list-of-dicts shape (canonical slug as ``name``; existing entries
    are never clobbered, colliding names are suffixed). The token's run span is
    then replaced by a single clean ``{{ <prefix>_<slug> }}`` run with the other
    runs in the span cleared. The resolved ``.docx`` is saved as a NEW file (the
    original upload is preserved) and its path returned.
    """
    cfg = cfg if isinstance(cfg, dict) else {}
    template_path = str(template_path)

    # Re-extract tokens so approved.token_index lines up with the .docx.
    tokens = extract_placeholders(template_path)

    # Open a fresh document to rewrite (keeps the original upload untouched).
    doc = Document(template_path)
    paragraphs = list(_iter_all_paragraphs(doc))

    # Map each paragraph (by identity) to its run-relative tokens so we can match
    # an extracted token back to a concrete paragraph + run span. We rebuild the
    # flat token list in the same order to align indices.
    flat_index = 0
    token_locations = {}  # token_index -> (paragraph, [run indices])
    for para in paragraphs:
        for tok in _tokens_in_paragraph(para):
            token_locations[flat_index] = (para, list(tok.location.runs))
            flat_index += 1

    # Persist each approved spec into its config section, then rewrite its token.
    for prop in approved:
        kind = prop.get("kind", "chart")
        spec = dict(prop.get("spec") or {})

        if kind == "narrative":
            _set_narrative_slot(cfg, prop)
            canonical = None
        elif kind == "metadata":
            _set_metadata(cfg, prop)
            canonical = None
        else:
            section, prefix = _KIND_SECTION.get(kind, _KIND_SECTION["chart"])
            base_slug = _slugify(prop.get("name") or spec.get("name") or kind)
            final_name = _write_spec(cfg, section, spec, base_slug)
            canonical = f"{{{{ {prefix}{final_name} }}}}"

        if canonical is None:
            continue

        loc = token_locations.get(prop.get("token_index"))
        if not loc:
            continue
        para, run_indices = loc
        _rewrite_run_span(para, run_indices, canonical)

    resolved = _resolved_path(template_path)
    doc.save(resolved)
    return cfg, resolved


def _rewrite_run_span(paragraph, run_indices: List[int], text: str):
    """Replace the token's run span with a single clean run carrying ``text``.

    The first run in the span gets ``text``; the remaining spanned runs are
    cleared. So the placeholder ends up as exactly one non-empty XML run.
    """
    if not run_indices:
        return
    runs = paragraph.runs
    first = run_indices[0]
    if first >= len(runs):
        return
    runs[first].text = text
    for idx in run_indices[1:]:
        if idx < len(runs):
            runs[idx].text = ""
