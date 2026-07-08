import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import pandas as pd
from docx.shared import Inches
from docxtpl import DocxTemplate, InlineImage
from jinja2.sandbox import SandboxedEnvironment
from src.data.transform import load_processed_data, apply_local_scope, aggregate_repeat, join_repeat_to_main, apply_computed_columns, build_views
from src.reports.charts import generate_chart, build_bullet_list_text, CHART_DIR
from src.reports.indicators import compute_indicators, build_traffic_light_table, build_equity_charts
from src.reports.narrator import generate_narrative
from src.reports.summaries import compute_summaries
from src.utils.provenance import build_provenance, data_mtime
from src.utils.config import get_palette
from src.reports.logframe import build_logframe
from src.reports.data_quality import build_data_quality

log = logging.getLogger(__name__)

import re as _re
from docx import Document as _Document

def _strip_residual_brackets(path: Path) -> None:
    """Remove [[ and ]] delimiters left by unresolved Express Fill tokens.

    docxtpl only resolves {{ }} Jinja2 tags; any [[token]] that was not matched
    by the Express Fill pipeline survives into the saved .docx verbatim.  This
    pass reopens the file, strips the bracket pair from every run's text while
    preserving the inner content, and saves in-place.
    """
    _OPEN = "[["; _CLOSE = "]]"

    def _clean_runs(para):
        runs = para.runs
        if not runs:
            return
        joined = "".join(run.text for run in runs)
        if _OPEN not in joined and _CLOSE not in joined:
            return
        # Join the paragraph's run text before pattern-matching so a delimiter
        # split across runs — even mid-character (e.g. "[" | "[NOM]" | "]") — is
        # matched as a whole. The current per-run replace only catches a
        # delimiter that is intact within a single run.
        cleaned = joined.replace(_OPEN, "").replace(_CLOSE, "")
        if cleaned == joined:
            return
        # Redistribute the cleaned text: keep it in the first run (preserving its
        # formatting) and blank the remaining runs. Only the surviving delimiter
        # tokens are removed; all other text is preserved verbatim.
        runs[0].text = cleaned
        for run in runs[1:]:
            run.text = ""

    def _walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _clean_runs(para)
                for nested in cell.tables:
                    _walk_table(nested)

    doc = _Document(str(path))
    for para in doc.paragraphs:
        _clean_runs(para)
    for table in doc.tables:
        _walk_table(table)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf:
                for para in hf.paragraphs:
                    _clean_runs(para)
    doc.save(str(path))


def sandboxed_jinja_env() -> SandboxedEnvironment:
    """Jinja2 environment used to render Word templates.

    Templates are user-uploaded (.docx via the web UI), so their content is
    untrusted. A SandboxedEnvironment blocks access to dunder/internal attributes
    (``__class__``, ``__globals__`` …) and unsafe callables, neutralising the
    SSTI gadget chains that would otherwise turn template rendering into remote
    code execution, while still allowing the legitimate ``{{ var }}`` / ``{% for %}``
    constructs the report templates rely on.
    """
    return SandboxedEnvironment()


def _table_display_frame(df, questions, opts) -> "pd.DataFrame":
    """Shape a resolved DataFrame into the table's display frame.

    Mirrors the legacy `table` chart breakdown: a value-count of the first
    question column into ``<label> | Count | Percent`` (top-N), so the rows and
    columns are identical to before — now destined for native cells rather than
    a PNG. If no question column resolves, the frame is passed through as-is.
    """
    from src.reports.charts import _top, _t  # local import: avoid a cycle at module load

    cols = [q for q in (questions or []) if q in df.columns]
    if not cols:
        return df
    c = cols[0]
    top_n = opts.get("top_n", 15) if isinstance(opts, dict) else 15
    counts = _top(df[c].dropna(), top_n).reset_index()
    count_col = _t(opts, "Count")
    pct_col = _t(opts, "Percent")
    counts.columns = [c, count_col]
    total = counts[count_col].sum()
    counts[pct_col] = (
        (counts[count_col] / total * 100).round(1).astype(str) + "%"
        if total
        else "0%"
    )
    return counts


def _apply_manual_table_borders(table) -> None:
    """Add a single, thin, black grid to a table via a direct ``w:tblBorders``
    element — ``add_table()`` (and a style-less template) leaves a table with no
    borders, so this is the fallback when no ``Table Grid`` style is available."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _insert_table_before_paragraph(doc, paragraph, display_df) -> None:
    """Build a native ``w:tbl`` for ``display_df`` and insert it into ``doc``
    immediately before ``paragraph``.

    One header row (column labels) + one row per record, populated as text
    cells. Styled ``Table Grid`` when the document defines that style, else a
    manual ``w:tblBorders`` grid is applied so borders are always visible
    (``add_table()`` adds none by default).
    """
    cols = [str(c) for c in display_df.columns]
    n_cols = max(len(cols), 1)

    # Build the table at the end of the document, then move it before the target
    # paragraph (python-docx exposes no direct "insert table at position" API).
    table = doc.add_table(rows=1, cols=n_cols)

    # Prefer the document's 'Table Grid' style for borders; fall back to manual.
    styled = False
    try:
        style_names = {s.name for s in doc.styles}
    except Exception:
        style_names = set()
    if "Table Grid" in style_names:
        try:
            table.style = "Table Grid"
            styled = True
        except Exception:
            styled = False
    if not styled:
        _apply_manual_table_borders(table)

    # Header row.
    header_cells = table.rows[0].cells
    for i, label in enumerate(cols):
        header_cells[i].text = label

    # One row per record.
    for _, record in display_df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(display_df.columns):
            row_cells[i].text = "" if pd.isna(record[col]) else str(record[col])

    # Move the freshly-appended table to just before the placeholder paragraph.
    paragraph._element.addprevious(table._tbl)


def _pick_df(
    questions: List[str],
    main_df: "pd.DataFrame",
    repeat_tables: Dict,
    source: Optional[str] = None,
) -> "pd.DataFrame":
    """Return the DataFrame that contains the requested question columns.

    When *source* is provided it takes precedence:
      - "main"      → always return main_df
      - other str   → look up repeat_tables[source], fall back to main_df if missing

    Without a source, falls back to the auto-heuristic: pick the DataFrame that
    contains the most of the requested columns (main first, then repeat tables).
    """
    if source:
        if source == "main":
            return main_df
        rdf = repeat_tables.get(source)
        if rdf is not None:
            return rdf
        # AI-suggested charts often use the leaf repeat-group name; repeat_tables
        # keys use the full slash-replaced path. Try a suffix match before falling back.
        matches = [k for k in repeat_tables if k.endswith(f"_{source}") or k == source]
        if len(matches) == 1:
            return repeat_tables[matches[0]]
        log.warning(f"source '{source}' not found in repeat_tables — falling back to auto-select")

    if not repeat_tables:
        return main_df
    best_df = main_df
    best_hits = sum(1 for q in questions if q in main_df.columns)
    for rdf in repeat_tables.values():
        hits = sum(1 for q in questions if q in rdf.columns)
        if hits > best_hits:
            best_hits = hits
            best_df = rdf
    return best_df


def _resolve_chart_df(
    chart_cfg: Dict,
    df: "pd.DataFrame",
    repeat_tables: Dict,
    per_form: Optional[Dict] = None,
) -> "tuple":
    """Return (chart_df, chart_repeat_tables) for a single chart config.

    When the chart carries a ``form:`` key the DataFrame is sourced from the
    matching ``per_form`` bundle instead of the default ``df``.  An unknown alias
    (or a ``form:`` key with no ``per_form`` bundle at all) raises ``ValueError``
    rather than silently rendering against wrong data (AC4).
    """
    form_alias = chart_cfg.get("form")
    if form_alias:
        if not per_form or form_alias not in per_form:
            raise ValueError(
                f"chart '{chart_cfg.get('name', '?')}' references unknown form alias "
                f"'{form_alias}'. Available aliases: {list(per_form or {})}"
            )
        bundle = per_form[form_alias]
        return bundle["df"], bundle.get("repeat_tables", {})
    return df, repeat_tables


def _filter_repeat_tables_by_split(
    df: "pd.DataFrame",
    repeat_tables: Dict,
    split_col: str,
    val,
) -> Dict:
    """When splitting reports, keep only repeat rows whose parent submission is in the split."""
    if not repeat_tables:
        return {}
    id_col = next((c for c in ("_id", "_index", "_uuid") if c in df.columns), None)
    if not id_col:
        return repeat_tables
    surviving_ids = set(df[df[split_col] == val][id_col])
    return {
        name: rdf[rdf["_parent_index"].isin(surviving_ids)]
        for name, rdf in repeat_tables.items()
    }


class ReportBuilder:
    def __init__(self, cfg: Dict, *, strict: bool = False):
        self.cfg = cfg
        self.report_cfg = cfg.get("report", {})
        self.charts_cfg: List[Dict] = cfg.get("charts", [])
        # Equity / inclusion lens (ME-1): a single `equity_dimensions:` line
        # expands into one disaggregation chart per indicator x dimension,
        # appended to the rendered chart set without disturbing explicit charts.
        self.charts_cfg = self.charts_cfg + build_equity_charts(cfg)
        self.tables_cfg: List[Dict] = cfg.get("tables", [])
        self.lists_cfg: List[Dict] = cfg.get("lists", [])
        self.strict = strict

    def build(self, sample_size: Optional[int] = None, split_by: Optional[str] = None, random_sample: bool = False, split_sample: Optional[int] = None, session: Optional[str] = None, period: Optional[str] = None, compare: Optional[List[str]] = None) -> List[Path]:
        from src.utils.periods import parse_period_arg
        resolved_period = parse_period_arg(self.cfg, period)
        df, repeat_tables = load_processed_data(self.cfg, sample_size=sample_size, random_sample=random_sample, session=session, period=resolved_period)
        df = apply_computed_columns(df, self.cfg, repeat_tables, strict=self.strict)
        # Clear prior report outputs so each build's result is exactly the
        # current run's set (default, split-by, and --split-sample alike).
        # Only *.docx reports are removed; other files + the charts dir are
        # untouched. Done once here, not in _render (which runs per split value).
        out_dir = Path(self.report_cfg.get("output_dir", "reports"))
        out_dir.mkdir(parents=True, exist_ok=True)
        for stale in out_dir.glob("*.docx"):
            stale.unlink()
        split_col = split_by or self.report_cfg.get("split_by")
        if split_col:
            if split_col not in df.columns:
                log.warning(f"split_by column '{split_col}' not found — building single report")
                return [self._render(df, repeat_tables, suffix="", compare=compare, split_by=None)]
            unique_vals = sorted(df[split_col].dropna().unique())
            if split_sample and split_sample < len(unique_vals):
                log.info(f"Split sample: limiting to first {split_sample} of {len(unique_vals)} value(s)")
                unique_vals = unique_vals[:split_sample]
            log.info(f"Split by '{split_col}': {len(unique_vals)} value(s) → {len(unique_vals)} report(s)")
            paths = []
            for val in unique_vals:
                safe = str(val).replace("/", "_").replace(" ", "_")
                # Filter repeat tables to rows whose parent submission survived the split
                filtered_repeats = _filter_repeat_tables_by_split(df, repeat_tables, split_col, val)
                paths.append(self._render(df[df[split_col] == val], filtered_repeats, suffix=f"_{safe}", split_value=str(val), compare=compare, split_by=split_col))
            return paths
        suffix = f"_sample{sample_size}" if sample_size else ""
        return [self._render(df, repeat_tables, suffix=suffix, compare=compare, split_by=None)]

    def _render(self, df: "pd.DataFrame", repeat_tables: Dict, suffix: str = "", split_value: Optional[str] = None, compare: Optional[List[str]] = None, split_by: Optional[str] = None) -> Path:
        template_path = Path(self.report_cfg.get("template","templates/report_template.docx"))
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}\nRun generate-template or see TEMPLATE_GUIDE.md")
        tpl = DocxTemplate(template_path)

        # Apply PII redaction (consent gating + column rules) before any
        # downstream rendering — charts, indicators, summaries, views all see
        # the redacted data.  No-op when no pii: block is configured.
        from src.utils.pii import apply_pii
        df, repeat_tables = apply_pii(df, repeat_tables, self.cfg)

        # Compute named virtual views and inject into repeat_tables so all
        # consumers (charts, summaries, indicators) can reference them by name.
        # Called here so views see the already split-filtered df and repeat_tables.
        views = build_views(self.cfg, df, repeat_tables, strict=self.strict)
        if views:
            repeat_tables = {**repeat_tables, **views}

        stats_table = self._stats_table(df)

        from src.utils.periods import all_periods, baseline_period
        registry = all_periods(self.cfg)
        per_period = None
        if registry and len(registry) > 1:
            base = baseline_period(self.cfg)
            base_slug = base["slug"] if base else None
            per_period = {}
            for entry in registry:
                try:
                    p_df, p_repeats = load_processed_data(self.cfg, period=entry)
                    per_period[entry["slug"]] = {
                        "df": p_df,
                        "repeat_tables": p_repeats,
                        "label": entry["label"],
                        "is_baseline": entry["slug"] == base_slug,
                    }
                except FileNotFoundError:
                    continue
        if compare:
            from src.utils.periods import slugify, all_periods, baseline_period
            c_registry = all_periods(self.cfg)
            label_to_slug = {e["label"]: e["slug"] for e in c_registry}
            slugs = [label_to_slug.get(lbl, slugify(lbl)) for lbl in compare]
            if per_period:
                # Filter per_period to just the requested slugs, preserving the compare order.
                filtered = {}
                for s in slugs:
                    if s in per_period:
                        filtered[s] = per_period[s]
                per_period = filtered or None
            if per_period:
                base = baseline_period(self.cfg)
                base_slug = base["slug"] if base else slugs[0]
                for s in per_period:
                    per_period[s]["is_baseline"] = (s == base_slug)

        self._last_per_period = per_period   # exposed for Task 18 (chart payload)

        # ME-7: build per_form bundle so charts can use `form: <alias>` selector.
        self._per_form: Optional[Dict] = None
        forms_list = (self.cfg.get("api") or {}).get("forms")
        if forms_list:
            self._per_form = {}
            for form_entry in forms_list:
                alias = form_entry.get("alias")
                if not alias:
                    continue
                alias_cfg = {**self.cfg, "form": {"alias": alias, "uid": form_entry.get("uid", "")}}
                try:
                    f_df, f_repeats = load_processed_data(alias_cfg)
                    self._per_form[alias] = {"df": f_df, "repeat_tables": f_repeats}
                except FileNotFoundError:
                    log.warning(f"No data found for form alias '{alias}' — skipping per_form entry")

        indicators  = compute_indicators(
            self.cfg.get("indicators", []), df, repeat_tables, per_period=per_period,
            per_form=self._per_form,
        )
        logframe = build_logframe(self.cfg, indicators)
        # ME-2: red/amber/green progress table + below-threshold flags.
        indicators_cfg = self.cfg.get("indicators", [])
        traffic_light = build_traffic_light_table(indicators_cfg, indicators)
        flagged_indicators = [
            r for r in traffic_light["rows"] if r["status"] in ("warning", "critical")
        ]
        data_quality = build_data_quality(self.cfg, df, repeat_tables)
        summaries   = compute_summaries(
            self.cfg.get("summaries", []), df, self.cfg.get("ai"),
            repeat_tables,
        )

        narrative = generate_narrative(
            ai_cfg        = self.cfg.get("ai"),
            report_cfg    = self.report_cfg,
            df            = df,
            stats_table   = stats_table,
            indicators    = indicators,
            charts_cfg    = self.charts_cfg,
            summaries     = summaries,
            split_value   = split_value,
            questions_cfg = self.cfg.get("questions"),
        )

        provenance = build_provenance(
            self.cfg,
            df,
            data_downloaded_at=data_mtime(
                Path(self.cfg.get("export", {}).get("output_dir", "data/processed")),
                self.cfg.get("form", {}).get("alias", "form"),
            ),
            compared_periods=compare,
        )

        now = datetime.today()
        # Native-table sentinels (tables: recipes AND MNT-33 `table`-type charts)
        # accumulate here across _generate_charts + _generate_tables, then get
        # swapped by _insert_native_tables after render. Reset once per build so
        # a split_by iteration never inherits the previous slice's tables.
        self._pending_tables = {}
        context = {
            "report_title":  self.report_cfg.get("title", "Report"),
            "period":        self.report_cfg.get("period", datetime.today().strftime("%B %Y")),
            "n_submissions": len(df),
            "generated_at":  now.strftime("%d/%m/%Y %H:%M"),
            "year":          now.strftime("%Y"),
            "month":         now.strftime("%m"),
            "day":           now.strftime("%d"),
            "split_value":   split_value or "",
            "split_by":      split_by or "",
            "provenance":    provenance,
            "logframe":      logframe,
            "traffic_light": traffic_light,
            "flagged_indicators": flagged_indicators,
            "data_quality":  data_quality,
            **narrative,
            "stats_table":   stats_table,
            **indicators,
            **summaries,
            **self._generate_charts(tpl, df, repeat_tables, per_form=self._per_form),
            **self._generate_tables(tpl, df, repeat_tables),
            **self._generate_lists(tpl, df, repeat_tables),
        }
        tpl.render(context, jinja_env=sandboxed_jinja_env())
        # Swap {{ table_<name> }} sentinels for native python-docx tables.
        self._insert_native_tables(tpl)
        out_dir = Path(self.report_cfg.get("output_dir","reports"))
        out_dir.mkdir(parents=True, exist_ok=True)
        alias = self.cfg.get("form",{}).get("alias","form")
        out_path = out_dir / f"{alias}_report{suffix}_{datetime.today().strftime('%Y%m%d')}.docx"
        tpl.save(out_path)
        _strip_residual_brackets(out_path)
        log.info(f"Report saved → {out_path}")
        return out_path

    def _generate_charts(self, tpl, df, repeat_tables: Dict, per_form: Optional[Dict] = None):
        CHART_DIR.mkdir(parents=True, exist_ok=True)
        _language = (self.cfg.get("ai") or {}).get("language") or "English"
        key_to_label = {
            q["kobo_key"]: q.get("export_label") or q.get("label") or q["kobo_key"]
            for q in self.cfg.get("questions", [])
        }
        images = {}
        for c in self.charts_cfg:
            name = c.get("name")
            if not name:
                continue

            # ME-7: resolve the base DataFrame for this chart (form: selector).
            chart_base_df, chart_base_repeats = _resolve_chart_df(c, df, repeat_tables, per_form)

            # Resolve question names (kobo_key → export_label)
            resolved_questions = [
                key_to_label.get(q, q) if q not in chart_base_df.columns else q
                for q in c.get("questions", [])
            ]

            # group_by: inject as second question and auto-upgrade chart type
            group_by = c.get("group_by")
            if group_by:
                group_col = key_to_label.get(group_by, group_by)
                if group_col not in resolved_questions:
                    resolved_questions = resolved_questions[:1] + [group_col] + resolved_questions[1:]
                chart_type = c.get("type", "bar")
                _AUTO_UPGRADE = {"bar": "grouped_bar", "horizontal_bar": "grouped_bar", "histogram": "box_plot"}
                upgraded_type = _AUTO_UPGRADE.get(chart_type, chart_type)
                resolved = {**c, "questions": resolved_questions, "type": upgraded_type}
            else:
                resolved = {**c, "questions": resolved_questions}

            # 1. Select explicit source or auto-pick (within the resolved base tables)
            source = c.get("source")
            chart_df = _pick_df(resolved_questions, chart_base_df, chart_base_repeats, source=source)

            # 1b. Join parent fields into repeat table if requested
            join_parent = c.get("join_parent")
            if join_parent and source and source != "main":
                chart_df = join_repeat_to_main(chart_df, chart_base_df, join_parent)

            # 2. Apply per-chart filter and sample
            filter_expr = c.get("filter")
            sample_n = c.get("sample")
            if filter_expr or sample_n:
                chart_df = apply_local_scope(
                    chart_df, {}, filter_expr=filter_expr, sample_n=sample_n
                )

            # 3. Apply repeat-group aggregation if requested
            agg_spec = c.get("aggregate")
            if agg_spec and source and source != "main":
                chart_df = aggregate_repeat(chart_df, agg_spec)

            # Enrich period_bar / period_line with per-period computed values
            if c.get("type") in ("period_bar", "period_line") and getattr(self, "_last_per_period", None):
                from src.reports.indicators import compute_indicators
                metric = (c.get("options", {}) or {}).get("metric") or c.get("metric")
                periods_payload = []
                if metric:
                    for slug, bundle in self._last_per_period.items():
                        ind_cfg = {
                            "name":     metric,
                            "stat":     c.get("stat", "count"),
                            "question": c.get("question"),
                        }
                        try:
                            result = compute_indicators([ind_cfg], bundle["df"], bundle.get("repeat_tables", {}))
                            value  = result.get(f"ind_{metric}", "0")
                        except Exception:
                            value  = "0"
                        periods_payload.append({
                            "slug":  slug,
                            "label": bundle["label"],
                            "value": value,
                        })
                # Inject the payload into a COPY of resolved so the original config isn't mutated.
                enriched_opts = {**(resolved.get("options", {}) or {}), "periods": periods_payload}
                resolved = {**resolved, "options": enriched_opts}

            # bullet_list is a text-injection render type — it bypasses the
            # matplotlib/InlineImage pipeline entirely and fills a
            # {{ list_<name> }} text placeholder instead of {{ chart_<name> }}.
            if resolved.get("type") == "bullet_list":
                images[f"list_{name}"] = build_bullet_list_text(
                    chart_df, resolved_questions, resolved.get("options") or {})
                continue

            # MNT-33: the legacy `table` CHART TYPE renders as a NATIVE Word table,
            # not a flattened PNG — the same w:tbl path the tables: section uses
            # (MNT-30). A PNG table is unsearchable, non-editable, and inaccessible.
            # Route it through the shared native-table sentinel: fill the chart's
            # own {{ chart_<name> }} placeholder with a sentinel that
            # _insert_native_tables swaps for a real table after tpl.render. This
            # bridges legacy `charts: [{type: table}]` configs with no migration.
            if resolved.get("type") == "table":
                display_df = _table_display_frame(
                    chart_df, resolved_questions, resolved.get("options") or {})
                self._pending_tables = getattr(self, "_pending_tables", None) or {}
                sentinel = f"@@DBNATIVE_TABLE::chart::{name}::@@"
                self._pending_tables[sentinel] = display_df
                images[f"chart_{name}"] = sentinel
                continue

            png = generate_chart(resolved, chart_df, language=_language, palette=get_palette(self.cfg))
            width = Inches(c.get("options", {}).get("width_inches", 5.5))
            images[f"chart_{name}"] = InlineImage(tpl, str(png), width=width) if png and png.exists() else ""
        return images

    def _generate_tables(self, tpl, df, repeat_tables: Dict):
        """Resolve each cfg['tables'] recipe into a NATIVE python-docx table.

        A table shares a chart's source/filter/aggregate handling, but instead of
        a flattened PNG it becomes a real ``w:tbl`` — selectable, editable,
        accessible text with visible borders. Charts (bar/pie/line/…) stay as
        InlineImage PNGs; only tables are native.

        The ``{{ table_<name> }}`` placeholder is rendered by docxtpl to a unique
        text sentinel; ``_insert_native_tables`` (run after ``tpl.render``) then
        swaps each sentinel paragraph for the built native table. This keeps the
        existing plain ``{{ }}`` placeholder contract intact (docxtpl only unwraps
        the enclosing paragraph for the ``{{p }}`` subdoc syntax, which
        ``generate-template`` does not emit).
        """
        key_to_label = {
            q["kobo_key"]: q.get("export_label") or q.get("label") or q["kobo_key"]
            for q in self.cfg.get("questions", [])
        }
        # Pending (sentinel → display frame) swaps applied after render. Preserve
        # any sentinels already registered by _generate_charts (MNT-33 table-type
        # charts run first in the context dict) rather than wiping them.
        self._pending_tables = getattr(self, "_pending_tables", None) or {}
        sentinels = {}
        for t in self.tables_cfg:
            name = t.get("name")
            if not name:
                continue

            # Resolve question names (kobo_key → export_label), same as charts.
            resolved_questions = [
                key_to_label.get(q, q) if q not in df.columns else q
                for q in t.get("questions", [])
            ]

            # 1. Select explicit source or auto-pick.
            source = t.get("source")
            table_df = _pick_df(resolved_questions, df, repeat_tables, source=source)

            # 1b. Join parent fields into repeat table if requested.
            join_parent = t.get("join_parent")
            if join_parent and source and source != "main":
                table_df = join_repeat_to_main(table_df, df, join_parent)

            # 2. Apply per-table filter and sample.
            filter_expr = t.get("filter")
            sample_n = t.get("sample")
            if filter_expr or sample_n:
                table_df = apply_local_scope(
                    table_df, {}, filter_expr=filter_expr, sample_n=sample_n
                )

            # 3. Apply repeat-group aggregation if requested.
            agg_spec = t.get("aggregate")
            if agg_spec and source and source != "main":
                table_df = aggregate_repeat(table_df, agg_spec)

            # 4. Shape into the display frame (header labels + one row per record),
            #    matching the legacy `table` chart's count/percent breakdown so the
            #    resolved rows/columns are unchanged — now as native cells.
            display_df = _table_display_frame(table_df, resolved_questions, t.get("options") or {})
            sentinel = f"@@DBNATIVE_TABLE::{name}::@@"
            self._pending_tables[sentinel] = display_df
            sentinels[f"table_{name}"] = sentinel
        return sentinels

    def _insert_native_tables(self, tpl) -> None:
        """Replace each rendered table sentinel with a native ``w:tbl``.

        Walks the rendered document (body + table cells) for the paragraph whose
        text holds a sentinel emitted by ``_generate_tables`` and swaps that whole
        paragraph for a bordered python-docx table built from the display frame.
        """
        pending = getattr(self, "_pending_tables", None)
        if not pending:
            return
        doc = tpl.docx

        def _walk(paragraphs, get_tables):
            for para in list(paragraphs):
                text = para.text
                for sentinel, display_df in pending.items():
                    if sentinel in text:
                        _insert_table_before_paragraph(doc, para, display_df)
                        # Drop the placeholder paragraph entirely.
                        para._element.getparent().remove(para._element)
                        break

        # Body-level paragraphs.
        _walk(doc.paragraphs, None)
        # Paragraphs nested inside existing table cells.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _walk(cell.paragraphs, None)

    def _generate_lists(self, tpl, df, repeat_tables: Dict):
        """Render each cfg['lists'] recipe into a {{ list_<name> }} text context value.

        A list is a text-injection render type — like the bullet_list chart type it
        bypasses the matplotlib/InlineImage pipeline entirely, rendering the named
        `question` column's raw row values (via build_bullet_list_text) as a plain
        string rather than an image.
        """
        key_to_label = {
            q["kobo_key"]: q.get("export_label") or q.get("label") or q["kobo_key"]
            for q in self.cfg.get("questions", [])
        }
        images = {}
        for l in self.lists_cfg:
            name = l.get("name")
            if not name:
                continue

            question = l.get("question")
            resolved_question = key_to_label.get(question, question) if question not in df.columns else question

            source = l.get("source")
            list_df = _pick_df([resolved_question], df, repeat_tables, source=source)

            filter_expr = l.get("filter")
            if filter_expr:
                list_df = apply_local_scope(list_df, {}, filter_expr=filter_expr)

            images[f"list_{name}"] = build_bullet_list_text(
                list_df, [resolved_question], l.get("options") or {})
        return images

    def _stats_table(self, df):
        rows = []
        for q in self.cfg.get("questions",[]):
            if q.get("category") != "quantitative": continue
            label = q.get("export_label") or q.get("label") or q["kobo_key"]
            if label not in df.columns: continue
            s = pd.to_numeric(df[label], errors="coerce").dropna()
            if s.empty: continue
            rows.append({"label":label,"n":len(s),"mean":round(s.mean(),2),
                         "median":round(s.median(),2),"min":round(s.min(),2),"max":round(s.max(),2)})
        return rows
